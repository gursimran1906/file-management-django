import logging
import html
from html import escape as html_escape
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.db import transaction
from django.db.models import Q, F, OuterRef, Subquery, Max, CharField, TextField, BooleanField, Exists, Count, Sum, Case, When, Value, DateField
from django.db.models.functions import Cast, Coalesce, Greatest, Concat
from .models import WIP, Memo, NextWork, LastWork, MatterKeyDate, FileStatus, FileLocation, MatterType, PricingItem, ClientContactDetails, ClientKeyDocument, AuthorisedParties
from .models import LedgerAccountTransfers, Modifications, Invoices, RiskAssessment, PoliciesRead, OngoingMonitoring, CreditNote, CURRENT_VAT_RATE
from .models import OthersideDetails, MatterAttendanceNotes, MatterEmails, MatterLetters, PmtsSlips, Free30Mins, Free30MinsAttendees
from .models import Undertaking, Policy, PolicyVersion, Bundle, BundleSection, BundleDocument, MatterFileReview
from .forms import MemoForm, OpenFileForm, NextWorkFormWithoutFileNumber, NextWorkForm, LastWorkFormWithoutFileNumber, LastWorkForm, AttendanceNoteForm, AttendanceNoteFormHalf, LetterForm, LetterHalfForm, PolicyForm
from .forms import PmtsForm, PmtsHalfForm, LedgerAccountTransfersHalfForm, LedgerAccountTransfersForm, InvoicesForm, CreditNoteHalfForm, ClientForm, ClientKeyDocumentFormSet, MatterKeyDateForm, MatterClientKeyDocumentForm, AuthorisedPartyForm, RiskAssessmentForm, OngoingMonitoringForm, OtherSideForm
from .forms import Free30MinsForm, Free30MinsAttendeesForm, UndertakingForm, MatterFileReviewForm, PricingItemForm
from .utils import create_modification, parse_bundle_filename
from .audit import (
    audit_client_key_document_formset,
    build_form_field_changes,
    log_bundle_event,
    log_created,
    log_deleted_on_parent,
    snapshot_key_date,
    snapshot_key_document,
)
from .audit_display import build_change_items, enrich_file_logs
from django.utils import timezone
from users.models import CPDTrainingLog, CustomUser, HolidayRecord, SicknessRecord
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.core.paginator import Paginator
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
import json
import threading
from django.core.cache import cache
from weasyprint import HTML
from django.utils.safestring import mark_safe
from django.contrib.contenttypes.models import ContentType
import csv
import calendar as calendar_module
import holidays
from datetime import date, datetime, timedelta, time
from dateutil.relativedelta import relativedelta
import copy
from django.forms.models import model_to_dict
from decimal import Decimal, InvalidOperation
import ast
from .serializers import InvoicesSerializer
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from django.contrib.auth.decorators import login_required
from django.http import FileResponse, Http404
from django.conf import settings
from django.utils.dateparse import parse_date
import os
import PyPDF2
import zipfile
from io import BytesIO
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from backend.sharepoint.bundle_cache import BundleTempCache
from backend.sharepoint.client import SharePointClientError

logger = logging.getLogger('backend')
CURRENT_VAT_RATE_PERCENT = int(CURRENT_VAT_RATE * 100)

MATTER_FILE_REVIEW_SECTIONS = [
    {
        'title': 'Client Onboarding',
        'rows': [
            {
                'question': 'File Opening Checklist completed?',
                'answer_field': 'file_opening_checklist_completed',
                'comments_field': 'file_opening_checklist_completed_comments',
            },
            {
                'question': 'Engagement documents sent to the client and copies kept on file?',
                'bullets': [
                    'Client care letter',
                    'Terms of business (including complaints procedure and cancellation procedure)',
                    'Non-contentious Business Agreement (if applicable)',
                    'PEP',
                    'SOF',
                ],
                'answer_field': 'engagement_documents_sent_and_filed',
                'comments_field': 'engagement_documents_sent_and_filed_comments',
            },
            {
                'question': 'Accurate and up-to-date details of our charging rates and basis provided?',
                'answer_field': 'charging_rates_and_basis_provided',
                'comments_field': 'charging_rates_and_basis_provided_comments',
            },
            {
                'question': 'Initial estimate of costs provided?',
                'answer_field': 'initial_costs_estimate_provided',
                'comments_field': 'initial_costs_estimate_provided_comments',
            },
            {
                'question': "Signed Letter of Authority obtained to enable us to receive advice and instructions on the client's behalf (if applicable)?",
                'answer_field': 'letter_of_authority_obtained',
                'comments_field': 'letter_of_authority_obtained_comments',
            },
            {
                'question': 'Initial Risk Assessment completed?',
                'answer_field': 'initial_risk_assessment_completed',
                'comments_field': 'initial_risk_assessment_completed_comments',
            },
        ],
    },
    {
        'title': 'Matter Management',
        'rows': [
            {
                'question': 'All key dates recorded in shared calendar and WIP?',
                'answer_field': 'key_dates_recorded_in_calendar_and_wip',
                'comments_field': 'key_dates_recorded_in_calendar_and_wip_comments',
            },
            {
                'question': 'Key information and advice shared with the client?',
                'answer_field': 'key_information_and_advice_shared',
                'comments_field': 'key_information_and_advice_shared_comments',
            },
            {
                'question': 'Costs estimates updated as necessary?',
                'answer_field': 'costs_estimates_updated',
                'comments_field': 'costs_estimates_updated_comments',
            },
            {
                'question': 'Overall, the matter is progressing without any long unexplained periods of dormancy?',
                'answer_field': 'matter_progressing_without_dormancy',
                'comments_field': 'matter_progressing_without_dormancy_comments',
            },
            {
                'question': 'Overall, the file appears to be maintained in good order?',
                'answer_field': 'file_maintained_in_good_order',
                'comments_field': 'file_maintained_in_good_order_comments',
            },
        ],
    },
    {
        'title': 'Ongoing Monitoring',
        'rows': [
            {
                'question': 'Ongoing AML, financial crime prevention and sanctions monitoring carried out at appropriate intervals or on appropriate triggers (if applicable)?',
                'answer_field': 'ongoing_aml_sanctions_monitoring_carried_out',
                'comments_field': 'ongoing_aml_sanctions_monitoring_carried_out_comments',
            },
            {
                'question': 'Copies of all documents obtained from ongoing monitoring activities kept and correctly filed?',
                'answer_field': 'ongoing_monitoring_documents_kept_and_filed',
                'comments_field': 'ongoing_monitoring_documents_kept_and_filed_comments',
            },
            {
                'question': 'Further conflict checks and consideration carried out appropriately, if parties have changed?',
                'answer_field': 'further_conflict_checks_completed',
                'comments_field': 'further_conflict_checks_completed_comments',
            },
        ],
    },
    {
        'title': 'Finance, Costs And Accounting',
        'rows': [
            {
                'question': 'Money on account (as appropriate to the matter) has been requested at appropriate times and received before significant work undertaken?',
                'answer_field': 'money_on_account_requested_and_received',
                'comments_field': 'money_on_account_requested_and_received_comments',
            },
            {
                'question': 'All disbursements paid in a timely manner?',
                'answer_field': 'disbursements_paid_timely',
                'comments_field': 'disbursements_paid_timely_comments',
            },
            {
                'question': 'Costs and disbursements billed at appropriate intervals?',
                'answer_field': 'costs_and_disbursements_billed_timely',
                'comments_field': 'costs_and_disbursements_billed_timely_comments',
            },
            {
                'question': 'Overdue invoices?',
                'answer_field': 'overdue_invoices',
                'comments_field': 'overdue_invoices_comments',
            },
        ],
    },
    {
        'title': 'Legal Advice And Instructions',
        'rows': [
            {
                'question': 'Appropriate advice given to the client on all substantive issues to date?',
                'answer_field': 'appropriate_advice_given',
                'comments_field': 'appropriate_advice_given_comments',
            },
            {
                'question': 'Matter proceeding within the scope set out in the client care letter? If not, why?',
                'answer_field': 'matter_within_client_care_scope',
                'comments_field': 'matter_within_client_care_scope_comments',
            },
        ],
    },
    {
        'title': 'Specific Risk Issues',
        'rows': [
            {
                'question': 'Undertakings given by the firm have been discharged appropriately and on time, or released in writing (if not ongoing)?',
                'answer_field': 'undertakings_discharged_or_released',
                'comments_field': 'undertakings_discharged_or_released_comments',
            },
            {
                'question': 'Have any complaints been raised by the client? If so, has the firms complaints procedure been followed?',
                'answer_field': 'complaints_raised_and_process_followed',
                'comments_field': 'complaints_raised_and_process_followed_comments',
            },
            {
                'question': 'Does the file review raise any internal concerns?',
                'answer_field': 'internal_concerns_raised',
                'comments_field': 'internal_concerns_raised_comments',
            },
            {
                'question': 'Does the file review raise any concerns or suspicions about potential money laundering, sanctions breaches, or any other economic crime?',
                'answer_field': 'economic_crime_or_sanctions_concerns',
                'comments_field': 'economic_crime_or_sanctions_concerns_comments',
            },
        ],
    },
]


def build_matter_file_review_display_data(reviews):
    output = []
    for review in reviews:
        section_data = []
        for section in MATTER_FILE_REVIEW_SECTIONS:
            rows = []
            for row in section['rows']:
                rows.append({
                    'question': row['question'],
                    'bullets': row.get('bullets', []),
                    'answer': getattr(review, row['answer_field']) or '---',
                    'comments': getattr(review, row['comments_field']) or '---',
                })
            section_data.append({
                'title': section['title'],
                'rows': rows,
            })
        output.append({
            'review': review,
            'sections': section_data,
        })
    return output


def calculate_invoice_total_with_vat(invoice):
    our_costs = invoice.our_costs
    costs = ast.literal_eval(our_costs) if not isinstance(
        our_costs, list) else our_costs
    total_cost_invoice = Decimal('0')
    for cost in costs:
        total_cost_invoice += Decimal(str(cost))
    vat_inv = Decimal(str(invoice.vat or 0))
    total_cost_and_vat = total_cost_invoice + vat_inv
    return round(total_cost_invoice, 2), round(vat_inv, 2), round(total_cost_and_vat, 2)


def calculate_credit_note_breakdown(gross_amount):
    gross = round(Decimal(str(gross_amount or 0)), 2)
    denominator = Decimal('1.00') + CURRENT_VAT_RATE
    if denominator == 0:
        return gross, Decimal('0.00'), gross

    vat_amount = round((gross * CURRENT_VAT_RATE) / denominator, 2)
    net_amount = round(gross - vat_amount, 2)
    return net_amount, vat_amount, gross


def get_approved_credit_note_totals(invoice_ids):
    totals = {}
    if not invoice_ids:
        return totals

    rows = CreditNote.objects.filter(
        invoice_id__in=invoice_ids,
        status='F'
    ).values('invoice_id').annotate(total=Sum('amount'))
    for row in rows:
        totals[row['invoice_id']] = row['total'] or Decimal('0')
    return totals


def get_invoice_approved_credit_total(invoice):
    return CreditNote.objects.filter(
        invoice=invoice,
        status='F'
    ).aggregate(total=Sum('amount')).get('total') or Decimal('0')


def get_invoice_max_credit_amount(invoice, excluded_credit_note_id=None):
    max_allowed = Decimal(str(invoice.total_due_left or 0))
    if excluded_credit_note_id:
        excluded_note = CreditNote.objects.filter(
            id=excluded_credit_note_id,
            status='F',
            invoice=invoice
        ).first()
        if excluded_note:
            max_allowed += Decimal(str(excluded_note.amount))
    if max_allowed < 0:
        return Decimal('0')
    return round(max_allowed, 2)


def get_effective_invoice_due(invoice, approved_credit_total=None):
    effective_due = Decimal(str(invoice.total_due_left or 0))
    if effective_due <= 0:
        return Decimal('0')
    return round(effective_due, 2)


def get_invoice_outstanding_balance(invoice, approved_credit_total=None):
    """Amount still owed on a final invoice (handles unset total_due_left)."""
    if approved_credit_total is None:
        approved_credit_total = get_invoice_approved_credit_total(invoice)

    if invoice.total_due_left is not None:
        return get_effective_invoice_due(invoice, approved_credit_total)

    _, _, gross = calculate_invoice_total_with_vat(invoice)
    cash_allocated = Decimal('0')
    for slip in invoice.cash_allocated_slips.all():
        alloc_raw = slip.amount_allocated
        if not alloc_raw:
            continue
        try:
            data = json.loads(alloc_raw) if isinstance(
                alloc_raw, str) else alloc_raw
            if isinstance(data, dict):
                cash_allocated += Decimal(str(data.get(str(invoice.id), 0) or 0))
        except (json.JSONDecodeError, TypeError, InvalidOperation):
            continue

    balance = gross - approved_credit_total - cash_allocated
    if balance <= 0:
        return Decimal('0')
    return round(balance, 2)


def get_user_dashboard_wip_ids(user):
    """WIP ids for the dashboard, ordered by most recent user activity."""
    activity_scores = {}

    def touch(wip_id, ts):
        if not wip_id:
            return
        if wip_id not in activity_scores or ts > activity_scores[wip_id]:
            activity_scores[wip_id] = ts

    for row in NextWork.objects.filter(
        person=user, file_number__isnull=False
    ).values('file_number').annotate(latest=Max('timestamp')):
        touch(row['file_number'], row['latest'])

    for row in NextWork.objects.filter(
        created_by=user, file_number__isnull=False
    ).values('file_number').annotate(latest=Max('timestamp')):
        touch(row['file_number'], row['latest'])

    for row in LastWork.objects.filter(
        person=user, file_number__isnull=False
    ).values('file_number').annotate(latest=Max('timestamp')):
        touch(row['file_number'], row['latest'])

    for row in MatterAttendanceNotes.objects.filter(
        person_attended=user, file_number__isnull=False
    ).values('file_number').annotate(latest=Max('timestamp')):
        touch(row['file_number'], row['latest'])

    for row in MatterEmails.objects.filter(
        fee_earner=user, file_number__isnull=False
    ).values('file_number').annotate(latest=Max('time')):
        touch(row['file_number'], row['latest'])

    fee_earner_ids = WIP.objects.filter(
        fee_earner=user,
        file_status__status__in=['Open', 'To Be Closed'],
    ).values_list('id', flat=True)

    fallback_ts = timezone.make_aware(datetime.min)
    for wip_id in fee_earner_ids:
        touch(wip_id, activity_scores.get(wip_id, fallback_ts))

    if not activity_scores:
        return list(fee_earner_ids)

    return sorted(activity_scores.keys(), key=lambda wip_id: activity_scores[wip_id], reverse=True)


def get_user_dashboard_wips(user):
    wip_ids = get_user_dashboard_wip_ids(user)
    if not wip_ids:
        return WIP.objects.none()

    preserved_order = Case(
        *[When(id=wip_id, then=pos) for pos, wip_id in enumerate(wip_ids)],
        default=len(wip_ids),
    )
    return WIP.objects.filter(id__in=wip_ids).select_related(
        'client1', 'client2', 'file_status', 'fee_earner'
    ).order_by(preserved_order)


def build_dashboard_files(user, display_limit=40):
    """Recent active files with grouped overdue invoices, overdue files first."""
    all_wips = list(get_user_dashboard_wips(user))
    all_wip_ids = {wip.id for wip in all_wips}
    fee_earner_wip_ids = set(
        WIP.objects.filter(
            fee_earner=user,
            file_status__status__in=['Open', 'To Be Closed'],
        ).values_list('id', flat=True)
    )
    invoice_wip_ids = all_wip_ids | fee_earner_wip_ids

    invoice_candidates = Invoices.objects.filter(
        state='F',
        file_number_id__in=invoice_wip_ids,
    ).select_related('file_number').prefetch_related('cash_allocated_slips')

    approved_credit_totals = get_approved_credit_note_totals(
        [invoice.id for invoice in invoice_candidates]
    )

    overdue_invoices_by_wip = {}
    for invoice in invoice_candidates:
        balance = get_invoice_outstanding_balance(
            invoice, approved_credit_totals.get(invoice.id, Decimal('0'))
        )
        if balance > 0 and invoice.file_number_id:
            invoice.display_balance_due = balance
            overdue_invoices_by_wip.setdefault(
                invoice.file_number_id, []).append(invoice)

    wip_by_id = {wip.id: wip for wip in all_wips}
    display_ids = [wip.id for wip in all_wips[:display_limit]]
    for wip_id in overdue_invoices_by_wip:
        if wip_id not in display_ids:
            display_ids.append(wip_id)
            if wip_id not in wip_by_id:
                extra = WIP.objects.filter(id=wip_id).select_related(
                    'client1', 'client2', 'file_status', 'fee_earner'
                ).first()
                if extra:
                    wip_by_id[wip_id] = extra

    dashboard_files = []
    for wip_id in display_ids:
        wip = wip_by_id.get(wip_id)
        if not wip:
            continue
        overdue = overdue_invoices_by_wip.get(wip_id, [])
        dashboard_files.append({'wip': wip, 'overdue_invoices': overdue})

    dashboard_files.sort(
        key=lambda entry: (
            0 if entry['overdue_invoices'] else 1, display_ids.index(entry['wip'].id))
    )

    overdue_invoice_file_count = sum(
        1 for entry in dashboard_files if entry['overdue_invoices']
    )
    unsettled_invoices = [
        invoice
        for invoices in overdue_invoices_by_wip.values()
        for invoice in invoices
    ]

    return dashboard_files, overdue_invoice_file_count, unsettled_invoices


def serialize_kanban_task(task, request_user, *, is_completed=False):
    file_number = task.file_number.file_number if task.file_number else ''
    home_url = reverse('home', args=[file_number]) if file_number else ''
    task_text = task.task or ''
    if len(task_text) > 80:
        task_text = task_text[:80] + '...'

    is_owner = task.person_id == request_user.id
    can_edit = is_owner
    edit_url = ''
    if can_edit:
        if is_completed:
            edit_url = reverse('edit_last_work', args=[task.id])
        else:
            edit_url = reverse('edit_next_work', args=[task.id])

    payload = {
        'id': task.id,
        'file_number': file_number,
        'home_url': home_url,
        'task': task_text,
        'date': task.date.isoformat() if task.date else None,
        'timestamp': task.timestamp.isoformat(),
        'assigned_to': task.person.get_full_name() if task.person else 'Unassigned',
        'created_by': task.created_by.get_full_name() if task.created_by else 'Unknown',
        'is_created_by_me': task.created_by_id == request_user.id if task.created_by_id else False,
        'can_edit': can_edit,
        'edit_url': edit_url,
    }
    if not is_completed:
        payload['urgency'] = task.urgency
    return payload


def get_key_document_expiry_alerts(wips, warning_days=30):
    today = timezone.localdate()
    warning_date = today + timedelta(days=warning_days)
    client_file_numbers = {}

    rows = wips.filter(file_status__status__in=['Open', 'To Be Closed']).values(
        'file_number', 'client1_id', 'client2_id')
    for row in rows:
        if row['client1_id']:
            client_file_numbers.setdefault(
                row['client1_id'], set()).add(row['file_number'])
        if row['client2_id']:
            client_file_numbers.setdefault(
                row['client2_id'], set()).add(row['file_number'])

    documents = ClientKeyDocument.objects.filter(
        client_id__in=client_file_numbers.keys(),
        expiry_date__isnull=False,
        expiry_date__lte=warning_date,
    ).select_related('client').order_by('expiry_date', 'client__name', 'category')

    alerts = []
    for document in documents:
        alerts.append({
            'client_id': document.client_id,
            'client_name': document.client.name,
            'document_category': document.get_category_display(),
            'document_type': document.document_type,
            'document_reference': document.document_reference,
            'expiry_date': document.expiry_date,
            'status': 'expired' if document.expiry_date < today else 'due_soon',
            'file_numbers': sorted(client_file_numbers.get(document.client_id, [])),
        })

    return sorted(alerts, key=lambda alert: (
        alert['expiry_date'], alert['client_name'], alert['document_category'], alert['document_type']))


def get_missing_key_document_alerts(wips):
    active_wips = list(wips.filter(
        file_status__status__in=['Open', 'To Be Closed']
    ).select_related('client1', 'client2').order_by('file_number'))
    client_ids = set()
    for matter in active_wips:
        if matter.client1_id:
            client_ids.add(matter.client1_id)
        if matter.client2_id:
            client_ids.add(matter.client2_id)

    existing_documents = ClientKeyDocument.objects.filter(
        client_id__in=client_ids
    ).values('client_id', 'category')
    document_categories_by_client = {}
    for document in existing_documents:
        document_categories_by_client.setdefault(
            document['client_id'], set()).add(document['category'])

    required_categories = [
        ('proof_of_id', 'Proof of ID'),
        ('proof_of_address', 'Proof of Address'),
    ]
    alerts = []
    for matter in active_wips:
        matter_clients = [matter.client1]
        if matter.client2:
            matter_clients.append(matter.client2)

        for client in matter_clients:
            existing_categories = document_categories_by_client.get(
                client.id, set())
            for category, label in required_categories:
                if category not in existing_categories:
                    alerts.append({
                        'file_number': matter.file_number,
                        'client_id': client.id,
                        'client_name': client.name,
                        'document_category': label,
                    })

    return alerts


def get_dashboard_key_document_scope_wips(user, key_doc_scope):
    if key_doc_scope == 'all_active':
        return WIP.objects.filter(
            file_status__status__in=['Open', 'To Be Closed']
        ).distinct(), 'all_active'

    return get_user_dashboard_wips(user).filter(
        file_status__status__in=['Open', 'To Be Closed']
    ).distinct(), 'associated'


def get_matter_key_documents(matter):
    clients = [matter.client1]
    if matter.client2:
        clients.append(matter.client2)

    documents = ClientKeyDocument.objects.filter(
        client__in=clients
    ).select_related('client', 'verified_by').order_by(
        'client__name', 'category', 'expiry_date', 'document_type')

    today = timezone.localdate()
    key_documents = []
    for document in documents:
        key_documents.append({
            'document': document,
            'status': (
                'expired' if document.expiry_date and document.expiry_date < today
                else 'due_soon' if document.expiry_date and document.expiry_date <= today + timedelta(days=30)
                else 'current'
            )
        })

    return key_documents


def get_risk_assessments_due_queryset(wip_queryset):
    one_year_ago = timezone.localdate() - relativedelta(years=1)

    latest_assessment_subquery = RiskAssessment.objects.filter(
        matter=OuterRef('pk')
    ).order_by('-due_diligence_date').values('due_diligence_date')[:1]

    latest_monitoring_subquery = OngoingMonitoring.objects.filter(
        file_number=OuterRef('pk')
    ).order_by('-date_due_diligence_conducted').values('date_due_diligence_conducted')[:1]

    return wip_queryset.annotate(
        latest_assessment_date=Subquery(latest_assessment_subquery),
        latest_monitoring_date=Subquery(latest_monitoring_subquery)
    ).filter(
        Q(file_status__status__in=['Open', 'To Be Closed']) &
        (
            Q(latest_assessment_date__isnull=True) |
            (
                Q(latest_assessment_date__lte=one_year_ago) &
                (
                    Q(latest_monitoring_date__isnull=True) |
                    Q(latest_monitoring_date__lte=one_year_ago)
                )
            )
        )
    ).order_by('file_number')


def get_file_reviews_due_queryset(wip_queryset):
    three_months_ago = timezone.localdate() - relativedelta(months=3)

    latest_review_subquery = MatterFileReview.objects.filter(
        matter=OuterRef('pk')
    ).order_by('-date_review_completed').values('date_review_completed')[:1]

    latest_review_by_subquery = MatterFileReview.objects.filter(
        matter=OuterRef('pk')
    ).order_by('-date_review_completed').values('file_review_completed_by__first_name')[:1]

    return wip_queryset.annotate(
        latest_review_date=Subquery(latest_review_subquery),
        latest_review_by=Subquery(latest_review_by_subquery),
    ).filter(
        Q(file_status__status__in=['Open', 'To Be Closed']) &
        (
            Q(latest_review_date__isnull=True) |
            Q(latest_review_date__lte=three_months_ago)
        )
    ).order_by('file_number')


def get_dashboard_risk_scope_wips(user, risk_scope):
    if risk_scope == 'all_active':
        return WIP.objects.filter(
            file_status__status__in=['Open', 'To Be Closed']
        ).distinct(), 'all_active'

    associated_wips = get_user_dashboard_wips(user).filter(
        file_status__status__in=['Open', 'To Be Closed']
    ).distinct()
    return associated_wips, 'associated'


@login_required
def display_data_index_page(request):
    if 'valToSearch' in request.POST:
        search_by = request.POST.get('searchBy', 'FileNumber')
        val_to_search = request.POST.get('valToSearch', '')
        show_archived = 'showArchived' in request.POST

        data = get_index_search_data(
            search_by=search_by,
            val_to_search=val_to_search,
            show_archived=show_archived
        )

        context = {
            'search_by': search_by,
            'val_to_search': val_to_search,
            'show_archived': show_archived,
            'data': data,
        }

        return render(request, 'index.html', context)

    return render(request, 'index.html')


@login_required
def download_search_report(request):
    if 'valToSearch' in request.POST:
        search_by = request.POST.get('searchBy', 'FileNumber')
        val_to_search = request.POST.get('valToSearch', '')
        show_archived = 'showArchived' in request.POST

        data = get_index_search_data(
            search_by=search_by,
            val_to_search=val_to_search,
            show_archived=show_archived
        )

        context = {
            'search_by': search_by,
            'val_to_search': val_to_search,
            'show_archived': show_archived,
            'data': data,
            'user': request.user
        }

        html_string = render_to_string(
            'download_templates/search_report.html', context)

        pdf_file = HTML(string=html_string).write_pdf()

        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="search_report_{search_by}+{val_to_search}.pdf"'
        return response


def get_index_search_filter(search_by, val_to_search, show_archived):
    file_status_field_name = 'file_status__status'

    if show_archived:
        file_status_list = ['Open', 'Archived']
        filter_factor = Q(
            **{f"{file_status_field_name}__in": file_status_list})
    elif search_by == 'ToBeClosed':
        filter_factor = Q(**{file_status_field_name: 'To Be Closed'})
    else:
        filter_factor = Q(**{file_status_field_name: 'Open'})

    if search_by == 'ClientName':
        filter_factor &= Q(client1__name__icontains=val_to_search) | Q(
            client2__name__icontains=val_to_search
        )
    elif search_by == 'FeeEarner':
        if val_to_search == "DC":
            filter_factor &= Q(fee_earner=None)
        else:
            filter_factor &= Q(fee_earner__username__icontains=val_to_search)
    else:
        filter_factor &= Q(file_number__icontains=val_to_search)

    return filter_factor


def get_index_search_data(search_by, val_to_search, show_archived):
    filter_factor = get_index_search_filter(
        search_by, val_to_search, show_archived)

    last_work_subquery = LastWork.objects.filter(
        file_number=OuterRef('pk')
    ).order_by('-timestamp')

    latest_email_subquery = MatterEmails.objects.filter(
        file_number=OuterRef('pk')
    ).annotate(
        email_activity_date=Coalesce(
            Cast('time', output_field=DateField()),
            Cast('timestamp', output_field=DateField())
        ),
        email_activity_desc=Coalesce(
            'description',
            'subject',
            Value('Email activity', output_field=TextField()),
            output_field=TextField()
        )
    ).order_by('-email_activity_date', '-timestamp')

    latest_attendance_subquery = MatterAttendanceNotes.objects.filter(
        file_number=OuterRef('pk'),
        date__isnull=False
    ).order_by('-date', '-timestamp')

    fallback_date = date(1900, 1, 1)

    data = WIP.objects.filter(filter_factor).annotate(
        latest_work_task_raw=Subquery(last_work_subquery.values('task')[
                                      :1], output_field=TextField()),
        latest_work_person_raw=Subquery(
            last_work_subquery.values('person__username')[:1]),
        latest_work_entry_date=Subquery(last_work_subquery.values('date')[
                                        :1], output_field=DateField()),
        latest_email_task_raw=Subquery(latest_email_subquery.values(
            'email_activity_desc')[:1], output_field=TextField()),
        latest_email_person_raw=Subquery(
            latest_email_subquery.values('fee_earner__username')[:1]),
        latest_email_date=Subquery(latest_email_subquery.values(
            'email_activity_date')[:1], output_field=DateField()),
        latest_attendance_task_raw=Subquery(latest_attendance_subquery.values(
            'subject_line')[:1], output_field=TextField()),
        latest_attendance_is_charged_raw=Subquery(latest_attendance_subquery.values(
            'is_charged')[:1], output_field=BooleanField()),
        latest_attendance_person_raw=Subquery(
            latest_attendance_subquery.values('person_attended__username')[:1]),
        latest_attendance_note_date=Subquery(latest_attendance_subquery.values('date')[
                                             :1], output_field=DateField())
    ).annotate(
        work_date_for_compare=Coalesce('latest_work_entry_date', Value(
            fallback_date, output_field=DateField())),
        email_date_for_compare=Coalesce('latest_email_date', Value(
            fallback_date, output_field=DateField())),
        attendance_date_for_compare=Coalesce(
            'latest_attendance_note_date', Value(fallback_date, output_field=DateField()))
    ).annotate(
        latest_last_work_date=Case(
            When(
                latest_work_entry_date__isnull=True,
                latest_email_date__isnull=True,
                latest_attendance_note_date__isnull=True,
                then=Value(None, output_field=DateField())
            ),
            default=Greatest(
                'work_date_for_compare',
                'email_date_for_compare',
                'attendance_date_for_compare'
            ),
            output_field=DateField()
        ),
        latest_activity_source=Case(
            When(
                latest_work_entry_date__isnull=True,
                latest_email_date__isnull=True,
                latest_attendance_note_date__isnull=True,
                then=Value('none')
            ),
            When(
                Q(work_date_for_compare__gte=F('email_date_for_compare')) &
                Q(work_date_for_compare__gte=F('attendance_date_for_compare')),
                then=Value('work')
            ),
            When(
                email_date_for_compare__gte=F('attendance_date_for_compare'),
                then=Value('email')
            ),
            default=Value('attendance'),
            output_field=CharField()
        )
    ).annotate(
        latest_last_work_task=Case(
            When(
                latest_activity_source='work',
                then=Coalesce('latest_work_task_raw', Value(
                    'Work activity', output_field=TextField()), output_field=TextField())
            ),
            When(
                latest_activity_source='email',
                then=Coalesce('latest_email_task_raw', Value(
                    'Email activity', output_field=TextField()), output_field=TextField())
            ),
            When(
                latest_activity_source='attendance',
                then=Case(
                    When(
                        latest_attendance_is_charged_raw=False,
                        then=Concat(
                            Coalesce('latest_attendance_task_raw', Value(
                                'Attendance note', output_field=TextField()), output_field=TextField()),
                            Value(' (N/C)', output_field=TextField()),
                            output_field=TextField()
                        )
                    ),
                    default=Coalesce('latest_attendance_task_raw', Value(
                        'Attendance note', output_field=TextField()), output_field=TextField()),
                    output_field=TextField()
                )
            ),
            default=Value(None, output_field=TextField()),
            output_field=TextField()
        ),
        latest_last_work_person=Case(
            When(latest_activity_source='work', then=Coalesce(
                'latest_work_person_raw', Value('-'))),
            When(latest_activity_source='email', then=Coalesce(
                'latest_email_person_raw', Value('-'))),
            When(latest_activity_source='attendance', then=Coalesce(
                'latest_attendance_person_raw', Value('-'))),
            default=Value(None, output_field=CharField()),
            output_field=CharField()
        )
    ).values(
        'file_number',
        'fee_earner__username',
        'matter_description',
        'client1__name',
        'client2__name',
        'comments',
        'latest_last_work_date',
        'latest_last_work_person',
        'latest_last_work_task',
    ).order_by('file_number')

    return data


@login_required
def user_dashboard(request):
    user = CustomUser.objects.get(username=request.user)
    risk_scope = request.GET.get('risk_scope', 'associated')
    key_doc_scope = request.GET.get('key_doc_scope', 'associated')
    user_next_works = NextWork.objects.filter(
        Q(person=user) & Q(completed=False)).order_by('date')
    user_last_works = LastWork.objects.filter(person=user).order_by('-date')

    # Check if user has pending tasks (for cookie logic)
    has_pending_tasks = NextWork.objects.filter(
        person=user,
        status__in=['to_do', 'in_progress']
    ).exists()
    # Check if user is manager and show pending holiday requests
    if user.is_manager:
        pending_holiday_requests = HolidayRecord.objects.filter(
            approved=False,
            checked_by__isnull=True
        ).count()

        if pending_holiday_requests > 0:
            if pending_holiday_requests == 1:
                messages.info(
                    request, f'You have {pending_holiday_requests} holiday request pending your approval.')
            else:
                messages.info(
                    request, f'You have {pending_holiday_requests} holiday requests pending your approval.')

    now = timezone.now()
    unique_wips = get_user_dashboard_wips(user)
    dashboard_files, overdue_invoice_file_count, unsettled_invoices = build_dashboard_files(
        user)
    risk_scope_wips, validated_risk_scope = get_dashboard_risk_scope_wips(
        user, risk_scope
    )
    key_doc_scope_wips, validated_key_doc_scope = get_dashboard_key_document_scope_wips(
        user, key_doc_scope
    )
    risk_assessments_due = get_risk_assessments_due_queryset(risk_scope_wips)
    file_reviews_due = get_file_reviews_due_queryset(unique_wips)

    eleven_months_ago = timezone.now() - relativedelta(months=11)
    unique_aml_checks_due = get_aml_checks_due_from_wips(
        unique_wips, eleven_months_ago, user=user)
    key_document_expiry_alerts = get_key_document_expiry_alerts(
        key_doc_scope_wips)
    missing_key_document_alerts = get_missing_key_document_alerts(
        key_doc_scope_wips)
    last_100_emails = MatterEmails.objects.filter(
        fee_earner=user).order_by('-time')[:100]

    latest_version_subquery = PolicyVersion.objects.filter(
        policy=OuterRef('pk')
    ).order_by('-version_number').values('pk')[:1]

    unread_policies_exist = Policy.objects.annotate(
        latest_version_id=Subquery(latest_version_subquery),
        is_read=Exists(
            PoliciesRead.objects.filter(
                policy=OuterRef('pk'),
                policy_version_id=OuterRef('latest_version_id'),
                read_by=user
            )
        )
    ).filter(is_read=False).exists()

    context = {
        'now': now,
        'user_next_works': user_next_works,
        'user_last_works': user_last_works,
        'risk_assessments_due_files': risk_assessments_due,
        'aml_checks_due': unique_aml_checks_due,
        'key_document_expiry_alerts': key_document_expiry_alerts,
        'missing_key_document_alerts': missing_key_document_alerts,
        'unsettled_invoices': unsettled_invoices,
        'dashboard_files': dashboard_files,
        'overdue_invoice_file_count': overdue_invoice_file_count,
        'last_100_emails': last_100_emails,
        'files': unique_wips,
        'unread_policies_exist': unread_policies_exist,
        'has_pending_tasks': has_pending_tasks,
        'risk_scope': validated_risk_scope,
        'key_doc_scope': validated_key_doc_scope,
        'file_reviews_due_files': file_reviews_due,
    }

    return render(request, 'dashboard.html', context)


@login_required
@require_POST
def update_task_status(request):
    """Update task status via AJAX"""
    try:
        data = json.loads(request.body)
        task_id = data.get('task_id')
        new_status = data.get('status')

        if not task_id or not new_status:
            return JsonResponse({
                'success': False,
                'error': 'Missing task_id or status'
            })

        # Allow both the assigned person and the creator to update the task
        task = get_object_or_404(NextWork, Q(id=task_id) & (
            Q(person=request.user) | Q(created_by=request.user)))
        task.status = new_status
        task.save()

        return JsonResponse({'success': True})

    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
@require_POST
def load_initial_tasks(request):
    """Load initial tasks for all statuses via AJAX"""
    try:
        data = json.loads(request.body)
        count = data.get('count', 5)  # Default to 5 tasks per status
        filter_created_by_me = data.get('filter_created_by_me', False)

        # Handle "all" case - when count is "all", we don't limit the queryset
        show_all = count == 'all'

        task_data = {}
        total_counts = {}

        # Base query filters
        base_filter_nextwork = {'created_by': request.user} if filter_created_by_me else {
            'person': request.user}
        base_filter_lastwork = {'created_by': request.user} if filter_created_by_me else {
            'person': request.user}

        # Load To Do tasks with smart ordering (urgency priority, then due date)
        to_do_tasks = NextWork.objects.filter(
            **base_filter_nextwork,
            status='to_do'
        ).select_related('person', 'created_by', 'file_number').extra(
            select={
                'urgency_order': "CASE urgency WHEN 'urgent' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 WHEN 'low' THEN 4 ELSE 5 END"
            }
        ).order_by('urgency_order', 'date')

        total_counts['to_do'] = to_do_tasks.count()
        to_do_limited = to_do_tasks if show_all else to_do_tasks[:count]

        task_data['to_do'] = [
            serialize_kanban_task(task, request.user)
            for task in to_do_limited
        ]

        # Load In Progress tasks with smart ordering (urgency priority, then due date)
        in_progress_tasks = NextWork.objects.filter(
            **base_filter_nextwork,
            status='in_progress'
        ).select_related('person', 'created_by', 'file_number').extra(
            select={
                'urgency_order': "CASE urgency WHEN 'urgent' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 WHEN 'low' THEN 4 ELSE 5 END"
            }
        ).order_by('urgency_order', 'date')

        total_counts['in_progress'] = in_progress_tasks.count()
        in_progress_limited = in_progress_tasks if show_all else in_progress_tasks[:count]

        task_data['in_progress'] = [
            serialize_kanban_task(task, request.user)
            for task in in_progress_limited
        ]

        # Load Completed tasks (last 7 days)
        completed_tasks = LastWork.objects.filter(
            **base_filter_lastwork,
            timestamp__gte=timezone.now() - timezone.timedelta(days=7)
        ).select_related('person', 'created_by', 'file_number').order_by('-timestamp')

        total_counts['completed'] = completed_tasks.count()
        completed_limited = completed_tasks if show_all else completed_tasks[:count]

        task_data['completed'] = [
            serialize_kanban_task(task, request.user, is_completed=True)
            for task in completed_limited
        ]

        return JsonResponse({
            'success': True,
            'tasks': task_data,
            'total_counts': total_counts,
            'loaded_count': count,
            'filter_applied': filter_created_by_me
        })

    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
@require_POST
def load_more_tasks(request):
    """Load more tasks for a specific status via AJAX"""
    try:
        data = json.loads(request.body)
        status = data.get('status')
        offset = data.get('offset', 0)
        count = data.get('count', 3)
        filter_created_by_me = data.get('filter_created_by_me', False)

        if not status:
            return JsonResponse({
                'success': False,
                'error': 'Missing status'
            })

        tasks_data = []
        total_count = 0

        # Base query filters
        base_filter_nextwork = {'created_by': request.user} if filter_created_by_me else {
            'person': request.user}
        base_filter_lastwork = {'created_by': request.user} if filter_created_by_me else {
            'person': request.user}

        if status in ['to_do', 'in_progress']:
            all_tasks = NextWork.objects.filter(
                **base_filter_nextwork,
                status=status
            ).select_related('person', 'created_by', 'file_number').extra(
                select={
                    'urgency_order': "CASE urgency WHEN 'urgent' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 WHEN 'low' THEN 4 ELSE 5 END"
                }
            ).order_by('urgency_order', 'date')

            total_count = all_tasks.count()
            tasks = all_tasks[offset:offset + count]

            tasks_data = [
                serialize_kanban_task(task, request.user)
                for task in tasks
            ]

        elif status == 'completed':
            # For completed tasks, we show from LastWork
            all_tasks = LastWork.objects.filter(
                **base_filter_lastwork,
                timestamp__gte=timezone.now() - timezone.timedelta(days=7)
            ).select_related('person', 'created_by', 'file_number').order_by('-timestamp')

            total_count = all_tasks.count()
            tasks = all_tasks[offset:offset + count]

            tasks_data = [
                serialize_kanban_task(task, request.user, is_completed=True)
                for task in tasks
            ]

        has_more = (offset + count) < total_count

        return JsonResponse({
            'success': True,
            'tasks': tasks_data,
            'has_more': has_more,
            'total_count': total_count,
            'loaded_count': len(tasks_data)
        })

    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
def get_files(request):
    """Get list of files for modal dropdown"""
    try:
        files = WIP.objects.filter(
            file_status__status='Open'
        ).values('id', 'file_number', 'matter_description').order_by('file_number')

        files_data = []
        for file in files:
            files_data.append({
                'id': file['id'],
                'file_number': file['file_number'],
                'description': file['matter_description'] or 'No description',
                'display_text': f"{file['file_number']} - {file['matter_description'] or 'No description'}"
            })

        return JsonResponse({
            'success': True,
            'files': files_data
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
def get_users(request):
    """Get list of users for modal dropdown"""
    try:
        users = CustomUser.objects.all().values('id', 'first_name', 'last_name')

        users_data = []
        for user in users:
            name = f"{user['first_name']} {user['last_name']}".strip()
            if not name:
                name = f"User {user['id']}"
            users_data.append({
                'id': user['id'],
                'name': name
            })

        return JsonResponse({
            'success': True,
            'users': users_data
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
def get_risk_assessments_due_data(request):
    """Get risk assessments due for dashboard without page reload."""
    try:
        user = CustomUser.objects.get(username=request.user)
        risk_scope = request.GET.get('risk_scope', 'associated')
        risk_scope_wips, validated_risk_scope = get_dashboard_risk_scope_wips(
            user, risk_scope
        )
        risk_assessments_due = get_risk_assessments_due_queryset(
            risk_scope_wips
        ).select_related('client1', 'client2')

        data = []
        for assessment in risk_assessments_due:
            reason_due = 'No risk assessment completed'
            if assessment.latest_assessment_date:
                reason_due = 'No ongoing monitoring in the last year'

            data.append({
                'file_number': assessment.file_number,
                'matter_description': assessment.matter_description or '',
                'client1_name': assessment.client1.name if assessment.client1 else '',
                'client2_name': assessment.client2.name if assessment.client2 else '',
                'latest_assessment_date_display': (
                    assessment.latest_assessment_date.strftime('%d/%m/%Y')
                    if assessment.latest_assessment_date else None
                ),
                'latest_monitoring_date_display': (
                    assessment.latest_monitoring_date.strftime('%d/%m/%Y')
                    if assessment.latest_monitoring_date else None
                ),
                'reason_due': reason_due,
                'home_url': reverse('home', args=[assessment.file_number])
            })

        return JsonResponse({
            'success': True,
            'risk_scope': validated_risk_scope,
            'count': len(data),
            'items': data
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
@require_POST
def create_task(request):
    """Create a new NextWork task"""
    try:
        data = json.loads(request.body)

        # Create new NextWork instance
        task = NextWork(
            file_number_id=data.get('file_number'),
            person_id=data.get('person'),
            task=data.get('task'),
            date=data.get('date') if data.get('date') else None,
            urgency=data.get('urgency', 'medium'),
            status=data.get('status', 'to_do'),
            created_by=request.user
        )
        task.save()

        return JsonResponse({
            'success': True,
            'task_id': task.id
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
def display_data_home_page(request, file_number):
    try:
        logger.info(
            f'User {request.user.username} accessing matter file {file_number}')
        matter = WIP.objects.select_related(
            'fee_earner',
            'client1',
            'client1__created_by',
            'client2',
            'client2__created_by',
            'matter_type',
            'file_status',
            'file_location',
            'other_side',
            'authorised_party1',
            'authorised_party2',
            'created_by',
        ).get(file_number=file_number)
        matter_file_reviews = MatterFileReview.objects.filter(
            matter=matter
        ).select_related(
            'file_reviewed_by',
            'file_review_completed_by',
            'created_by',
        ).order_by('-date_review_completed', '-date_reviewed', '-timestamp')
        matter_key_dates = MatterKeyDate.objects.filter(matter=matter)
        matter_key_date_form = MatterKeyDateForm()
        matter_key_document_form = MatterClientKeyDocumentForm(matter=matter)
        matter_key_documents = get_matter_key_documents(matter)
        today = timezone.localdate()
        matter_key_date_alerts = matter_key_dates.filter(
            date__gte=today,
            date__lte=today + timedelta(days=30)
        )
        matter_key_document_alerts = [
            item for item in matter_key_documents if item['status'] == 'expired'
        ]
        undertakings = Undertaking.objects.filter(file_number=matter).select_related(
            'given_by', 'discharged_by')
        next_work_form = NextWorkFormWithoutFileNumber()
        next_work = NextWork.objects.filter(
            file_number=matter, completed=False).select_related(
            'person', 'created_by').order_by('date')
        last_work = LastWork.objects.filter(
            file_number=matter).select_related('person', 'created_by').order_by('-date')
        last_work_form = LastWorkFormWithoutFileNumber()
        ongoing_monitorings = OngoingMonitoring.objects.filter(
            file_number=matter.id).select_related(
            'signed_by', 'created_by').order_by('-timestamp')
        risk_assessment = RiskAssessment.objects.filter(
            matter=matter
        ).select_related('due_diligence_signed_by').order_by('-due_diligence_date')
        eleven_months_ago = (timezone.now() - relativedelta(months=11)).date()
        if risk_assessment.exists():
            risk_assessment = risk_assessment[0]
            if risk_assessment.due_diligence_date <= eleven_months_ago:
                if ongoing_monitorings.exists():
                    if ongoing_monitorings[0].date_due_diligence_conducted <= eleven_months_ago:
                        eleven_months_since_last_risk_assessment = True
                    else:
                        eleven_months_since_last_risk_assessment = False
                else:
                    eleven_months_since_last_risk_assessment = True
            else:
                eleven_months_since_last_risk_assessment = False
        else:
            risk_assessment = ""
            eleven_months_since_last_risk_assessment = False
        if matter.file_status.status == 'Archived':
            messages.error(
                request, "ARCHIVED MATTER. Please note this matter is archived.")
        bundles = Bundle.objects.filter(
            created_by=request.user,
            file_number=matter,
        ).order_by('-created_at')
        activity_logs, log_meta = enrich_file_logs(
            get_file_logs(file_number, limit=300))
        return render(request, 'home.html', {'matter': matter,
                                             'bundles': bundles,
                                             'undertakings': undertakings,
                                             'file_number': file_number,
                                             'next_work_form': next_work_form, 'next_work': next_work,
                                             'last_work': last_work, 'last_work_form': last_work_form,
                                             'ongoing_monitorings': ongoing_monitorings,
                                             'risk_assessment': risk_assessment, 'eleven_months_since_last_risk_assessment': eleven_months_since_last_risk_assessment,
                                             'matter_file_reviews': build_matter_file_review_display_data(matter_file_reviews),
                                             'matter_key_dates': matter_key_dates,
                                             'matter_key_date_form': matter_key_date_form,
                                             'matter_key_document_form': matter_key_document_form,
                                             'matter_key_documents': matter_key_documents,
                                             'matter_key_date_alerts': matter_key_date_alerts,
                                             'matter_key_document_alerts': matter_key_document_alerts,
                                             'logs': activity_logs,
                                             'log_meta': log_meta,
                                             'log_limit': 300})
    except WIP.DoesNotExist:
        logger.warning(
            f'User {request.user.username} attempted to access non-existent matter file {file_number}')
        messages.error(request, 'Matter file not found')
        return render(request, 'home.html', {'error': 'Matter file not found'})
    except Exception as e:
        logger.error(
            f'Error displaying matter file {file_number} for user {request.user.username}: {str(e)}', exc_info=True)
        messages.error(
            request, 'An error occurred while loading the matter file')
        return redirect('index')


@login_required
@require_POST
def add_matter_key_date(request, file_number):
    matter = get_object_or_404(WIP, file_number=file_number)
    form = MatterKeyDateForm(request.POST)
    if form.is_valid():
        key_date = form.save(commit=False)
        key_date.matter = matter
        key_date.created_by = request.user
        key_date.save()
        log_created(
            request.user,
            key_date,
            snapshot_key_date(key_date),
        )
        messages.success(request, 'Key date added.')
    else:
        messages.error(request, 'Please correct the key date form.')

    return redirect('home', file_number=file_number)


@login_required
@require_POST
def add_matter_key_document(request, file_number):
    matter = get_object_or_404(WIP, file_number=file_number)
    form = MatterClientKeyDocumentForm(matter=matter, data=request.POST)
    if form.is_valid():
        document = form.save(commit=False)
        if document.verified_on:
            document.verified_by = request.user
        document.save()
        log_created(
            request.user,
            document,
            snapshot_key_document(document),
        )
        messages.success(request, 'Key document added.')
    else:
        messages.error(request, 'Please correct the key document form.')

    return redirect('home', file_number=file_number)


@login_required
def edit_matter_key_date(request, id):
    key_date = get_object_or_404(MatterKeyDate, id=id)
    if request.method == 'POST':
        form = MatterKeyDateForm(request.POST, instance=key_date)
        if form.is_valid():
            duplicate_obj = copy.deepcopy(key_date)
            changed_fields = form.changed_data
            form.save()
            changes = build_form_field_changes(
                duplicate_obj, key_date, changed_fields)
            if changes:
                create_modification(
                    user=request.user,
                    modified_obj=key_date,
                    changes=changes,
                )
            messages.success(request, 'Key date updated.')
            return redirect('home', file_number=key_date.matter.file_number)
        messages.error(request, 'Please correct the key date form.')
    else:
        form = MatterKeyDateForm(instance=key_date)

    return render(request, 'edit_models.html', {
        'form': form,
        'title': 'Matter Key Date'
    })


@login_required
@require_POST
def delete_matter_key_date(request, id):
    key_date = get_object_or_404(MatterKeyDate, id=id)
    file_number = key_date.matter.file_number
    log_deleted_on_parent(
        request.user,
        key_date.matter,
        'key_date',
        snapshot_key_date(key_date),
    )
    key_date.delete()
    messages.success(request, 'Key date deleted.')
    return redirect('home', file_number=file_number)


def _month_bounds(year, month):
    last_day = calendar_module.monthrange(year, month)[1]
    return date(year, month, 1), date(year, month, last_day)


def _coerce_calendar_month(request):
    today = timezone.localdate()
    try:
        year = int(request.GET.get('year', today.year))
        month = int(request.GET.get('month', today.month))
        first_day, last_day = _month_bounds(year, month)
    except (TypeError, ValueError):
        year = today.year
        month = today.month
        first_day, last_day = _month_bounds(year, month)

    return year, month, first_day, last_day


def _range_query_string(request, start_date, end_date, view_months=1):
    params = request.GET.copy()
    params['year'] = start_date.year
    params['month'] = start_date.month
    params['date_from'] = start_date.isoformat()
    params['date_to'] = end_date.isoformat()
    params['view_months'] = view_months
    return params.urlencode()


KEY_DATE_SOURCE_CHOICES = [
    ('matter_key_date', 'Matter key dates'),
    ('risk_assessment_due', 'Risk assessments due'),
    ('ongoing_monitoring_due', 'Ongoing monitoring due'),
    ('aml_due', 'AML checks due'),
    ('id_expiry', 'ID/address expiry or review'),
    ('staff_leave', 'Staff leave'),
    ('sickness', 'Sickness'),
    ('office_closure', 'Office closures'),
    ('bank_holiday', 'Bank holidays'),
]


DEFAULT_KEY_DATE_SOURCES = [
    'matter_key_date',
    'staff_leave',
    'sickness',
    'office_closure',
    'bank_holiday',
]


VIEW_MONTH_CHOICES = [1, 2, 6]


def _make_central_calendar_event(
    *,
    event_date,
    title,
    matter=None,
    source,
    source_label,
    subtitle='',
    time_value=None,
    edit_url='',
    detail_url='',
    person_label='',
    location='',
    notes='',
):
    if not detail_url and matter:
        detail_url = reverse('home', args=[matter.file_number])
    return {
        'date': event_date,
        'time': time_value,
        'title': title,
        'matter': matter,
        'source': source,
        'source_label': source_label,
        'subtitle': subtitle,
        'edit_url': edit_url,
        'detail_url': detail_url,
        'person_label': person_label,
        'location': location,
        'notes': notes,
    }


def _event_date(value):
    if not value:
        return None
    if isinstance(value, datetime):
        return timezone.localtime(value).date()
    return value


def _iter_event_dates(start_value, end_value, lower_bound, upper_bound):
    start_date = _event_date(start_value)
    end_date = _event_date(end_value) or start_date
    if not start_date:
        return

    current = max(start_date, lower_bound)
    last_date = min(end_date, upper_bound)
    while current <= last_date:
        yield current
        current += timedelta(days=1)


def _source_counts_for_day(events):
    count_lookup = {}
    for event in events:
        count_lookup[event['source']] = count_lookup.get(
            event['source'], 0) + 1

    counts = []
    for source, label in KEY_DATE_SOURCE_CHOICES:
        count = count_lookup.get(source)
        if count:
            counts.append({
                'source': source,
                'label': label,
                'count': count,
            })
    return counts


def _build_central_key_dates_context(request):
    year, month, month_start, month_end = _coerce_calendar_month(request)

    try:
        selected_view_months = int(request.GET.get('view_months', 1))
    except (TypeError, ValueError):
        selected_view_months = 1
    if selected_view_months not in VIEW_MONTH_CHOICES:
        selected_view_months = 1

    today = timezone.localdate()
    has_custom_date_range = 'date_from' in request.GET or 'date_to' in request.GET
    default_end = (
        month_start + relativedelta(months=selected_view_months)) - timedelta(days=1)
    if selected_view_months == 1 and not has_custom_date_range:
        default_start = today - timedelta(days=15)
        default_end = today + timedelta(days=15)
    else:
        default_start = month_start

    date_from = parse_date(request.GET.get('date_from') or '') or default_start
    date_to = parse_date(request.GET.get('date_to') or '') or default_end
    if date_from > date_to:
        date_from, date_to = date_to, date_from
    open_status = FileStatus.objects.filter(status__iexact='Open').first()
    selected_file_status = request.GET.get('file_status')
    if selected_file_status is None and open_status:
        selected_file_status = str(open_status.id)

    selected_fee_earner = request.GET.get('fee_earner', '')
    selected_matter_type = request.GET.get('matter_type', '')
    selected_matter = request.GET.get('matter', '')
    selected_date_type = request.GET.get('date_type', '')
    matter_search = request.GET.get('matter_search', '').strip()
    client_search = request.GET.get('client_search', '').strip()
    selected_sources = request.GET.getlist('source')
    if 'sources_submitted' not in request.GET:
        selected_sources = DEFAULT_KEY_DATE_SOURCES.copy()

    base_matters = WIP.objects.select_related(
        'fee_earner',
        'matter_type',
        'file_status',
        'client1',
        'client2',
        'authorised_party1',
        'authorised_party2',
    )

    if selected_fee_earner:
        base_matters = base_matters.filter(fee_earner_id=selected_fee_earner)
    if selected_matter_type:
        base_matters = base_matters.filter(matter_type_id=selected_matter_type)
    if selected_file_status and selected_file_status != 'all':
        base_matters = base_matters.filter(file_status_id=selected_file_status)
    if selected_matter:
        base_matters = base_matters.filter(id=selected_matter)
    if matter_search:
        base_matters = base_matters.filter(
            Q(file_number__icontains=matter_search)
            | Q(matter_description__icontains=matter_search)
        )
    if client_search:
        base_matters = base_matters.filter(
            Q(client1__name__icontains=client_search)
            | Q(client2__name__icontains=client_search)
        )

    matter_ids = base_matters.values_list('id', flat=True)

    key_dates = MatterKeyDate.objects.select_related(
        'matter',
        'matter__fee_earner',
        'matter__matter_type',
        'matter__file_status',
        'matter__client1',
        'matter__client2',
        'created_by',
    ).filter(
        matter_id__in=matter_ids,
        date__gte=date_from,
        date__lte=date_to,
    )

    if selected_date_type:
        key_dates = key_dates.filter(date_type=selected_date_type)

    key_dates = key_dates.order_by(
        'date', 'time', 'matter__file_number', 'title')

    events = []
    if 'matter_key_date' in selected_sources:
        for key_date in key_dates:
            events.append(_make_central_calendar_event(
                event_date=key_date.date,
                time_value=key_date.time,
                title=key_date.title,
                matter=key_date.matter,
                source='matter_key_date',
                source_label=key_date.get_date_type_display(),
                subtitle=key_date.get_date_type_display(),
                edit_url=reverse('edit_matter_key_date', args=[key_date.id]),
                location=key_date.location,
                notes=key_date.notes,
            ))

    latest_assessment_subquery = RiskAssessment.objects.filter(
        matter=OuterRef('pk')
    ).order_by('-due_diligence_date').values('due_diligence_date')[:1]
    latest_monitoring_subquery = OngoingMonitoring.objects.filter(
        file_number=OuterRef('pk')
    ).order_by('-date_due_diligence_conducted').values('date_due_diligence_conducted')[:1]

    annotated_matters = base_matters.annotate(
        latest_assessment_date=Subquery(latest_assessment_subquery),
        latest_monitoring_date=Subquery(latest_monitoring_subquery),
    )

    if 'risk_assessment_due' in selected_sources:
        for matter in annotated_matters:
            if matter.latest_assessment_date:
                continue
            due_date = timezone.localdate()
            if date_from <= due_date <= date_to:
                events.append(_make_central_calendar_event(
                    event_date=due_date,
                    title='Risk assessment due',
                    matter=matter,
                    source='risk_assessment_due',
                    source_label='Risk due',
                    subtitle='Initial matter risk assessment',
                    edit_url=reverse('add_risk_assessment',
                                     args=[matter.file_number]),
                ))

    if 'ongoing_monitoring_due' in selected_sources:
        for matter in annotated_matters:
            last_review_date = matter.latest_monitoring_date or matter.latest_assessment_date
            if not last_review_date:
                continue
            due_date = last_review_date + relativedelta(years=1)
            if date_from <= due_date <= date_to:
                events.append(_make_central_calendar_event(
                    event_date=due_date,
                    title='Ongoing monitoring due',
                    matter=matter,
                    source='ongoing_monitoring_due',
                    source_label='Monitoring',
                    subtitle='Annual ongoing monitoring review',
                    edit_url=reverse('add_ongoing_monitoring',
                                     args=[matter.file_number]),
                ))

    if 'aml_due' in selected_sources:
        for matter in base_matters:
            parties = [
                (matter.client1, 'Client'),
                (matter.client2, 'Client'),
                (matter.authorised_party1, 'Authorised party'),
                (matter.authorised_party2, 'Authorised party'),
            ]
            for party, party_type in parties:
                if not party:
                    continue
                due_date = (
                    party.date_of_last_aml + relativedelta(years=1)
                    if party.date_of_last_aml
                    else timezone.localdate()
                )
                if date_from <= due_date <= date_to:
                    events.append(_make_central_calendar_event(
                        event_date=due_date,
                        title=f'AML check due - {party.name}',
                        matter=matter,
                        source='aml_due',
                        source_label='AML',
                        subtitle=party_type,
                        person_label=party.name,
                        edit_url=reverse('edit_client', args=[party.id]) if party_type == 'Client' else reverse(
                            'edit_authorised_party', args=[party.id]),
                    ))

    if 'id_expiry' in selected_sources:
        client_ids = set(base_matters.values_list('client1_id', flat=True))
        client_ids.update(
            base_matters.exclude(client2_id__isnull=True).values_list(
                'client2_id', flat=True)
        )
        client_ids.discard(None)
        documents = ClientKeyDocument.objects.filter(
            client_id__in=client_ids,
            expiry_date__gte=date_from,
            expiry_date__lte=date_to,
        ).select_related('client').order_by('expiry_date', 'client__name', 'category')
        client_matters = {}
        for matter in base_matters:
            client_matters.setdefault(matter.client1_id, []).append(matter)
            if matter.client2_id:
                client_matters.setdefault(matter.client2_id, []).append(matter)
        for document in documents:
            for matter in client_matters.get(document.client_id, []):
                events.append(_make_central_calendar_event(
                    event_date=document.expiry_date,
                    title=f'{document.get_category_display()} expiring - {document.client.name}',
                    matter=matter,
                    source='id_expiry',
                    source_label='ID expiry',
                    subtitle=document.document_type or document.get_category_display(),
                    person_label=document.client.name,
                    edit_url=reverse('edit_client', args=[document.client_id]),
                    notes=document.document_reference,
                ))

        annual_review_documents = ClientKeyDocument.objects.filter(
            client_id__in=client_ids,
            verified_on__isnull=False,
        ).select_related('client').order_by('verified_on', 'client__name', 'category')
        for document in annual_review_documents:
            review_due_date = document.verified_on + relativedelta(years=1)
            if not (date_from <= review_due_date <= date_to):
                continue
            for matter in client_matters.get(document.client_id, []):
                events.append(_make_central_calendar_event(
                    event_date=review_due_date,
                    title=f'{document.get_category_display()} review due - {document.client.name}',
                    matter=matter,
                    source='id_expiry',
                    source_label='ID review',
                    subtitle=document.document_type or document.get_category_display(),
                    person_label=document.client.name,
                    edit_url=reverse('edit_client', args=[document.client_id]),
                    notes=document.document_reference,
                ))

    selected_staff_id = selected_fee_earner or None

    if 'staff_leave' in selected_sources or 'office_closure' in selected_sources:
        holiday_records = HolidayRecord.objects.select_related('employee').filter(
            start_date__date__lte=date_to,
            end_date__date__gte=date_from,
            approved=True,
        ).order_by('start_date', 'employee__username')
        if selected_staff_id:
            holiday_records = holiday_records.filter(
                employee_id=selected_staff_id)

        for holiday in holiday_records:
            is_office_closure = holiday.reason == 'Office Closure'
            source = 'office_closure' if is_office_closure else 'staff_leave'
            if source not in selected_sources:
                continue

            if is_office_closure:
                title = f'Office closure - {holiday.employee}'
                source_label = 'Office closure'
                subtitle = holiday.reason
            else:
                leave_type = 'Annual leave' if holiday.type == 'Paid' else 'Unpaid leave'
                title = f'{holiday.employee} off'
                source_label = leave_type
                subtitle = holiday.reason or leave_type

            for event_date in _iter_event_dates(holiday.start_date, holiday.end_date, date_from, date_to):
                events.append(_make_central_calendar_event(
                    event_date=event_date,
                    title=title,
                    source=source,
                    source_label=source_label,
                    subtitle=subtitle,
                    person_label=str(holiday.employee),
                    edit_url=reverse('profile_page'),
                    detail_url=reverse('profile_page'),
                ))

    if 'sickness' in selected_sources:
        sickness_records = SicknessRecord.objects.select_related('employee').filter(
            start_date__date__lte=date_to,
        ).filter(
            Q(end_date__isnull=True) | Q(end_date__date__gte=date_from)
        ).order_by('start_date', 'employee__username')
        if selected_staff_id:
            sickness_records = sickness_records.filter(
                employee_id=selected_staff_id)
        elif not request.user.is_manager:
            sickness_records = sickness_records.filter(employee=request.user)

        for sickness in sickness_records:
            for event_date in _iter_event_dates(sickness.start_date, sickness.end_date or sickness.start_date, date_from, date_to):
                events.append(_make_central_calendar_event(
                    event_date=event_date,
                    title=f'{sickness.employee} off sick',
                    source='sickness',
                    source_label='Sickness',
                    subtitle=sickness.description,
                    person_label=str(sickness.employee),
                    edit_url=reverse('profile_page'),
                    detail_url=reverse('profile_page'),
                    notes=sickness.description,
                ))

    if 'bank_holiday' in selected_sources:
        holiday_list = holidays.country_holidays(
            'GB',
            subdiv='ENG',
            years=range(date_from.year, date_to.year + 1),
        )
        for holiday_date, holiday_name in holiday_list.items():
            if date_from <= holiday_date <= date_to:
                events.append(_make_central_calendar_event(
                    event_date=holiday_date,
                    title=f'Bank holiday - {holiday_name}',
                    source='bank_holiday',
                    source_label='Bank holiday',
                    subtitle=holiday_name,
                    notes=holiday_name,
                ))

    events = sorted(events, key=lambda event: (
        event['date'],
        event['time'] or time.min,
        event['matter'].file_number if event['matter'] else '',
        event['title'],
    ))

    for index, event in enumerate(events, start=1):
        event['list_id'] = f'calendar-entry-{index}'

    calendar_events = {}
    for event in events:
        if date_from <= event['date'] <= date_to:
            calendar_events.setdefault(event['date'], []).append(event)

    calendar_start = date_from - timedelta(days=date_from.weekday())
    calendar_end = date_to + timedelta(days=(6 - date_to.weekday()))
    calendar_weeks = []
    current_day = calendar_start
    while current_day <= calendar_end:
        week_days = []
        for offset in range(7):
            day = current_day + timedelta(days=offset)
            is_in_range = date_from <= day <= date_to
            day_events = calendar_events.get(day, []) if is_in_range else []
            week_days.append({
                'date': day,
                'is_current_month': is_in_range,
                'is_today': day == timezone.localdate(),
                'events': day_events,
                'visible_events': day_events[:2],
                'hidden_event_count': max(len(day_events) - 2, 0),
                'source_counts': _source_counts_for_day(day_events),
            })
        calendar_weeks.append(week_days)
        current_day += timedelta(days=7)

    calendar_months = [{
        'label': f"{date_from.strftime('%d %b %Y')} - {date_to.strftime('%d %b %Y')}",
        'weeks': calendar_weeks,
    }]

    if selected_view_months == 1:
        period_delta = (date_to - date_from) + timedelta(days=1)
        previous_start = date_from - period_delta
        previous_end = date_to - period_delta
        next_start = date_from + period_delta
        next_end = date_to + period_delta
    else:
        previous_start = date_from - relativedelta(months=selected_view_months)
        previous_end = date_to - relativedelta(months=selected_view_months)
        next_start = date_from + relativedelta(months=selected_view_months)
        next_end = date_to + relativedelta(months=selected_view_months)
    month_label = f"{date_from.strftime('%d %b %Y')} - {date_to.strftime('%d %b %Y')}"

    context = {
        'key_dates': events,
        'calendar_months': calendar_months,
        'current_year': date_from.year,
        'current_month': date_from.month,
        'month_label': month_label,
        'previous_month_query': _range_query_string(request, previous_start, previous_end, selected_view_months),
        'next_month_query': _range_query_string(request, next_start, next_end, selected_view_months),
        'fee_earners': CustomUser.objects.filter(
            is_matter_fee_earner=True,
            is_active=True,
        ).order_by('first_name', 'last_name', 'username'),
        'matter_types': MatterType.objects.all().order_by('type'),
        'file_statuses': FileStatus.objects.all().order_by('status'),
        'matters': WIP.objects.select_related('client1', 'client2').order_by('file_number'),
        'date_type_choices': MatterKeyDate.DATE_TYPE_CHOICES,
        'source_choices': KEY_DATE_SOURCE_CHOICES,
        'view_month_choices': VIEW_MONTH_CHOICES,
        'filters': {
            'fee_earner': selected_fee_earner,
            'matter_type': selected_matter_type,
            'file_status': selected_file_status or 'all',
            'matter': selected_matter,
            'date_type': selected_date_type,
            'sources': selected_sources,
            'matter_search': matter_search,
            'client_search': client_search,
            'date_from': date_from.isoformat(),
            'date_to': date_to.isoformat(),
            'view_months': selected_view_months,
        },
    }
    return context


@login_required
def central_key_dates(request):
    context = _build_central_key_dates_context(request)
    return render(request, 'key_dates.html', context)


def _key_date_event_row(event):
    matter = event['matter']
    client = ''
    fee_earner = ''
    matter_description = ''
    file_number = ''
    if matter:
        file_number = matter.file_number
        matter_description = matter.matter_description or ''
        client = matter.client1.name
        if matter.client2:
            client = f'{client} / {matter.client2.name}'
        fee_earner = str(matter.fee_earner or '')

    return [
        event['date'].isoformat(),
        event['date'].strftime('%A'),
        event['time'].strftime('%H:%M') if event['time'] else '',
        event['source_label'],
        event['title'],
        file_number,
        matter_description,
        client,
        fee_earner,
        event['person_label'],
        event['location'],
        event['subtitle'],
        event['notes'],
    ]


@login_required
def download_central_key_dates(request, export_format, export_kind):
    if export_format not in {'csv', 'pdf'} or export_kind not in {'calendar', 'list'}:
        raise Http404

    context = _build_central_key_dates_context(request)
    filename_base = f"key-dates-{export_kind}-{context['filters']['date_from']}-to-{context['filters']['date_to']}"

    if export_format == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename_base}.csv"'
        writer = csv.writer(response)
        if export_kind == 'calendar':
            writer.writerow(['Date', 'Day', 'Total Entries',
                            'Source Counts', 'Entries'])
            for calendar_month in context['calendar_months']:
                for week in calendar_month['weeks']:
                    for day in week:
                        if not day['events']:
                            continue
                        source_counts = ', '.join(
                            f"{source_count['label']}: {source_count['count']}"
                            for source_count in day['source_counts']
                        )
                        entries = '; '.join(
                            f"{event['source_label']} - {event['title']}"
                            for event in day['events']
                        )
                        writer.writerow([
                            day['date'].isoformat(),
                            day['date'].strftime('%A'),
                            len(day['events']),
                            source_counts,
                            entries,
                        ])
        else:
            writer.writerow([
                'Date',
                'Day',
                'Time',
                'Source',
                'Title',
                'Matter',
                'Matter Description',
                'Client',
                'Fee Earner',
                'Person',
                'Location',
                'Detail',
                'Notes',
            ])
            for event in context['key_dates']:
                writer.writerow(_key_date_event_row(event))
        return response

    context['export_kind'] = export_kind
    html = render_to_string('key_dates_export.html', context, request=request)
    pdf = HTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf()
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename_base}.pdf"'
    return response


def get_modifications_by_object(model, object_ids):
    if not object_ids:
        return {}

    content_type = ContentType.objects.get_for_model(model)
    modifications = Modifications.objects.filter(
        content_type=content_type,
        object_id__in=object_ids
    ).select_related('modified_by').order_by('timestamp')

    modifications_by_object = {}
    for modification in modifications:
        modifications_by_object.setdefault(
            modification.object_id, []).append(modification)

    return modifications_by_object


def get_file_logs(file_number, limit=None):
    file = WIP.objects.select_related(
        'created_by',
        'client1__created_by',
        'client2__created_by',
        'authorised_party1__created_by',
        'authorised_party2__created_by',
        'other_side__created_by',
    ).filter(file_number=file_number).first()
    """
    Log: {datetime:datetime, description, user, type_of_data}
    """
    if not file:
        return []

    logs = []
    logs.append({'timestamp': file.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': 'Matter created.',
                 'user': file.created_by,
                 'type': 'file_info'})

    wip_modifications = get_modifications_by_object(WIP, [file.id])
    for modification in wip_modifications.get(file.id, []):
        logs.append({
            'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
            'desc': 'File updated.',
            'changes_list': build_change_items(modification.changes),
            'user': modification.modified_by.username if modification.modified_by else None,
            'type': 'file_info'
        })

    client_ids = [file.client1_id]
    if file.client2_id:
        client_ids.append(file.client2_id)
    client_modifications = get_modifications_by_object(
        ClientContactDetails, client_ids)

    for modification in client_modifications.get(file.client1_id, []):
        logs.append({
            'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
            'desc': f'Client ({file.client1}) updated.',
            'changes_list': build_change_items(modification.changes),
            'user': modification.modified_by.username if modification.modified_by else None,
            'type': 'client_info'
        })
    logs.append({'timestamp': file.client1.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                 'desc': f'Client ({file.client1}) Created.',
                 'user': file.client1.created_by,
                 'type': 'client_info'})

    if file.client2:
        logs.append({'timestamp': file.client2.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Client ({file.client2}) Created.',
                     'user': file.client2.created_by,
                     'type': 'client_info'})
        for modification in client_modifications.get(file.client2_id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Client ({file.client2}) updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'client_info'
            })

    authorised_party_ids = [
        party_id for party_id in [file.authorised_party1_id, file.authorised_party2_id]
        if party_id
    ]
    authorised_party_modifications = get_modifications_by_object(
        AuthorisedParties, authorised_party_ids)

    if file.authorised_party1:
        logs.append({'timestamp': file.authorised_party1.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Authorised Party {file.authorised_party1} Created.',
                     'user': file.authorised_party1.created_by,
                     'type': 'authorised_party_info'})
        for modification in authorised_party_modifications.get(file.authorised_party1_id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Authorised Party ({file.authorised_party1}) updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'authorised_party_info'
            })

    if file.authorised_party2:
        logs.append({'timestamp': file.authorised_party2.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Authorised Party ({file.authorised_party2.name}) Created.',
                     'user': file.authorised_party2.created_by,
                     'type': 'authorised_party_info'})
        for modification in authorised_party_modifications.get(file.authorised_party2_id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Authorised Party ({file.authorised_party2}) updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'authorised_party_info'
            })

    if file.other_side:
        logs.append({'timestamp': file.other_side.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Other Side ({file.other_side}) Created.',
                     'user': file.other_side.created_by,
                     'type': 'other_side_info'})
        other_side_modifications = get_modifications_by_object(
            OthersideDetails, [file.other_side_id])
        for modification in other_side_modifications.get(file.other_side_id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Other Side ({file.other_side}) updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'other_side_info'
            })
    emails = MatterEmails.objects.filter(file_number=file.id)
    for email in emails:
        sender = json.loads(email.sender)
        receiver = json.loads(email.receiver)
        desc = f'Email to {
            receiver} - {email.subject}' if email.is_sent else f'Email from {sender} - {email.subject}'
        logs.append({'timestamp': email.time.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': desc, 'user': 'Automatic System',
                    'type': 'email'})

    attendance_notes = list(MatterAttendanceNotes.objects.filter(
        file_number=file.id).select_related('created_by'))
    attendance_note_modifications = get_modifications_by_object(
        MatterAttendanceNotes, [note.id for note in attendance_notes])
    for note in attendance_notes:
        note_subject_with_charge_status = (
            f'{note.subject_line} (N/C)' if not note.is_charged else note.subject_line
        )
        logs.append({'timestamp': note.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Attendance note created - {note_subject_with_charge_status}',
                     'user': note.created_by,
                    'type': 'attendance_note'})

        for modification in attendance_note_modifications.get(note.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Attendance note ({note.date.strftime('%d/%m/%Y')}) updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'attendance_note'
            })

    letters = list(MatterLetters.objects.filter(
        file_number=file.id).select_related('created_by'))
    letter_modifications = get_modifications_by_object(
        MatterLetters, [letter.id for letter in letters])
    for letter in letters:
        logs.append({'timestamp': letter.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Letter entered - {letter.subject_line}',
                     'user': letter.created_by,
                    'type': 'letter'})

        for modification in letter_modifications.get(letter.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Letter ({letter.date.strftime('%d/%m/%Y')}) updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'letter'
            })

    next_work = list(NextWork.objects.filter(
        file_number=file.id).select_related('person', 'created_by'))
    next_work_modifications = get_modifications_by_object(
        NextWork, [work.id for work in next_work])
    for work in next_work:
        logs.append({'timestamp': work.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Next work created - {work.task} - for {work.person}',
                     'user': work.created_by,
                    'type': 'next_work'
                     })
        for modification in next_work_modifications.get(work.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Next Work updated (task: {work.task}, completed: {work.completed}).',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'next_work'
            })

    last_work = list(LastWork.objects.filter(
        file_number=file.id).select_related('person', 'created_by'))
    last_work_modifications = get_modifications_by_object(
        LastWork, [work.id for work in last_work])
    for work in last_work:
        logs.append({'timestamp': work.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Last work created - {work.task} - done by {work.person}',
                     'user': work.created_by,
                    'type': 'last_work'
                     })
        for modification in last_work_modifications.get(work.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Last Work updated (task: {work.task}).',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'last_work'
            })

    pmts_slips = list(PmtsSlips.objects.filter(
        file_number=file.id).select_related('created_by'))
    pmts_slip_modifications = get_modifications_by_object(
        PmtsSlips, [slip.id for slip in pmts_slips])
    for slip in pmts_slips:
        desc = f'Pink slip for amount £{
            slip.amount} - {slip.description}' if slip.is_money_out else f'Blue slip for amount £{slip.amount} - {slip.description}'
        logs.append({'timestamp': slip.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': desc,
                     'user': slip.created_by,
                    'type': 'pmts_slip'})
        for modification in pmts_slip_modifications.get(slip.id, []):
            slip_type = 'Pink slip' if slip.is_money_out else 'Blue slip'
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'{slip_type} updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'pmts_slip'
            })

    green_slips = list(LedgerAccountTransfers.objects.filter(
        Q(file_number_from=file.id) | Q(file_number_to=file.id)
    ).select_related('file_number_from', 'file_number_to', 'created_by'))
    green_slip_modifications = get_modifications_by_object(
        LedgerAccountTransfers, [slip.id for slip in green_slips])
    for slip in green_slips:
        desc = f'Green slip for amount £{
            slip.amount} - From: {slip.file_number_from} To: {slip.file_number_to} {slip.description}'
        logs.append({'timestamp': slip.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': desc,
                     'user': slip.created_by,
                    'type': 'green_slip'})
        for modification in green_slip_modifications.get(slip.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': 'Green slip updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'green_slip'
            })

    invoices = list(Invoices.objects.filter(
        file_number=file.id).select_related('created_by'))
    invoice_modifications = get_modifications_by_object(
        Invoices, [invoice.id for invoice in invoices])
    for invoice in invoices:
        desc = f'Invoice created for amount(s) {invoice.our_costs}'
        logs.append({'timestamp': invoice.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': desc,
                     'user': invoice.created_by,
                    'type': 'invoice'})
        for modification in invoice_modifications.get(invoice.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': 'Invoice updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'invoice'
            })

    credit_notes = list(CreditNote.objects.filter(
        file_number=file.id).select_related('invoice', 'created_by'))
    credit_note_modifications = get_modifications_by_object(
        CreditNote, [credit_note.id for credit_note in credit_notes])
    status_labels = dict(CreditNote.STATUSES)
    for credit_note in credit_notes:
        status_label = status_labels.get(
            credit_note.status, credit_note.status)
        invoice_number = credit_note.invoice.invoice_number or "Draft"
        desc = (
            f'Credit note for Invoice {invoice_number} of £{credit_note.amount} '
            f'created. Status: {status_label}'
        )
        logs.append({
            'timestamp': credit_note.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
            'desc': desc,
            'user': credit_note.created_by,
            'type': 'credit_note'
        })
        for modification in credit_note_modifications.get(credit_note.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': 'Credit note updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'credit_note'
            })

    risk_assessment = RiskAssessment.objects.filter(
        matter=file.id).select_related('due_diligence_signed_by').first()
    if risk_assessment:
        logs.append({'timestamp': risk_assessment.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                    'desc': f'Risk Assessment completed',
                     'user': risk_assessment.due_diligence_signed_by,
                     'type': 'risk_assessment'})
        risk_assessment_modifications = get_modifications_by_object(
            RiskAssessment, [risk_assessment.id])
        for modification in risk_assessment_modifications.get(risk_assessment.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': 'Risk Assessment updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'risk_assessment'
            })
    ongoing_monitoring = list(OngoingMonitoring.objects.filter(
        file_number=file.id).select_related('created_by'))
    ongoing_monitoring_modifications = get_modifications_by_object(
        OngoingMonitoring, [obj.id for obj in ongoing_monitoring])
    for obj in ongoing_monitoring:
        logs.append({'timestamp': obj.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                     'desc': f'Ongoing Monitoring done.',
                     'user': obj.created_by,
                     'type': 'ongoing_monitoring'})
        for modification in ongoing_monitoring_modifications.get(obj.id, []):
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': 'Ongoing Monitoring updated.',
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'ongoing_monitoring'
            })

    matter_key_dates = list(MatterKeyDate.objects.filter(
        matter=file.id).select_related('created_by'))
    key_date_modifications = get_modifications_by_object(
        MatterKeyDate, [key_date.id for key_date in matter_key_dates])
    for key_date in matter_key_dates:
        modifications = key_date_modifications.get(key_date.id, [])
        if not modifications:
            logs.append({
                'timestamp': key_date.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Key date added - {key_date.title}',
                'user': key_date.created_by,
                'type': 'key_date'
            })
        for modification in modifications:
            is_create = isinstance(
                modification.changes, dict) and modification.changes.get('created')
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': (
                    f'Key date added - {key_date.title}'
                    if is_create else f'Key date updated - {key_date.title}'
                ),
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'key_date'
            })

    key_document_client_ids = [file.client1_id]
    if file.client2_id:
        key_document_client_ids.append(file.client2_id)
    key_documents = list(ClientKeyDocument.objects.filter(
        client_id__in=key_document_client_ids).select_related('client', 'verified_by'))
    key_document_modifications = get_modifications_by_object(
        ClientKeyDocument, [document.id for document in key_documents])
    for document in key_documents:
        modifications = key_document_modifications.get(document.id, [])
        if not modifications:
            logs.append({
                'timestamp': document.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Key document added - {snapshot_key_document(document)}',
                'user': document.verified_by,
                'type': 'key_document'
            })
        for modification in modifications:
            is_create = isinstance(
                modification.changes, dict) and modification.changes.get('created')
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': (
                    f'Key document added - {snapshot_key_document(document)}'
                    if is_create else f'Key document updated - {snapshot_key_document(document)}'
                ),
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'key_document'
            })

    matter_file_reviews = list(MatterFileReview.objects.filter(
        matter=file.id).select_related('created_by'))
    matter_file_review_modifications = get_modifications_by_object(
        MatterFileReview, [review.id for review in matter_file_reviews])
    for review in matter_file_reviews:
        review_label = review.date_review_completed or review.date_reviewed or review.id
        modifications = matter_file_review_modifications.get(review.id, [])
        if not modifications:
            logs.append({
                'timestamp': review.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Matter file review added ({review_label})',
                'user': review.created_by,
                'type': 'matter_file_review'
            })
        for modification in modifications:
            is_create = isinstance(
                modification.changes, dict) and modification.changes.get('created')
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': (
                    f'Matter file review added ({review_label})'
                    if is_create else f'Matter file review updated ({review_label})'
                ),
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'matter_file_review'
            })

    bundles = list(Bundle.objects.filter(
        file_number=file.id).select_related('created_by'))
    bundle_modifications = get_modifications_by_object(
        Bundle, [bundle.id for bundle in bundles])
    for bundle in bundles:
        modifications = bundle_modifications.get(bundle.id, [])
        if not modifications:
            logs.append({
                'timestamp': bundle.created_at.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': f'Bundle created - {bundle.name}',
                'user': bundle.created_by,
                'type': 'bundle'
            })
        for modification in modifications:
            is_create = (
                isinstance(modification.changes, dict)
                and modification.changes.get('event', {}).get('new_value') == 'Bundle created'
            )
            logs.append({
                'timestamp': modification.timestamp.strftime("%d/%m/%Y %H:%M:%S"),
                'desc': (
                    f'Bundle created - {bundle.name}'
                    if is_create else f'Bundle updated - {bundle.name}'
                ),
                'changes_list': build_change_items(modification.changes),
                'user': modification.modified_by.username if modification.modified_by else None,
                'type': 'bundle'
            })

    sorted_logs = sorted(logs, key=lambda x: datetime.strptime(
        x['timestamp'], '%d/%m/%Y %H:%M:%S'), reverse=True)

    if limit:
        return sorted_logs[:limit]
    return sorted_logs


def add_new_client(request_post_copy, client_prefix, user):
    name = request_post_copy[f'ClientName{client_prefix}']
    is_business = f'Client{client_prefix}IsBusiness' in request_post_copy
    dob = request_post_copy[f'Client{client_prefix}DOB']
    occupation = request_post_copy[f'Client{client_prefix}Occupation']

    address_line1 = request_post_copy[f'Client{client_prefix}AddressLine1']
    address_line2 = request_post_copy[f'Client{client_prefix}AddressLine2']
    county = request_post_copy[f'Client{client_prefix}County']
    postcode = request_post_copy[f'Client{client_prefix}Postcode']
    email = request_post_copy[f'Client{client_prefix}Email']
    contact_number = request_post_copy[f'Client{client_prefix}ContactNumber']
    date_of_last_aml = request_post_copy[f'Client{client_prefix}AMLCheckDate']
    id_verified = f'IDVer{client_prefix}' in request_post_copy
    terms_of_engagement_signed = f'Client{client_prefix}TermsOfEngagementSigned' in request_post_copy
    ncba_signed = f'Client{client_prefix}NCBASigned' in request_post_copy
    pep_signed = f'Client{client_prefix}PEPSigned' in request_post_copy
    source_of_funds_signed = f'Client{client_prefix}SourceOfFundsSigned' in request_post_copy

    client_contact = ClientContactDetails(
        name=name,
        is_business=is_business,
        dob=dob if dob != '' else None,
        occupation=occupation,
        address_line1=address_line1,
        address_line2=address_line2,
        county=county,
        postcode=postcode,
        email=email,
        contact_number=contact_number,
        date_of_last_aml=date_of_last_aml if date_of_last_aml != '' else None,
        id_verified=id_verified,
        terms_of_engagement_signed=terms_of_engagement_signed,
        ncba_signed=ncba_signed,
        pep_signed=pep_signed,
        source_of_funds_signed=source_of_funds_signed,
        created_by=user
    )

    client_contact.save()
    add_client_key_documents_from_post(
        request_post_copy, client_prefix, client_contact, user)

    return client_contact.id


def add_client_key_documents_from_post(request_post_copy, client_prefix, client, user):
    document_configs = [
        ('ProofOfID', 'proof_of_id'),
        ('ProofOfAddress', 'proof_of_address'),
    ]

    for field_prefix, category in document_configs:
        document_type = request_post_copy.get(
            f'Client{client_prefix}{field_prefix}DocumentType', '').strip()
        document_reference = request_post_copy.get(
            f'Client{client_prefix}{field_prefix}DocumentReference', '').strip()
        issue_date = request_post_copy.get(
            f'Client{client_prefix}{field_prefix}IssueDate', '')
        expiry_date = request_post_copy.get(
            f'Client{client_prefix}{field_prefix}ExpiryDate', '')
        verified_on = request_post_copy.get(
            f'Client{client_prefix}{field_prefix}VerifiedOn', '')
        notes = request_post_copy.get(
            f'Client{client_prefix}{field_prefix}Notes', '').strip()

        if not any([document_type, document_reference, issue_date, expiry_date, verified_on, notes]):
            continue

        document = ClientKeyDocument.objects.create(
            client=client,
            category=category,
            document_type=document_type,
            document_reference=document_reference,
            issue_date=issue_date if issue_date != '' else None,
            expiry_date=expiry_date if expiry_date != '' else None,
            verified_on=verified_on if verified_on != '' else None,
            verified_by=user if verified_on != '' else None,
            notes=notes,
        )
        log_created(user, document, snapshot_key_document(document))


def add_new_authorised_party(request_post_copy, ap_prefix, user):
    name = request_post_copy[f'APName{ap_prefix}']
    relationship_to_client = request_post_copy[f'AP{ap_prefix}RelationshipToC']
    address_line1 = request_post_copy[f'AP{ap_prefix}AddressLine1']
    address_line2 = request_post_copy[f'AP{ap_prefix}AddressLine2']
    county = request_post_copy[f'AP{ap_prefix}County']
    postcode = request_post_copy[f'AP{ap_prefix}Postcode']
    email = request_post_copy[f'AP{ap_prefix}Email']
    contact_number = request_post_copy[f'AP{ap_prefix}ContactNumber']
    id_check = f'AP{ap_prefix}IDCheck' in request_post_copy
    date_of_id_check = request_post_copy[f'AP{ap_prefix}IDCheckDate']
    date_of_last_aml = request_post_copy[f'AP{ap_prefix}AMLCheckDate']

    try:
        authorised_party = AuthorisedParties(
            name=name,
            relationship_to_client=relationship_to_client,
            address_line1=address_line1,
            address_line2=address_line2,
            county=county,
            postcode=postcode,
            email=email,
            contact_number=contact_number,
            id_check=id_check,
            date_of_id_check=date_of_id_check if date_of_id_check != '' else None,
            date_of_last_aml=date_of_last_aml if date_of_last_aml != '' else None,
            created_by=user
        )
    except Exception as e:
        print(f'Error in adding Authorised Party {ap_prefix}: {str(e)}')

    authorised_party.save()

    return authorised_party.id


def add_new_otherside_details(request_post_copy, user):

    name = request_post_copy['OSName']
    address_line1 = request_post_copy['OSAddressLine1']
    address_line2 = request_post_copy['OSAddressLine2']
    county = request_post_copy['OSCounty']
    postcode = request_post_copy['OSPostcode']
    email = request_post_copy['OSEmail']
    contact_number = request_post_copy['OSContactNumber']
    solicitors = request_post_copy['OSSolicitors']
    solicitors_email = request_post_copy['OSSolicitorsEmail']

    otherside_details = OthersideDetails(
        name=name,
        address_line1=address_line1,
        address_line2=address_line2,
        county=county,
        postcode=postcode,
        email=email,
        contact_number=contact_number,
        solicitors=solicitors,
        solicitors_email=solicitors_email,
        created_by=user
    )

    otherside_details.save()

    # Return the ID of the newly added OthersideDetails
    return otherside_details.id


def preprocess_form_data(post_data):
    post_copy = post_data.copy()

    post_copy['client2'] = None if post_copy['client2'] == '0' else post_copy['client2']

    post_copy['authorised_party1'] = None if post_copy['authorised_party1'] == '0' else post_copy['authorised_party1']
    post_copy['authorised_party2'] = None if post_copy['authorised_party2'] == '0' else post_copy['authorised_party2']
    post_copy['other_side'] = None if post_copy['other_side'] == '0' else post_copy['other_side']

    return post_copy


def update_checkbox_values(data, *fields):
    for field in fields:
        data[field] = field in data

    return data


def get_aml_checks_due_from_wips(wips, threshold_date, user=None, sort_by='date'):
    base_filter = Q(file_status__status='Open')
    if user is not None:
        base_filter &= Q(fee_earner=user)

    relation_configs = [
        ('client1', 'Client', 'edit_client'),
        ('client2', 'Client', 'edit_client'),
        ('authorised_party1', 'Authorised Party', 'edit_authorised_party'),
        ('authorised_party2', 'Authorised Party', 'edit_authorised_party'),
    ]

    results = {}
    for relation, entity_type, edit_url_name in relation_configs:
        relation_results = wips.filter(
            base_filter & Q(
                **{f'{relation}__date_of_last_aml__lte': threshold_date})
        ).annotate(
            entity_id=F(f'{relation}__id'),
            entity_name=F(f'{relation}__name'),
            date_of_last_aml=F(f'{relation}__date_of_last_aml'),
        ).values('entity_id', 'entity_name', 'date_of_last_aml')

        for result in relation_results:
            key = (entity_type, result['entity_id'])
            results[key] = {
                'entity_id': result['entity_id'],
                'entity_name': result['entity_name'],
                'entity_type': entity_type,
                'date_of_last_aml': result['date_of_last_aml'],
                'edit_url': reverse(edit_url_name, args=[result['entity_id']]),
            }

    if sort_by == 'name':
        return sorted(results.values(), key=lambda x: (x['entity_name'] or '', x['entity_type']))
    return sorted(results.values(), key=lambda x: x['date_of_last_aml'])


def get_standard_data():
    fee_earners = CustomUser.objects.filter(is_matter_fee_earner=True)
    file_status = FileStatus.objects.all()
    file_locations = FileLocation.objects.all()
    matter_types = MatterType.objects.all().order_by('type')
    clients = ClientContactDetails.objects.all().order_by('name')
    authorised_parties = AuthorisedParties.objects.all().order_by('name')
    othersides = OthersideDetails.objects.all().order_by('name')

    form_data = {
        'fee_earners': fee_earners,
        'file_status': file_status,
        'file_locations': file_locations,
        'matter_types': matter_types,
        'clients': clients,
        'authorised_parties': authorised_parties,
        'othersides': othersides,
    }

    return form_data


@login_required
def open_new_file_page(request):
    form_data = get_standard_data()

    if request.method == 'POST':
        request_post_copy = preprocess_form_data(request.POST)
        try:
            if request_post_copy['client1'] == '-1':
                request_post_copy['client1'] = add_new_client(
                    request_post_copy, 1, request.user)

            if request_post_copy['client2'] == '-1':
                request_post_copy['client2'] = add_new_client(
                    request_post_copy, 2, request.user)

            if request_post_copy['authorised_party1'] == '-1':
                request_post_copy['authorised_party1'] = add_new_authorised_party(
                    request_post_copy, 1, request.user)

            if request_post_copy['authorised_party2'] == '-1':
                request_post_copy['authorised_party2'] = add_new_authorised_party(
                    request_post_copy, 2, request.user)

            if request_post_copy['other_side'] == '-1':
                request_post_copy['other_side'] = add_new_otherside_details(
                    request_post_copy, request.user)

            request_post_copy['created_by'] = request.user

            form = OpenFileForm(request_post_copy)
            if form.is_valid():
                instance = form.save()
                messages.success(request, 'File opened successfully.')
                return redirect('index')
            else:
                # Add error messages to be displayed in the template
                for field, errors in form.errors.items():
                    for error in errors:
                        messages.error(
                            request, f"{form[field].label}: {error}")
                return render(request, 'open_file.html', {'form_data': form_data, 'form': form})

        except Exception as e:
            messages.error(request, f"Error during file opening: {str(e)}")
            return render(request, 'open_file.html', {'form_data': form_data})
    else:
        return render(request, 'open_file.html', {'form_data': form_data})


@login_required
def add_risk_assessment(request, file_number):
    try:
        matter = WIP.objects.get(file_number=file_number)
    except WIP.DoesNotExist:
        messages.error(
            request, 'Matter with the given file number does not exist.')
        return redirect('index')
    if request.method == 'POST':
        post_data = request.POST.copy()

        form = RiskAssessmentForm(post_data)

        if form.is_valid():
            risk_assessment = form.save()
            log_created(
                request.user,
                risk_assessment,
                f'Risk assessment for {risk_assessment.matter.file_number}',
            )
            messages.success(request, 'Risk Assessment successfully added.')
            return redirect('home', risk_assessment.matter.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:

        form = RiskAssessmentForm(initial={'matter': matter.id})

    return render(request, 'risk_assessment.html', {'form': form, 'file_number': file_number, 'title': 'Add'})


@login_required
def add_matter_file_review(request, file_number):
    matter = get_object_or_404(WIP, file_number=file_number)

    if request.method == 'POST':
        form = MatterFileReviewForm(request.POST)
        if form.is_valid():
            review = form.save(commit=False)
            review.matter = matter
            review.created_by = request.user
            review.save()
            log_created(
                request.user,
                review,
                f'Matter file review ({review.date_review_completed or review.date_reviewed or review.id})',
            )
            messages.success(request, 'Matter file review added successfully.')
            return redirect('home', file_number=matter.file_number)

        for field, errors in form.errors.items():
            for error in errors:
                messages.error(request, f"{form[field].label}: {error}")
    else:
        form = MatterFileReviewForm(initial={
            'client_matter_reference': matter.file_number,
            'file_reviewed_by': request.user.id,
            'date_reviewed': timezone.localdate(),
            'file_review_completed_by': request.user.id,
            'date_review_completed': timezone.localdate(),
        })

    return render(request, 'matter_file_review.html', {
        'form': form,
        'file_number': matter.file_number,
        'matter': matter,
        'title': 'Add',
        'matter_file_review_sections': MATTER_FILE_REVIEW_SECTIONS,
    })


@login_required
def edit_matter_file_review(request, id):
    review = get_object_or_404(MatterFileReview, id=id)
    duplicate_obj = copy.deepcopy(review)

    if request.method == 'POST':
        form = MatterFileReviewForm(request.POST, instance=review)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }

            form.save()

            for field in changed_fields:
                changes[field]['new_value'] = str(getattr(review, field))

            if changes:
                create_modification(
                    user=request.user,
                    modified_obj=review,
                    changes=changes
                )

            messages.success(
                request, 'Matter file review updated successfully.')
            return redirect('home', file_number=review.matter.file_number)

        for field, errors in form.errors.items():
            for error in errors:
                messages.error(request, f"{form[field].label}: {error}")
    else:
        form = MatterFileReviewForm(instance=review)

    return render(request, 'matter_file_review.html', {
        'form': form,
        'file_number': review.matter.file_number,
        'matter': review.matter,
        'title': 'Edit',
        'matter_file_review_sections': MATTER_FILE_REVIEW_SECTIONS,
    })


@login_required
def download_matter_file_review(request, id):
    review = get_object_or_404(MatterFileReview, id=id)
    sections = []
    for section in MATTER_FILE_REVIEW_SECTIONS:
        rows = []
        for row in section['rows']:
            rows.append({
                'question': row['question'],
                'bullets': row.get('bullets', []),
                'answer': getattr(review, row['answer_field']) or '---',
                'comments': getattr(review, row['comments_field']) or '---',
            })
        sections.append({
            'title': section['title'],
            'rows': rows,
        })

    html_string = render_to_string(
        'download_templates/matter_file_review.html', {
            'review': review,
            'sections': sections,
        })

    pdf_file = HTML(string=html_string).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = (
        f'attachment; filename="file_review_{review.matter.file_number}_{id}.pdf"'
    )
    return response


@login_required
def internal_pricing(request):
    pricing_content_type = ContentType.objects.get_for_model(PricingItem)

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'create':
            if not request.user.is_manager:
                messages.error(
                    request, 'Only managers can add pricing entries.')
                return redirect('internal_pricing')

            form = PricingItemForm(request.POST, user=request.user)
            if form.is_valid():
                pricing = form.save(commit=False)
                pricing.created_by = request.user
                pricing.updated_by = request.user
                pricing.save()
                create_modification(
                    user=request.user,
                    modified_obj=pricing,
                    changes={
                        'created': {
                            'old_value': '',
                            'new_value': 'Pricing entry created',
                        }
                    }
                )
                messages.success(request, 'Pricing entry added.')
                return redirect('internal_pricing')

            messages.error(request, 'Please correct the pricing entry form.')

        elif action == 'update':
            pricing = get_object_or_404(
                PricingItem, id=request.POST.get('pricing_id'))

            if not pricing.can_edit(request.user):
                messages.error(
                    request, 'Only managers can edit this pricing entry.')
                return redirect('internal_pricing')

            duplicate_obj = copy.deepcopy(pricing)
            form = PricingItemForm(
                request.POST, instance=pricing, user=request.user)
            if form.is_valid():
                changed_fields = form.changed_data
                changes = {}
                for field in changed_fields:
                    changes[field] = {
                        'old_value': str(getattr(duplicate_obj, field)),
                        'new_value': None,
                    }

                pricing = form.save(commit=False)
                pricing.updated_by = request.user
                pricing.save()

                for field in changed_fields:
                    changes[field]['new_value'] = str(getattr(pricing, field))

                if changes:
                    create_modification(
                        user=request.user,
                        modified_obj=pricing,
                        changes=changes
                    )
                    messages.success(request, 'Pricing entry updated.')
                else:
                    messages.info(request, 'No pricing changes were made.')

                return redirect('internal_pricing')

            messages.error(request, 'Please correct the pricing update form.')

    pricing_entries = PricingItem.objects.select_related(
        'matter_type', 'created_by', 'updated_by'
    )
    if not request.user.is_manager:
        pricing_entries = pricing_entries.filter(is_active=True)

    pricing_entries = list(pricing_entries)

    audit_entries = Modifications.objects.filter(
        content_type=pricing_content_type,
        object_id__in=[pricing.id for pricing in pricing_entries]
    ).select_related('modified_by').order_by('-timestamp')

    audit_by_pricing = {}
    for audit in audit_entries:
        audit_by_pricing.setdefault(audit.object_id, []).append(audit)

    for pricing in pricing_entries:
        pricing.edit_form = PricingItemForm(
            instance=pricing, user=request.user)
        pricing.audit_entries = audit_by_pricing.get(pricing.id, [])
        pricing.user_can_edit = pricing.can_edit(request.user)

    return render(request, 'internal_pricing.html', {
        'pricing_entries': pricing_entries,
        'create_form': PricingItemForm(user=request.user),
        'can_create': request.user.is_manager,
        'current_vat_rate_percent': CURRENT_VAT_RATE_PERCENT,
        'fee_earners': CustomUser.objects.filter(
            is_matter_fee_earner=True,
            is_active=True,
        ).select_related('hourly_rate').order_by('first_name', 'last_name', 'username'),
    })


matter_pricing = internal_pricing


@login_required
def edit_client(request, id):
    client = ClientContactDetails.objects.get(id=id)
    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(client)
        form = ClientForm(request.POST, instance=client)
        key_document_formset = ClientKeyDocumentFormSet(
            request.POST, instance=client, prefix='key_documents')
        if form.is_valid() and key_document_formset.is_valid():

            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()
            key_documents = key_document_formset.save(commit=False)
            for deleted_document in key_document_formset.deleted_objects:
                deleted_document.delete()
            for document in key_documents:
                if document.verified_on and not document.verified_by:
                    document.verified_by = request.user
                document.save()
            key_document_formset.save_m2m()

            for field in changed_fields:
                changes[field]['new_value'] = str(
                    getattr(client, field))

            if changes:
                create_modification(
                    user=request.user,
                    modified_obj=client,
                    changes=changes
                )

            audit_client_key_document_formset(
                request.user, key_document_formset, client)

            messages.success(
                request, 'Successfully updated Client. Please search for File Number.')
            return redirect('index')
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            if key_document_formset.errors:
                error_message += '\nPlease check the key document rows.'
            messages.error(request, error_message)
    else:
        form = ClientForm(instance=client)
        key_document_formset = ClientKeyDocumentFormSet(
            instance=client, prefix='key_documents')
    return render(request, 'edit_models.html', {
        'form': form,
        'key_document_formset': key_document_formset,
        'title': 'Client Information'
    })


@login_required
def edit_authorised_party(request, id):
    ap = AuthorisedParties.objects.get(id=id)
    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(ap)
        form = AuthorisedPartyForm(request.POST, instance=ap)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()
            for field in changed_fields:
                changes[field]['new_value'] = str(
                    getattr(ap, field))
            create_modification(
                user=request.user,
                modified_obj=ap,
                changes=changes
            )
            messages.success(
                request, 'Successfully updated Authorised Party. Please search for File Number.')
            return redirect('index')
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = AuthorisedPartyForm(instance=ap)
    return render(request, 'edit_models.html', {'form': form, 'title': 'Authorised Party Information'})


@login_required
def edit_otherside(request, id):
    os = OthersideDetails.objects.get(id=id)
    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(os)
        form = OtherSideForm(request.POST, instance=os)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()
            for field in changed_fields:
                changes[field]['new_value'] = str(
                    getattr(os, field))
            create_modification(
                user=request.user,
                modified_obj=os,
                changes=changes
            )
            messages.success(
                request, 'Successfully updated Other Side Details on all applicable files. Please search for a File Number you were working on.')
            return redirect('index')
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = OtherSideForm(instance=os)
    return render(request, 'edit_models.html', {'form': form, 'title': 'Other Side Details'})


@login_required
def edit_file(request, file_number):
    file = WIP.objects.filter(file_number=file_number).first()
    form_data = get_standard_data()

    if request.method == 'POST':
        try:

            request_post_copy = preprocess_form_data(request.POST)

            if request_post_copy['client2'] == '-1':
                request_post_copy['client2'] = add_new_client(
                    request_post_copy, 2, request.user)

            if request_post_copy['authorised_party1'] == '-1':
                request_post_copy['authorised_party1'] = add_new_authorised_party(
                    request_post_copy, 1, request.user)

            if request_post_copy['authorised_party2'] == '-1':
                request_post_copy['authorised_party2'] = add_new_authorised_party(
                    request_post_copy, 2, request.user)

            if request_post_copy['other_side'] == '-1':
                request_post_copy['other_side'] = add_new_otherside_details(
                    request_post_copy, request.user)
            request_post_copy['created_by'] = file.created_by

            form = OpenFileForm(request_post_copy, instance=file)
            duplicate_obj = copy.deepcopy(file)
            if form.is_valid():
                changed_fields = form.changed_data
                changes = {}
                for field in changed_fields:
                    if field != 'created_by':
                        changes[field] = {
                            'old_value': str(getattr(duplicate_obj, field)),
                            'new_value': None
                        }
                form.save()

                for field in changed_fields:
                    if field != 'created_by':
                        changes[field]['new_value'] = str(getattr(file, field))
                if changes:
                    create_modification(
                        user=request.user,
                        modified_obj=file,
                        changes=changes
                    )

                messages.success(request, 'File successfully updated.')
                return redirect('home', file_number=file_number)
            else:

                for field, errors in form.errors.items():
                    for error in errors:
                        messages.error(
                            request, f"{form[field].label}: {error}")
                        print(f"{form[field].label}: {error}")
                return render(request, 'edit_file.html', {'form_data': form_data, 'form': form, 'file_number': file_number})

        except Exception as e:
            messages.error(request, f"Error during file editing: {str(e)}")
            print(f"Error during file editing: {e}")
            return render(request, 'edit_file.html', {'form_data': form_data, 'file_number': file_number})

    else:

        form = OpenFileForm(instance=file)

    return render(request, 'edit_file.html', {'form': form, 'form_data': form_data,
                                              'file_number': file_number})


@login_required
def add_new_work_file(request, file_number):
    if request.method == 'POST':
        request_post_copy = request.POST.copy()
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy['file_number'] = file_number_id
        request_post_copy['created_by'] = request.user
        form = NextWorkForm(request_post_copy)
        if form.is_valid():
            form.save()
            messages.success(request, 'Next work successfully added.')
            return redirect('home', file_number=file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')
    return redirect('home', file_number=file_number)


@login_required
def edit_next_work(request, id):
    nextwork_instance = get_object_or_404(NextWork, pk=id)

    if nextwork_instance.person != request.user and nextwork_instance.created_by != request.user:
        messages.error(request, 'You can only edit tasks assigned to you.')
        return redirect('user_dashboard')

    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(nextwork_instance)
        form = NextWorkForm(request.POST, instance=nextwork_instance)
        if form.is_valid():

            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:

                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()

            for field in changed_fields:

                if field == 'completed':
                    if getattr(nextwork_instance, field):
                        link = reverse('attendance_note_view', args=[
                                       nextwork_instance.file_number.file_number])
                        add_attendance_note_link = f"<a href='{link}' class='link'>add an attendance note</a>"
                        messages.info(request, mark_safe(
                            f'Please remember to {add_attendance_note_link} for work just completed.'))
                changes[field]['new_value'] = str(
                    getattr(nextwork_instance, field))

            create_modification(
                user=request.user,
                modified_obj=nextwork_instance,
                changes=changes
            )

            messages.success(request, 'Successfully updated next work.')
            return redirect('home', nextwork_instance.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = NextWorkForm(instance=nextwork_instance)

    # Render the template with the form
    return render(request, 'edit_models.html', {'form': form, 'title': 'Next Work', 'file_number': nextwork_instance.file_number.file_number})


@login_required
def add_last_work_file(request, file_number):
    if request.method == 'POST':
        request_post_copy = request.POST.copy()
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy['file_number'] = file_number_id
        request_post_copy['created_by'] = request.user
        form = LastWorkForm(request_post_copy)
        if form.is_valid():
            form.save()
            link = reverse('attendance_note_view', args=[file_number])
            add_attendance_note_link = f"<a href='{link}' class='link '>add an attendance note</a>"
            messages.info(request, mark_safe(
                f'Please remember to {add_attendance_note_link} for work just added.'))
            messages.success(request, 'Last work successfully added.')
            return redirect('home', file_number=file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')
    return redirect('home', file_number=file_number)


@login_required
def edit_last_work(request, id):
    lastwork_instance = get_object_or_404(LastWork, pk=id)

    if lastwork_instance.person != request.user and lastwork_instance.created_by != request.user:
        messages.error(request, 'You can only edit tasks assigned to you.')
        return redirect('user_dashboard')

    duplicate_obj = copy.deepcopy(lastwork_instance)
    if request.method == 'POST':
        form = LastWorkForm(request.POST, instance=lastwork_instance)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()

            for field in changed_fields:
                changes[field]['new_value'] = str(
                    getattr(lastwork_instance, field))

            create_modification(
                user=request.user,
                modified_obj=lastwork_instance,
                changes=changes
            )
            messages.success(request, 'Successfully updated last work.')
            return redirect('home', lastwork_instance.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = LastWorkForm(instance=lastwork_instance)

    # Render the template with the form
    return render(request, 'edit_models.html', {'form': form, 'title': 'Last Work', 'file_number': lastwork_instance.file_number.file_number})


@login_required
def attendance_note_view(request, file_number):
    form = AttendanceNoteFormHalf()
    matter = WIP.objects.select_related(
        'fee_earner', 'matter_type', 'file_status'
    ).filter(file_number=file_number).first()
    if not matter:
        messages.error(request, 'Matter file not found')
        return redirect('user_dashboard')
    attendance_notes = MatterAttendanceNotes.objects.filter(
        file_number=matter.id).order_by('-date')
    return render(request, 'attendance_notes.html', {
        'form': form,
        'file_number': file_number,
        'matter': matter,
        'attendance_notes': attendance_notes,
    })


@login_required
def download_attendance_notes_bulk_template(request, file_number):
    if not WIP.objects.filter(file_number=file_number).exists():
        messages.error(request, f'File number "{file_number}" not found.')
        return redirect('user_dashboard')

    response = HttpResponse(content_type='text/csv')
    response[
        'Content-Disposition'] = f'attachment; filename="attendance_notes_template_{file_number}.csv"'
    writer = csv.writer(response)
    writer.writerow([
        'date',
        'start_time',
        'finish_time',
        'subject_line',
        'content',
        'person',
        'is_charged'
    ])
    return response


def _get_row_value(row, keys):
    for key in keys:
        if key in row and row[key] is not None:
            value = str(row[key]).strip()
            if value:
                return value
    return ""


def _parse_bulk_note_date(value):
    value = (value or "").strip()
    formats = ['%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y']
    for fmt in formats:
        try:
            return datetime.strptime(value, fmt).date()
        except ValueError:
            continue
    raise ValueError("invalid date format")


def _parse_bulk_note_time(value):
    value = (value or "").strip()
    formats = ['%H:%M', '%H:%M:%S', '%I:%M %p', '%I:%M%p']
    for fmt in formats:
        try:
            return datetime.strptime(value, fmt).time()
        except ValueError:
            continue
    raise ValueError("invalid time format")


def _parse_bulk_note_bool(value, default=True):
    if value is None:
        return default

    normalized = str(value).strip().lower()
    if normalized == "":
        return default
    if normalized in {"1", "true", "t", "yes", "y"}:
        return True
    if normalized in {"0", "false", "f", "no", "n"}:
        return False
    raise ValueError("invalid boolean value")


def _to_quill_json(content):
    content = (content or "").replace('\r\n', '\n').replace('\r', '\n').strip()
    if not content:
        return json.dumps({"delta": "", "html": ""})

    if content.startswith('{') and '"delta"' in content and '"html"' in content:
        try:
            json.loads(content)
            return content
        except json.JSONDecodeError:
            pass

    escaped_lines = [html.escape(line) for line in content.split('\n')]
    html_content = "<p>" + "<br>".join(escaped_lines) + "</p>"
    delta_json = json.dumps({"ops": [{"insert": f"{content}\n"}]})
    return json.dumps({"delta": delta_json, "html": html_content})


def _resolve_user_for_bulk_note(raw_value, users_by_key, users_by_full_name):
    candidate = (raw_value or "").strip()
    if not candidate:
        return None, "missing initials/user"

    normalized = candidate.lower()
    if normalized in users_by_key:
        return users_by_key[normalized], None

    matches = users_by_full_name.get(normalized, [])
    if len(matches) == 1:
        return matches[0], None
    if len(matches) > 1:
        return None, f"ambiguous user '{candidate}'"
    return None, f"user '{candidate}' not found"


@login_required
def bulk_upload_attendance_notes(request, file_number):
    if request.method != 'POST':
        messages.error(request, 'Invalid request method.')
        return redirect('attendance_note_view', file_number=file_number)

    wip = WIP.objects.filter(file_number=file_number).first()
    if not wip:
        messages.error(request, f'File number "{file_number}" not found.')
        return redirect('attendance_note_view', file_number=file_number)

    uploaded_file = request.FILES.get('bulk_attendance_file')
    if not uploaded_file:
        messages.error(request, 'Please upload a CSV file.')
        return redirect('attendance_note_view', file_number=file_number)

    try:
        decoded = uploaded_file.read().decode('utf-8-sig')
    except UnicodeDecodeError:
        messages.error(
            request, 'Could not read file. Please upload UTF-8 CSV.')
        return redirect('attendance_note_view', file_number=file_number)

    reader = csv.DictReader(io.StringIO(decoded))
    if not reader.fieldnames:
        messages.error(request, 'CSV is missing a header row.')
        return redirect('attendance_note_view', file_number=file_number)

    users = CustomUser.objects.all()
    users_by_key = {}
    users_by_full_name = {}
    for user in users:
        if user.username:
            users_by_key[user.username.strip().lower()] = user
        if user.email:
            users_by_key[user.email.strip().lower()] = user
        users_by_key[str(user.id)] = user

        full_name = f"{(user.first_name or '').strip()} {(user.last_name or '').strip()}".strip(
        ).lower()
        if full_name:
            users_by_full_name.setdefault(full_name, []).append(user)

    created_count = 0
    errors = []

    for line_number, row in enumerate(reader, start=2):
        if not any((value or "").strip() for value in row.values()):
            continue

        date_str = _get_row_value(row, ['date'])
        start_time_str = _get_row_value(
            row, ['start_time', 'start', 'starttime'])
        finish_time_str = _get_row_value(
            row, ['finish_time', 'finish', 'end_time', 'end'])
        subject_line = _get_row_value(row, ['subject_line', 'subject'])
        content_value = _get_row_value(row, ['content', 'note', 'notes'])
        user_value = _get_row_value(
            row, ['person_attended', 'person', 'initials', 'username', 'user', 'fee_earner'])
        charged_value = _get_row_value(
            row, ['is_charged', 'charged', 'billable'])

        missing_fields = []
        if not date_str:
            missing_fields.append('date')
        if not start_time_str:
            missing_fields.append('start_time')
        if not finish_time_str:
            missing_fields.append('finish_time')
        if not subject_line:
            missing_fields.append('subject_line')
        if not content_value:
            missing_fields.append('content')
        if not user_value:
            missing_fields.append('person/initials')

        if missing_fields:
            errors.append(
                f'Line {line_number}: missing {", ".join(missing_fields)}.')
            continue

        try:
            parsed_date = _parse_bulk_note_date(date_str)
            parsed_start_time = _parse_bulk_note_time(start_time_str)
            parsed_finish_time = _parse_bulk_note_time(finish_time_str)
            parsed_is_charged = _parse_bulk_note_bool(
                charged_value, default=True)
        except ValueError as exc:
            errors.append(f'Line {line_number}: {str(exc)}.')
            continue

        resolved_user, user_error = _resolve_user_for_bulk_note(
            user_value, users_by_key, users_by_full_name)
        if user_error:
            errors.append(f'Line {line_number}: {user_error}.')
            continue

        row_payload = {
            'file_number': wip.id,
            'date': parsed_date.isoformat(),
            'start_time': parsed_start_time.strftime('%H:%M'),
            'finish_time': parsed_finish_time.strftime('%H:%M'),
            'subject_line': subject_line,
            'content': _to_quill_json(content_value),
            'is_charged': str(parsed_is_charged),
            'person_attended': resolved_user.id
        }

        form = AttendanceNoteForm(row_payload)
        if not form.is_valid():
            field_errors = []
            for field, field_messages in form.errors.items():
                for field_message in field_messages:
                    field_errors.append(f'{field}: {field_message}')
            errors.append(
                f'Line {line_number}: {"; ".join(field_errors) or "invalid row"}.')
            continue

        instance = form.save(commit=False)
        instance.created_by = request.user
        instance.save()
        created_count += 1

    if created_count:
        messages.success(
            request, f'Created {created_count} attendance note(s) from bulk upload.')

    if errors:
        preview_errors = " ".join(errors[:5])
        suffix = f' (and {len(errors) - 5} more)' if len(errors) > 5 else ''
        messages.warning(
            request, f'{len(errors)} row(s) failed. {preview_errors}{suffix}')

    if created_count == 0 and not errors:
        messages.warning(request, 'No rows found in CSV.')

    return redirect('attendance_note_view', file_number=file_number)


@login_required
def add_attendance_note(request, file_number):
    if request.method == 'POST':
        request_post_copy = request.POST.copy()
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy['file_number'] = file_number_id

        request_post_copy['created_by'] = request.user
        form = AttendanceNoteForm(request_post_copy)
        if form.is_valid():
            form.save()
            messages.success(request, 'Attendance Note successfully added.')
            return redirect('attendance_note_view', file_number=file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')

    return redirect(attendance_note_view, file_number=file_number)


@login_required
def download_attendance_note(request, id):
    a_n = get_object_or_404(MatterAttendanceNotes, id=id)

    context = {
        'client1_name': a_n.file_number.client1.name,
        'client2_name': a_n.file_number.client2.name if a_n.file_number.client2 else '',
        'file_number': a_n.file_number.file_number,
        'date': a_n.date.strftime('%d/%m/%Y'),
        'start_time': a_n.start_time,
        'finish_time': a_n.finish_time,
        'content': mark_safe(a_n.content.html),
        'is_charged': a_n.is_charged,
        'unit': a_n.unit,
        'person_attended': a_n.person_attended,
        'user': request.user
    }

    # Render the template with the provided context
    html_string = render_to_string(
        'download_templates/attendance_note.html', context)

    # Generate PDF from the rendered HTML using WeasyPrint
    pdf_file = HTML(string=html_string).write_pdf()

    return HttpResponse(pdf_file, content_type='application/pdf')


@login_required
def download_attendance_notes_bulk(request, file_number):
    wip = WIP.objects.filter(file_number=file_number).first()
    if not wip:
        messages.error(request, f'File number "{file_number}" not found.')
        return redirect('user_dashboard')

    attendance_notes = MatterAttendanceNotes.objects.filter(
        file_number=wip.id
    ).order_by('-date', '-start_time', '-id')

    if not attendance_notes.exists():
        messages.warning(request, 'No attendance notes found to download.')
        return redirect('attendance_note_view', file_number=file_number)

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for note in attendance_notes:
            context = {
                'client1_name': note.file_number.client1.name,
                'client2_name': note.file_number.client2.name if note.file_number.client2 else '',
                'file_number': note.file_number.file_number,
                'date': note.date.strftime('%d/%m/%Y'),
                'start_time': note.start_time,
                'finish_time': note.finish_time,
                'content': mark_safe(note.content.html),
                'is_charged': note.is_charged,
                'unit': note.unit,
                'person_attended': note.person_attended,
                'user': request.user
            }

            html_string = render_to_string(
                'download_templates/attendance_note.html', context)
            pdf_file = HTML(string=html_string).write_pdf()
            file_name = f'attendance_note_{note.date.strftime("%Y%m%d")}_{note.id}.pdf'
            zip_file.writestr(file_name, pdf_file)

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer.getvalue(),
                            content_type='application/zip')
    response['Content-Disposition'] = (
        f'attachment; filename="attendance_notes_{file_number}.zip"'
    )
    return response


@login_required
def edit_attendance_note(request, id):
    attendance_note_instance = get_object_or_404(MatterAttendanceNotes, pk=id)
    if request.method == 'POST':

        duplicate_obj = copy.deepcopy(attendance_note_instance)
        form = AttendanceNoteForm(
            request.POST, instance=attendance_note_instance)

        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                if field == 'content':
                    changes[field] = {
                        'old_value': duplicate_obj.content.html,
                        'new_value': None
                    }
                else:
                    changes[field] = {
                        'old_value': str(getattr(duplicate_obj, field)),
                        'new_value': None
                    }
            form.save()

            for field in changed_fields:
                if field == 'content':
                    changes[field]['new_value'] = attendance_note_instance.content.html
                else:
                    changes[field]['new_value'] = str(
                        getattr(attendance_note_instance, field))

            create_modification(
                user=request.user,
                modified_obj=attendance_note_instance,
                changes=changes
            )
            messages.success(request, 'Successfully updated Attendance Note.')
            return redirect('attendance_note_view', attendance_note_instance.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = AttendanceNoteForm(instance=attendance_note_instance)

    return render(request, 'edit_models.html', {'form': form, 'title': 'Attendance Note', 'file_number': attendance_note_instance.file_number.file_number})


@login_required
def correspondence_view(request, file_number):
    letter_form = LetterHalfForm()
    matter = WIP.objects.select_related(
        'fee_earner', 'matter_type', 'file_status'
    ).filter(file_number=file_number).first()
    if not matter:
        messages.error(request, 'Matter file not found')
        return redirect('user_dashboard')
    emails = MatterEmails.objects.filter(
        file_number=matter.id).order_by('-time')
    letters = MatterLetters.objects.filter(
        file_number=matter.id).order_by('-date')
    return render(request, 'correspondence.html', {
        'letter_form': letter_form,
        'file_number': file_number,
        'matter': matter,
        'emails': emails,
        'letters': letters,
    })


@login_required
def add_letter(request, file_number):
    if request.method == 'POST':
        request_post_copy = request.POST.copy()
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy['file_number'] = file_number_id
        request_post_copy['created_by'] = request.user
        form = LetterForm(request_post_copy)
        if form.is_valid():
            form.save()
            messages.success(request, 'Letter successfully added.')
            return redirect('correspondence_view', file_number=file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')

    return redirect('correspondence_view', file_number=file_number)


@login_required
def edit_letter(request, id):

    letter_instance = get_object_or_404(MatterLetters, pk=id)

    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(letter_instance)
        form = LetterForm(request.POST, instance=letter_instance)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()

            for field in changed_fields:
                changes[field]['new_value'] = getattr(letter_instance, field)

            create_modification(
                user=request.user,
                modified_obj=letter_instance,
                changes=changes
            )
            messages.success(request, 'Successfully updated Letter.')
            return redirect('correspondence_view', letter_instance.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = LetterForm(instance=letter_instance)

    return render(request, 'edit_models.html', {'form': form, 'title': 'Letter', 'file_number': letter_instance.file_number.file_number})


@login_required
def download_sowc(request, file_number):
    file = WIP.objects.filter(file_number=file_number).first()
    file_number_id = file.id
    emails = MatterEmails.objects.filter(file_number=file_number_id)
    attendance_notes = MatterAttendanceNotes.objects.filter(
        file_number=file_number_id)
    letters = MatterLetters.objects.filter(file_number=file_number_id)

    """
    {date, time, fee_earner, Description, Unit(s), Amount}
    """
    rows = []

    for note in attendance_notes:
        date = note.date.strftime('%d/%m/%Y')
        time = note.start_time.strftime('%H:%M')
        fee_earner = note.person_attended.username if note.person_attended != None else ''
        note_charge_status = " (N/C)" if not note.is_charged else ""
        desc = f"Attendance Note{note_charge_status} - {note.subject_line} from {note.start_time.strftime(
            '%I:%M %p')} to {note.finish_time.strftime('%I:%M %p')}"
        units = note.unit
        amount = ((note.person_attended.hourly_rate.hourly_amount/10) * units) if note.person_attended != None else (
            (note.file_number.fee_earner.hourly_rate.hourly_amount/10) * units)
        row = [date, time, fee_earner, desc, units, amount]
        rows.append(row)

    for email in emails:

        date = email.time.date().strftime('%d/%m/%Y')
        time = email.time.astimezone(
            timezone.get_current_timezone()).time().strftime('%H:%M')
        fee_earner = email.fee_earner.username if email.fee_earner != None else ''
        receiver = json.loads(email.receiver)
        sender = json.loads(email.sender)
        to_or_from = f"Email to {receiver[0]['emailAddress']['name']}" if email.is_sent else f"Perusal of email from {sender['emailAddress']['name']}"
        desc = to_or_from + f" @ {time}"
        units = email.units
        amount = ((email.fee_earner.hourly_rate.hourly_amount/10) * units) if email.fee_earner != None else (
            (email.file_number.fee_earner.hourly_rate.hourly_amount/10) * units)
        row = [date, time, fee_earner, desc, units, amount]
        rows.append(row)

    for letter in letters:
        date = letter.date.strftime('%d/%m/%Y')
        time = None
        fee_earner = letter.person_attended.username if letter.person_attended != None else ''
        to_or_from = f'Letter to {letter.to_or_from}' if letter.sent else f'Letter from {letter.to_or_from}'
        desc = f'{to_or_from} - {letter.subject_line}'
        units = 1
        amount = ((letter.person_attended.hourly_rate.hourly_amount/10) * units) if letter.person_attended != None else (
            (letter.file_number.fee_earner.hourly_rate.hourly_amount/10) * units)
        row = [date, time, fee_earner, desc, units, amount]
        rows.append(row)

    def sort_rows(rows):
        def get_sort_key(row):
            # Handling empty dates
            date_str = row[0] if row[0] else '01/01/0001'
            time_str = row[1] if row[1] else '00:00'
            date_time_str = f"{date_str} {time_str}"
            date_time = datetime.strptime(date_time_str, '%d/%m/%Y %H:%M')
            return date_time

        sorted_rows = sorted(rows, key=get_sort_key)
        return sorted_rows
    sorted_rows = sort_rows(rows)

    distinct_fee_earners = set(row[2] for row in rows)
    first_date = sorted_rows[0][0] if sorted_rows else None
    last_date = sorted_rows[-1][0] if sorted_rows else None

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="({file_number}) Schedule of Work and Costs from {first_date} to {last_date}.csv"'

    writer = csv.writer(response)

    writer.writerow(['', '', f'Client Name: {file.client1.name} Matter:{
                    file.matter_description}[{file.file_number}]'])
    writer.writerow(['', '', f'Schedule of Work and Costs from {
                    first_date} to {last_date}'])

    for fee_earner in distinct_fee_earners:
        user = CustomUser.objects.filter(username=fee_earner).first()
        if user != None:
            writer.writerow(
                ['', '', f'({user.first_name} {user.last_name}) {user.username} rate GBP{user.hourly_rate.hourly_amount} + VAT per hour, 6 minutes = 1 unit '])
    writer.writerow([])
    writer.writerow(['Date', 'Fee Earner', 'Description', 'Unit(s)', 'Amount'])
    for row in sorted_rows:
        writer.writerow([row[0], row[2], row[3], row[4], row[5]])
    writer.writerow([])

    sum_start_row = 4 + len(distinct_fee_earners) + 1
    sum_end_row = sum_start_row + len(sorted_rows)
    total_cost_row = sum_end_row + 1
    writer.writerow(['', '', 'Total Costs', '',
                    f'=sum(E{sum_start_row}:E{sum_end_row})'])
    writer.writerow(
        ['', '', f'VAT @{CURRENT_VAT_RATE_PERCENT}%', '', f'={CURRENT_VAT_RATE}*E{total_cost_row}'])
    writer.writerow(['', '', 'Total Costs and VAT', '',
                    f'=sum(E{total_cost_row}:E{total_cost_row+1})'])

    return response


@login_required
def finance_view(request, file_number):
    file_obj = WIP.objects.select_related(
        'fee_earner', 'matter_type', 'file_status'
    ).filter(file_number=file_number).first()
    if not file_obj:
        messages.error(request, 'File not found.')
        return redirect('index')

    file_number_id = file_obj.id
    pmts_slips = PmtsSlips.objects.filter(
        file_number=file_number_id).order_by('-date')
    pmts_form = PmtsHalfForm()
    green_slips_form = LedgerAccountTransfersHalfForm()
    credit_note_form = CreditNoteHalfForm()
    credit_note_form.fields['invoice'].queryset = Invoices.objects.filter(
        file_number=file_number_id, state='F').order_by('-invoice_number')
    green_slips = LedgerAccountTransfers.objects.filter(
        Q(file_number_from=file_number_id) | Q(file_number_to=file_number_id)).order_by('-date')
    invoices = Invoices.objects.filter(
        file_number=file_number_id).order_by('-date')
    credit_notes = CreditNote.objects.filter(
        file_number=file_number_id
    ).select_related('invoice', 'created_by', 'approved_by').order_by('-date', '-timestamp')

    credit_notes_by_invoice = {}
    approved_credit_totals = {}
    status_labels = dict(CreditNote.STATUSES)
    credit_notes_data = []
    for credit_note in credit_notes:
        credit_notes_by_invoice.setdefault(
            credit_note.invoice_id, []).append(credit_note)
        if credit_note.status == 'F':
            approved_credit_totals[credit_note.invoice_id] = approved_credit_totals.get(
                credit_note.invoice_id, Decimal('0')
            ) + credit_note.amount
        net_amount, vat_amount, gross_amount = calculate_credit_note_breakdown(
            credit_note.amount)

        credit_notes_data.append({
            'id': credit_note.id,
            'date': credit_note.date.strftime('%d/%m/%Y'),
            'invoice_number': credit_note.invoice.invoice_number,
            'amount': gross_amount,
            'net_amount': net_amount,
            'vat_amount': vat_amount,
            'reason': credit_note.reason,
            'status': credit_note.status,
            'status_display': status_labels.get(credit_note.status, credit_note.status),
            'created_by': credit_note.created_by,
            'approved_by': credit_note.approved_by,
            'approved_on': credit_note.approved_on,
            'can_review': request.user.is_manager and credit_note.status == 'P',
            'can_edit': (
                (credit_note.status == 'P' and (
                    credit_note.created_by_id == request.user.id or request.user.is_manager
                )) or
                (credit_note.status == 'F' and request.user.is_manager)
            ),
        })

    invoices_data = []

    total_invoices = Decimal('0')
    total_out = Decimal('0')
    total_approved_credit_notes = Decimal('0')
    for invoice in invoices:

        our_costs = invoice.our_costs

        costs = ast.literal_eval(our_costs) if type(
            our_costs) != type([]) else our_costs
        total_cost_invoice = Decimal('0')

        our_costs_desc_pre = invoice.our_costs_desc
        our_costs_desc = ast.literal_eval(our_costs_desc_pre) if type(
            our_costs_desc_pre) != type([]) else our_costs_desc_pre
        costs_display = "<div>"
        for i in range(len(costs)):
            total_cost_invoice = total_cost_invoice + Decimal(costs[i])
            costs_display = costs_display + \
                f"<b>{our_costs_desc[i]}</b>: £{costs[i]}<br>"
        _, vat_inv, total_cost_and_vat = calculate_invoice_total_with_vat(
            invoice)
        costs_display = costs_display + \
            f"Add VAT @{CURRENT_VAT_RATE_PERCENT}%: £{round(vat_inv, 2)}<br>"
        total_cost_and_vat = round(total_cost_and_vat, 2)
        costs_display = costs_display + \
            f"<b>Total Costs and VAT:</b> £{total_cost_and_vat}<br>"
        costs_display = costs_display + "</div>"

        blue_slips_display = "<div class='mt-2'><h5 class='text-xl font-medium' >Blues Slips attached</h5>"
        total_blue_slips = 0
        if invoice.moa_ids.exists():
            for slip in invoice.moa_ids.all():
                if isinstance(slip.amount_invoiced, str):
                    amount_invoiced = json.loads(slip.amount_invoiced)
                elif isinstance(slip.amount_invoiced, (bytes, bytearray)):
                    amount_invoiced = json.loads(
                        slip.amount_invoiced.decode('utf-8'))
                elif isinstance(slip.amount_invoiced, dict):
                    amount_invoiced = slip.amount_invoiced
                else:
                    raise ValueError(
                        "Unsupported type for slip.amount_invoiced")

                date = slip.date.strftime('%d/%m/%Y')
                amt = amount_invoiced[f"{invoice.id}"]['amt_invoiced']
                total_blue_slips = total_blue_slips + Decimal(amt)
                blue_slips_display = blue_slips_display + \
                    f"Payment from {slip.pmt_person} of <b>£{amt}</b> on <b>{date}</b><br>"
            blue_slips_display = blue_slips_display + \
                f"<b>Total Blue Slips:</b> £{round(total_blue_slips, 2)}<br>"
        else:
            blue_slips_display = blue_slips_display + "No Blue Slips Attached"

        blue_slips_display = blue_slips_display + "</div>"

        pink_slips_display = "<div><h5 class='text-xl font-medium' >Pink Slips attached</h5>"
        total_pink_slips = 0
        if invoice.disbs_ids.exists():
            for slip in invoice.disbs_ids.all():
                date = slip.date.strftime('%d/%m/%Y')
                total_pink_slips = total_pink_slips + slip.amount
                pink_slips_display = pink_slips_display + \
                    f"Payment to {slip.pmt_person} of £{
                        slip.amount} on {date}<br>"
            pink_slips_display = pink_slips_display + \
                f"<b>Total Pink Slips:</b> £{total_pink_slips}<br>"
        else:
            pink_slips_display = pink_slips_display + "No Pink Slips Attached"

        pink_slips_display = pink_slips_display + "</div>"

        green_slips_display = "<div><h5 class='text-xl font-medium' >Green Slips attached</h5>"
        total_green_slips = 0
        if invoice.green_slip_ids.exists():
            for slip in invoice.green_slip_ids.all():
                date = slip.date.strftime('%d/%m/%Y')
                if slip.file_number_from.file_number == file_number:
                    total_green_slips = total_green_slips - slip.amount
                    green_slips_display = green_slips_display + \
                        f"Transfer to {slip.file_number_to} of £{
                            slip.amount} on {date}<br>"
                else:

                    if isinstance(slip.amount_invoiced_to, str):
                        amount_invoiced = json.loads(slip.amount_invoiced_to)
                    elif isinstance(slip.amount_invoiced_to, (bytes, bytearray)):
                        amount_invoiced = json.loads(
                            slip.amount_invoiced_to.decode('utf-8'))
                    elif isinstance(slip.amount_invoiced_to, dict):
                        amount_invoiced = slip.amount_invoiced_to
                    else:
                        raise ValueError(
                            "Unsupported type for slip.amount_invoiced_to")

                    date = slip.date.strftime('%d/%m/%Y')
                    amt = amount_invoiced[f"{invoice.id}"]['amt_invoiced']

                    total_green_slips = total_green_slips + Decimal(amt)
                    green_slips_display = green_slips_display + \
                        f"Transfer from {slip.file_number_from} of £{
                            amt} on {date}<br>"
            green_slips_display = green_slips_display + \
                f"<b>Total Green Slips:</b> £{total_green_slips}<br>"
        else:
            green_slips_display = green_slips_display + "No Green Slips Attached"
        green_slips_display = green_slips_display + "</div>"

        cash_allocated_slips_display = "<div class='mt-2'><h5 class='text-xl font-medium' >Blue Slips attached (after invoice creation)</h5>"
        total_cash_allocated_slips = 0
        if invoice.cash_allocated_slips.exists():
            for slip in invoice.cash_allocated_slips.all():
                if isinstance(slip.amount_allocated, str):
                    amount_invoiced = json.loads(slip.amount_allocated)
                elif isinstance(slip.amount_allocated, (bytes, bytearray)):
                    amount_invoiced = json.loads(
                        slip.amount_allocated.decode('utf-8'))
                elif isinstance(slip.amount_allocated, dict):
                    amount_invoiced = slip.amount_allocated
                else:
                    raise ValueError(
                        "Unsupported type for slip.amount_invoiced")

                date = slip.date.strftime('%d/%m/%Y')
                invoice_id_str = f'{invoice.id}'
                if invoice_id_str in amount_invoiced:
                    amt = amount_invoiced[invoice_id_str]
                    total_cash_allocated_slips = total_cash_allocated_slips + \
                        Decimal(amt)
                    cash_allocated_slips_display = cash_allocated_slips_display + \
                        f"Payment from {slip.pmt_person} of <b>£{amt}</b> on <b>{date}</b><br>"

            cash_allocated_slips_display = cash_allocated_slips_display + \
                f"<b>Total Allocated Slips:</b> £{total_cash_allocated_slips}<br>"
        else:
            cash_allocated_slips_display = cash_allocated_slips_display + \
                "No Slips Attached After Invoice Creation"

        cash_allocated_slips_display = cash_allocated_slips_display + "</div>"

        invoice_credit_notes = credit_notes_by_invoice.get(invoice.id, [])
        approved_credit_total = approved_credit_totals.get(
            invoice.id, Decimal('0'))
        total_approved_credit_notes += approved_credit_total
        credit_notes_display = "<div><h5 class='text-xl font-medium'>Credit Notes</h5>"
        if invoice_credit_notes:
            for note in invoice_credit_notes:
                status_display = status_labels.get(note.status, note.status)
                net_amount, vat_amount, gross_amount = calculate_credit_note_breakdown(
                    note.amount)
                approved_meta = ""
                if note.approved_by:
                    approved_meta = (
                        f" (approved by {note.approved_by} on {note.approved_on.strftime('%d/%m/%Y %H:%M')})"
                        if note.approved_on else f" (approved by {note.approved_by})"
                    )
                credit_notes_display = credit_notes_display + (
                    f"{note.date.strftime('%d/%m/%Y')} - Ex VAT £{net_amount}, VAT £{vat_amount}, "
                    f"Total £{gross_amount} - {status_display}{approved_meta}<br>"
                )
            credit_notes_display = credit_notes_display + \
                f"<b>Total Final Credit Notes:</b> £{round(approved_credit_total, 2)}<br>"
        else:
            credit_notes_display = credit_notes_display + "No Credit Notes Issued"
        credit_notes_display = credit_notes_display + "</div>"

        balance = (total_cost_and_vat + total_pink_slips) - \
            total_green_slips - \
            (total_blue_slips + total_cash_allocated_slips) - approved_credit_total

        if balance >= 0:
            total_due_display = f"<div><b>Total Due: </b> £{round(balance, 2)}<br></div>"
        else:
            balance = balance * -1
            total_due_display = f"<div><b>Balance remaining on account:</b> £{round(balance, 2)}<br></div>"

        effective_due_left = get_effective_invoice_due(
            invoice, approved_credit_total)

        data = {'id': invoice.id,
                'state': invoice.state,
                'number': invoice.invoice_number,
                'total_cost_and_vat': total_cost_and_vat,
                'desc': invoice.description,
                'date': invoice.date.strftime('%d/%m/%Y'),
                'costs': mark_safe(costs_display),
                'pink_slips': mark_safe(pink_slips_display),
                'blue_slips': mark_safe(blue_slips_display),
                'green_slips': mark_safe(green_slips_display),
                'cash_allocated_slips': mark_safe(cash_allocated_slips_display),
                'credit_notes': mark_safe(credit_notes_display),
                'total_due': mark_safe(total_due_display),
                'total_due_left': effective_due_left,
                }

        invoices_data.append(data)

        total_invoices = total_invoices + total_cost_and_vat

    total_out = total_out + total_invoices
    total_in = Decimal('0')
    for slip in pmts_slips:
        if slip.is_money_out == True:
            total_out = total_out + slip.amount
        else:
            total_in = total_in + slip.amount

    for slip in green_slips:
        if slip.file_number_from.file_number == slip.file_number_to.file_number:
            continue
        if slip.file_number_from.file_number == file_number:
            total_out = total_out + slip.amount

        else:
            total_in = total_in + slip.amount
    total_in = total_in + total_approved_credit_notes
    total_in = round(total_in, 2)
    total_out = round(total_out, 2)
    total_balance = round(total_in - total_out, 2)

    colors = {'draft_invoice': "#F9EBDF",
              "invoice": "#FFFCC9",
              'credit_note': "#FFEAEA",
              'green': "#90EE90",
              'temp': "#CCD1D1"}

    return render(request, 'finances.html', {'total_monies_in': total_in, 'total_monies_out': total_out, 'total_monies_balance': total_balance,
                                             'pmts_slips': pmts_slips, 'file_number': file_number, 'file_number_id': file_number_id,
                                             'matter': file_obj,
                                             'colors': colors,
                                             'pmts_form': pmts_form, 'green_slip_form': green_slips_form,
                                             'green_slips': green_slips, 'invoices': invoices_data,
                                             'credit_notes': credit_notes_data, 'credit_note_form': credit_note_form,
                                             'current_vat_rate_percent': CURRENT_VAT_RATE_PERCENT})


@login_required
def add_credit_note(request, file_number):
    if request.method == 'POST':
        file_obj = WIP.objects.filter(file_number=file_number).first()
        if not file_obj:
            messages.error(request, 'File not found.')
            return redirect('index')

        request_post_copy = request.POST.copy()
        request_post_copy['file_number'] = file_obj.id
        form = CreditNoteHalfForm(request_post_copy)
        form.fields['invoice'].queryset = Invoices.objects.filter(
            file_number=file_obj.id, state='F'
        ).order_by('-invoice_number')
        if form.is_valid():
            credit_note = form.save(commit=False)
            if credit_note.invoice.file_number_id != file_obj.id:
                messages.error(
                    request, 'Selected invoice does not belong to this matter.')
                return redirect('finance_view', file_number=file_number)
            if credit_note.amount <= 0:
                messages.error(
                    request, 'Credit note amount must be greater than 0.')
                return redirect('finance_view', file_number=file_number)

            max_allowed_amount = get_effective_invoice_due(credit_note.invoice)
            if credit_note.amount > max_allowed_amount:
                messages.error(
                    request,
                    f'Credit note exceeds invoice due (£{max_allowed_amount}).'
                )
                return redirect('finance_view', file_number=file_number)

            credit_note.file_number = file_obj
            credit_note.created_by = request.user
            credit_note.status = 'P'
            credit_note.save()
            messages.success(
                request, 'Credit note submitted and is pending manager approval.')
            return redirect('finance_view', file_number=file_number)

        error_message = 'Form is not valid. Please correct the errors:'
        for field, errors in form.errors.items():
            error_message += f'\n{field}: {", ".join(errors)}'
        messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')

    return redirect('finance_view', file_number=file_number)


@login_required
def approve_credit_note(request, id):
    credit_note = get_object_or_404(CreditNote, id=id)
    if not request.user.is_manager:
        messages.error(request, 'Only managers can approve credit notes.')
        return redirect('finance_view', file_number=credit_note.file_number.file_number)

    if credit_note.status != 'P':
        messages.info(request, 'This credit note has already been reviewed.')
        return redirect('finance_view', file_number=credit_note.file_number.file_number)

    max_allowed_amount = get_effective_invoice_due(credit_note.invoice)
    if credit_note.amount > max_allowed_amount:
        messages.error(
            request,
            f'Credit note cannot be approved because invoice due is £{max_allowed_amount}.'
        )
        return redirect('finance_view', file_number=credit_note.file_number.file_number)

    invoice = credit_note.invoice
    prev_invoice_due = Decimal(str(invoice.total_due_left or 0))
    new_invoice_due = prev_invoice_due - Decimal(str(credit_note.amount))
    if new_invoice_due < 0:
        new_invoice_due = Decimal('0')

    with transaction.atomic():
        invoice.total_due_left = new_invoice_due
        invoice.save(update_fields=['total_due_left'])

        old_status = credit_note.status
        credit_note.status = 'F'
        credit_note.approved_by = request.user
        credit_note.approved_on = timezone.now()
        credit_note.save(
            update_fields=['status', 'approved_by', 'approved_on'])

        create_modification(
            request.user,
            invoice,
            {
                'total_due_left': {
                    'old_value': str(prev_invoice_due),
                    'new_value': str(invoice.total_due_left)
                },
                'reason': f'Approved Credit Note {credit_note.id}'
            }
        )

        create_modification(
            request.user,
            credit_note,
            {
                'status': {'old_value': old_status, 'new_value': 'F'},
                'approved_by': {'old_value': None, 'new_value': str(request.user)},
                'approved_on': {'old_value': None, 'new_value': str(credit_note.approved_on)},
            }
        )
    messages.success(request, 'Credit note approved and finalized.')
    return redirect('finance_view', file_number=credit_note.file_number.file_number)


@login_required
def reject_credit_note(request, id):
    credit_note = get_object_or_404(CreditNote, id=id)
    if not request.user.is_manager:
        messages.error(request, 'Only managers can reject credit notes.')
        return redirect('finance_view', file_number=credit_note.file_number.file_number)

    if credit_note.status != 'P':
        messages.info(request, 'This credit note has already been reviewed.')
        return redirect('finance_view', file_number=credit_note.file_number.file_number)

    old_status = credit_note.status
    credit_note.status = 'R'
    credit_note.approved_by = request.user
    credit_note.approved_on = timezone.now()
    credit_note.save(update_fields=['status', 'approved_by', 'approved_on'])

    create_modification(
        request.user,
        credit_note,
        {
            'status': {'old_value': old_status, 'new_value': 'R'},
            'approved_by': {'old_value': None, 'new_value': str(request.user)},
            'approved_on': {'old_value': None, 'new_value': str(credit_note.approved_on)},
        }
    )
    messages.success(request, 'Credit note rejected.')
    return redirect('finance_view', file_number=credit_note.file_number.file_number)


@login_required
def edit_credit_note(request, id):
    credit_note = get_object_or_404(CreditNote, id=id)

    can_edit_pending = credit_note.status == 'P' and (
        credit_note.created_by_id == request.user.id or request.user.is_manager
    )
    can_edit_approved = credit_note.status == 'F' and request.user.is_manager
    if not (can_edit_pending or can_edit_approved):
        messages.error(
            request, 'You do not have permission to edit this credit note.')
        return redirect('finance_view', file_number=credit_note.file_number.file_number)

    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(credit_note)
        form = CreditNoteHalfForm(request.POST, instance=credit_note)
        form.fields['invoice'].queryset = Invoices.objects.filter(
            file_number=credit_note.file_number_id, state='F'
        ).order_by('-invoice_number')
        if form.is_valid():
            edited_credit_note = form.save(commit=False)
            edited_credit_note.file_number = credit_note.file_number

            if edited_credit_note.amount <= 0:
                messages.error(
                    request, 'Credit note amount must be greater than 0.')
                return redirect('edit_credit_note', id=credit_note.id)

            if edited_credit_note.status == 'F':
                max_allowed_amount = get_invoice_max_credit_amount(
                    edited_credit_note.invoice, excluded_credit_note_id=edited_credit_note.id)
            else:
                max_allowed_amount = get_effective_invoice_due(
                    edited_credit_note.invoice)

            if edited_credit_note.amount > max_allowed_amount:
                messages.error(
                    request,
                    f'Credit note exceeds invoice due (£{max_allowed_amount}).'
                )
                return redirect('edit_credit_note', id=credit_note.id)

            with transaction.atomic():
                if duplicate_obj.status == 'F':
                    old_invoice = Invoices.objects.select_for_update().filter(
                        id=duplicate_obj.invoice_id
                    ).first()
                    new_invoice = Invoices.objects.select_for_update().filter(
                        id=edited_credit_note.invoice_id
                    ).first()

                    old_invoice_prev_due = Decimal(
                        str(old_invoice.total_due_left or 0))
                    old_invoice.total_due_left = old_invoice_prev_due + \
                        Decimal(str(duplicate_obj.amount))
                    old_invoice.save(update_fields=['total_due_left'])

                    if old_invoice.id == new_invoice.id:
                        new_invoice_prev_due = Decimal(
                            str(old_invoice.total_due_left or 0))
                    else:
                        new_invoice_prev_due = Decimal(
                            str(new_invoice.total_due_left or 0))

                    new_invoice_due = new_invoice_prev_due - \
                        Decimal(str(edited_credit_note.amount))
                    if new_invoice_due < 0:
                        new_invoice_due = Decimal('0')
                    new_invoice.total_due_left = new_invoice_due
                    new_invoice.save(update_fields=['total_due_left'])

                    create_modification(
                        user=request.user,
                        modified_obj=old_invoice,
                        changes={
                            'total_due_left': {
                                'old_value': str(old_invoice_prev_due),
                                'new_value': str(old_invoice.total_due_left)
                            },
                            'reason': f'Edited Credit Note {edited_credit_note.id} (reverse old amount)'
                        }
                    )
                    create_modification(
                        user=request.user,
                        modified_obj=new_invoice,
                        changes={
                            'total_due_left': {
                                'old_value': str(new_invoice_prev_due),
                                'new_value': str(new_invoice.total_due_left)
                            },
                            'reason': f'Edited Credit Note {edited_credit_note.id} (apply new amount)'
                        }
                    )

                edited_credit_note.save()

                changed_fields = form.changed_data
                changes = {}
                for field in changed_fields:
                    changes[field] = {
                        'old_value': str(getattr(duplicate_obj, field)),
                        'new_value': str(getattr(edited_credit_note, field))
                    }
                if changes:
                    create_modification(
                        user=request.user,
                        modified_obj=edited_credit_note,
                        changes=changes
                    )
            messages.success(request, 'Credit note updated.')
            return redirect('finance_view', file_number=edited_credit_note.file_number.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = CreditNoteHalfForm(instance=credit_note)
        form.fields['invoice'].queryset = Invoices.objects.filter(
            file_number=credit_note.file_number_id, state='F'
        ).order_by('-invoice_number')

    return render(request, 'edit_models.html', {
        'form': form,
        'title': 'Edit Credit Note',
        'file_number': credit_note.file_number.file_number
    })


@login_required
def add_pink_slip(request, file_number):
    if request.method == 'POST':
        request_post_copy = request.POST.copy()
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy['file_number'] = file_number_id
        request_post_copy['is_money_out'] = True
        request_post_copy['balance_left'] = request_post_copy['amount']
        request_post_copy['created_by'] = request.user
        form = PmtsForm(request_post_copy)
        if form.is_valid():
            form.save()
            messages.success(request, 'Pink Slip successfully added.')
            return redirect('finance_view', file_number=file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')

    return redirect('finance_view', file_number=file_number)


@login_required
def add_blue_slip(request, file_number):
    if request.method == 'POST':
        request_post_copy = request.POST.copy()
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy['file_number'] = file_number_id
        request_post_copy['is_money_out'] = False
        request_post_copy['balance_left'] = request_post_copy['amount']
        request_post_copy['created_by'] = request.user
        form = PmtsForm(request_post_copy)
        if form.is_valid():
            form.save()
            messages.success(request, 'Blue Slip successfully added.')
            return redirect('finance_view', file_number=file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')

    return redirect('finance_view', file_number=file_number)


@login_required
def add_green_slip(request, file_number):
    if request.method == 'POST':
        request_post_copy = request.POST.copy()
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy['file_number_from'] = file_number_id

        request_post_copy['balance_left_from'] = request_post_copy['amount']
        request_post_copy['balance_left_to'] = request_post_copy['amount']
        request_post_copy['created_by'] = request.user
        form = LedgerAccountTransfersForm(request_post_copy)
        if form.is_valid():
            form.save()
            messages.success(request, 'Green Slip successfully added.')
            return redirect('finance_view', file_number=file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')

    return redirect('finance_view', file_number=file_number)


@login_required
def edit_pmts_slip(request, id):

    pmt_instance = get_object_or_404(PmtsSlips, pk=id)

    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(pmt_instance)
        post_copy = request.POST.copy()

        form = PmtsForm(request.POST, instance=pmt_instance)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()

            for field in changed_fields:
                changes[field]['new_value'] = str(getattr(pmt_instance, field))

            create_modification(
                user=request.user,
                modified_obj=pmt_instance,
                changes=changes
            )
            messages.success(request, 'Successfully updated Slip.')
            return redirect('finance_view', pmt_instance.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = PmtsForm(instance=pmt_instance)

    return render(request, 'edit_models.html', {'form': form, 'title': 'Slip', 'file_number': pmt_instance.file_number.file_number})


@login_required
def download_pmts_slip(request, id):
    slip = get_object_or_404(PmtsSlips, id=id)

    # Render the template with the provided context
    html_string = render_to_string(
        'download_templates/pmts_slip.html', {"slip": slip})

    # Generate PDF from the rendered HTML using WeasyPrint
    pdf_file = HTML(string=html_string).write_pdf()

    return HttpResponse(pdf_file, content_type='application/pdf')


@login_required
def edit_green_slip(request, id):

    pmt_instance = get_object_or_404(LedgerAccountTransfers, pk=id)

    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(pmt_instance)
        form = LedgerAccountTransfersForm(request.POST, instance=pmt_instance)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()

            for field in changed_fields:
                changes[field]['new_value'] = str(getattr(pmt_instance, field))

            create_modification(
                user=request.user,
                modified_obj=pmt_instance,
                changes=changes
            )
            messages.success(request, 'Successfully updated Green Slip.')
            return redirect('finance_view', pmt_instance.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = LedgerAccountTransfersForm(instance=pmt_instance)

    return render(request, 'edit_models.html', {'form': form, 'title': 'Edit Green Slip', 'file_number': pmt_instance.file_number_from.file_number})


@login_required
def download_green_slip(request, id):
    slip = get_object_or_404(LedgerAccountTransfers, id=id)

    # Render the template with the provided context
    html_string = render_to_string(
        'download_templates/green_slip.html', {"slip": slip})

    # Generate PDF from the rendered HTML using WeasyPrint
    pdf_file = HTML(string=html_string).write_pdf()

    return HttpResponse(pdf_file, content_type='application/pdf')


@login_required
def add_invoice(request, file_number):
    if request.method == 'POST':
        file_number_id = WIP.objects.filter(file_number=file_number).first().id
        request_post_copy = request.POST.copy()

        if request_post_copy['state'] == 'F':
            largest_invoice_number = Invoices.objects.aggregate(Max('invoice_number'))[
                'invoice_number__max']
            request_post_copy['invoice_number'] = largest_invoice_number+1

        if 'by_email' in request_post_copy:
            request_post_copy['by_email'] = True
        if 'by_post' in request_post_copy:
            request_post_copy['by_post'] = True

        request_post_copy['our_costs_desc'] = json.dumps(
            request.POST.getlist('our_costs_desc[]'))
        request_post_copy['our_costs'] = json.dumps(
            request.POST.getlist('our_costs[]'))

        total_costs = 0
        for cost in request.POST.getlist('our_costs[]'):
            total_costs = total_costs + Decimal(cost)
        vat_amount = round(total_costs * CURRENT_VAT_RATE, 2)
        total_costs_and_vat = vat_amount + total_costs

        request_post_copy['created_by'] = request.user
        request_post_copy['file_number'] = file_number_id
        request_post_copy['vat'] = str(vat_amount)

        form = InvoicesForm(request_post_copy)

        if form.is_valid():

            invoice_instance = form.save(commit=False)
            invoice_instance.save()

            disbs_ids = [int(id_str)
                         for id_str in request.POST.getlist('pink_slips[]')]
            total_disbs = 0
            for id in disbs_ids:
                obj = PmtsSlips.objects.filter(id=id).first()
                total_disbs = total_disbs + obj.amount
                obj.amount_invoiced = json.dumps(str(obj.amount))
                obj.balance_left = 0
                obj.save()

            total_costs_and_disbs = total_costs_and_vat + total_disbs
            temp_costs = total_costs_and_disbs

            green_slips_ids = [int(id_str)
                               for id_str in request.POST.getlist('green_slips[]')]
            for id in green_slips_ids:
                obj = LedgerAccountTransfers.objects.filter(id=id).first()
                if obj.file_number_from.file_number == file_number:
                    obj.amount_invoiced_from = json.dumps(str(obj.amount))
                    obj.balance_left_from = 0
                else:
                    temp_costs = temp_costs - obj.balance_left_to
                    invoice_slip_obj = {str(invoice_instance.id): {
                        'amt_invoiced': str(obj.balance_left_to), 'balance_left': ''}}
                    if temp_costs < 0:
                        obj.balance_left_to = obj.balance_left_to - total_costs_and_disbs
                    elif temp_costs >= 0:
                        obj.balance_left_to = 0
                    invoice_slip_obj[str(invoice_instance.id)]['balance_left'] = str(
                        obj.balance_left_to)
                    prev_amount_invoiced_to_obj = json.loads(
                        obj.amount_invoiced_to) if obj.amount_invoiced_to != {} else {}
                    prev_amount_invoiced_to_obj.update(invoice_slip_obj)
                    obj.amount_invoiced_to = json.dumps(
                        prev_amount_invoiced_to_obj)
                    total_costs_and_disbs = temp_costs
                obj.save()

            moa_ids = [int(id_str)
                       for id_str in request.POST.getlist('blue_slips[]')]
            for id in moa_ids:
                obj = PmtsSlips.objects.filter(id=id).first()
                temp_costs = temp_costs - obj.balance_left
                invoice_slip_obj = {str(invoice_instance.id): {
                    'amt_invoiced': str(obj.balance_left), 'balance_left': ''}}
                if temp_costs < 0:
                    obj.balance_left = obj.balance_left - total_costs_and_disbs
                elif temp_costs >= 0:
                    obj.balance_left = 0

                invoice_slip_obj[str(invoice_instance.id)
                                 ]['balance_left'] = str(obj.balance_left)
                prev_amount_invoiced_to_obj = json.loads(
                    obj.amount_invoiced) if obj.amount_invoiced != {} else {}
                prev_amount_invoiced_to_obj.update(invoice_slip_obj)
                obj.amount_invoiced = json.dumps(prev_amount_invoiced_to_obj)
                total_costs_and_disbs = temp_costs
                obj.save()

            if temp_costs <= 0:
                invoice_instance.total_due_left = 0
            else:
                invoice_instance.total_due_left = temp_costs

            invoice_instance.disbs_ids.set(disbs_ids)
            invoice_instance.moa_ids.set(moa_ids)
            invoice_instance.green_slip_ids.set(green_slips_ids)

            invoice_instance.save()
            messages.success(request, f'Invoice {
                             invoice_instance.invoice_number} successfully added. ')
            return redirect('finance_view', file_number=file_number)
        else:
            error_message = f'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        messages.error(request, 'Invalid request method.')

    return redirect('finance_view', file_number=file_number)


@login_required
def allocate_monies(request):
    if request.method == 'POST':

        amt_to_allocate_raw = request.POST['amt_to_allocate']
        invoice_num = request.POST['invoice_num']
        slip_id = request.POST['slip_id']

        invoice = Invoices.objects.filter(invoice_number=invoice_num).first()
        slip = PmtsSlips.objects.filter(id=slip_id).first()
        if not invoice or not slip:
            messages.error(request, 'Invoice or slip not found.')
            if invoice:
                return redirect('finance_view', file_number=invoice.file_number.file_number)
            return redirect('index')

        try:
            amt_to_allocate = Decimal(amt_to_allocate_raw)
        except Exception:
            messages.error(
                request, 'Please provide a valid amount to allocate.')
            return redirect('finance_view', file_number=invoice.file_number.file_number)

        if amt_to_allocate <= 0:
            messages.error(
                request, 'Amount to allocate must be greater than 0.')
            return redirect('finance_view', file_number=invoice.file_number.file_number)

        if amt_to_allocate > slip.balance_left:
            messages.error(
                request, 'Amount to allocate cannot exceed slip balance.')
            return redirect('finance_view', file_number=invoice.file_number.file_number)

        due_left = invoice.total_due_left
        approved_credit_total = get_invoice_approved_credit_total(invoice)
        effective_due_left = get_effective_invoice_due(
            invoice, approved_credit_total)
        if amt_to_allocate > effective_due_left:
            messages.error(
                request,
                f'Amount to allocate cannot exceed invoice due after credit notes (£{effective_due_left}).'
            )
            return redirect('finance_view', file_number=invoice.file_number.file_number)

        balance = due_left - amt_to_allocate

        # Store previous values for modification tracking
        prev_slip_amount_allocated = slip.amount_allocated
        prev_slip_balance_left = slip.balance_left

        already_allocated = json.loads(
            slip.amount_allocated) if slip.amount_allocated != {} else {}
        already_allocated.update({str(invoice.id): str(amt_to_allocate_raw)})
        slip.amount_allocated = json.dumps(already_allocated)

        # Calculate total allocated amount from the slip
        total_allocated = sum(Decimal(str(amt))
                              for amt in already_allocated.values())

        # Fix balance calculation: original amount - total allocated
        slip.balance_left = slip.amount - total_allocated

        invoice.cash_allocated_slips.add(slip.id)
        if balance <= 0:
            invoice.total_due_left = 0
        else:
            invoice.total_due_left = balance

        # Create modification record for invoice
        invoice_changes = {'prev_total_due_left': str(due_left),
                           'after_total_due_left': str(invoice.total_due_left),
                           'amount_allocated': str(amt_to_allocate),
                           'approved_credit_notes_total': str(approved_credit_total)}

        create_modification(
            user=request.user,
            modified_obj=invoice,
            changes=invoice_changes
        )

        # Create modification record for slip
        slip_changes = {'prev_amount_allocated': prev_slip_amount_allocated,
                        'after_amount_allocated': slip.amount_allocated,
                        'prev_balance_left': str(prev_slip_balance_left),
                        'after_balance_left': str(slip.balance_left),
                        'amount_allocated_to_invoice': str(amt_to_allocate),
                        'invoice_number': invoice_num}

        create_modification(
            user=request.user,
            modified_obj=slip,
            changes=slip_changes
        )

        slip.save()
        invoice.save()
        messages.success(
            request, f"£{amt_to_allocate} successfully allocated to Invoice {invoice_num}.")

    return redirect('finance_view', invoice.file_number.file_number)


@login_required
def download_invoice(request, id):

    invoice = Invoices.objects.filter(id=id).first()

    our_costs = invoice.our_costs

    costs = ast.literal_eval(our_costs) if type(
        our_costs) != type([]) else our_costs
    total_cost_invoice = 0

    file_details_display = f"<tr><td><b>Our Ref:</b>{
        invoice.file_number.file_number}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td><b>Invoice No:</b>{
            invoice.invoice_number}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td><b>Date:</b>{invoice.date.strftime(
            '%d/%m/%Y')}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td>&nbsp;</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td><b>Private & Confidential</b></td><td></td></tr>"
    file_details_display = file_details_display + \
        f"""<tr><td class='d-flex flex-row'>
        <div class="me-4 " >{invoice.file_number.client1.name}<br>
        {invoice.file_number.client1.address_line1}<br>
        {invoice.file_number.client1.address_line2}<br>
        {invoice.file_number.client1.county}, {invoice.file_number.client1.postcode}
        </div>"""
    if invoice.file_number.client2:
        file_details_display = file_details_display + f"""
            <div class="border-start ps-4">{invoice.file_number.client2.name}<br>
            {invoice.file_number.client2.address_line1}<br>
            {invoice.file_number.client2.address_line2}<br>
            {invoice.file_number.client2.county}, {invoice.file_number.client2.postcode}
            </div>"""
    file_details_display = file_details_display + """ </td><td></td></tr>"""
    if invoice.payable_by == 'Client':
        payable_by = "&nbsp;"
    else:
        payable_by = invoice.payable_by
        payable_by = f"<tr><td><b>Payable by: </b>{
            payable_by}</td><td></td></tr>"

    file_details_display = file_details_display + payable_by
    if invoice.by_email == True and invoice.by_post == True:
        send_via = f'<b>By post and email to: </b>{
            invoice.file_number.client1.email}'
    elif invoice.by_email == True:
        send_via = f'<b>By email to: </b>{invoice.file_number.client1.email}'
    elif invoice.by_post == True:
        send_via = f'<b>By post</b>'
    else:
        send_via = ""

    file_details_display = file_details_display + \
        f"<tr><td><b>Re: </b>{invoice.file_number.matter_description}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td colspan='2' style='text-align: right;'>{send_via}</td></tr>"

    desc_and_cost_display = "<tr><td>&nbsp;</td><td></td></tr>"
    desc_and_cost_display = desc_and_cost_display + \
        f"<tr><td style='text-align: justify; text-justify: inter-word;' colspan='2'>{
            invoice.description}</td></tr>"
    desc_and_cost_display = desc_and_cost_display + \
        "<tr><td>&nbsp;</td><td></td></tr>"

    our_costs_desc_pre = invoice.our_costs_desc
    our_costs_desc = ast.literal_eval(our_costs_desc_pre) if type(
        our_costs_desc_pre) != type([]) else our_costs_desc_pre
    costs_display = ""
    for i in range(len(costs)):
        total_cost_invoice = total_cost_invoice + Decimal(costs[i])
        costs_display = costs_display + \
            f"<tr><td><b>{
                our_costs_desc[i]}</b>:</td><td style='text-align: center;"
        if i == 0:
            costs_display = costs_display + f"border-top: solid; border-top-width: thin;'"
        else:
            costs_display = costs_display + "'"
        costs_display = costs_display + \
            f">£{round(Decimal(costs[i]), 2)}</td></tr>"
    _, vat_inv, total_cost_and_vat = calculate_invoice_total_with_vat(invoice)
    costs_display = costs_display + \
        f"<tr><td >Add VAT @{CURRENT_VAT_RATE_PERCENT}%:</td><td style='text-align: center; border-top: solid; border-top-width: thin;'>£{
            round(vat_inv, 2)}</td></tr>"
    total_cost_and_vat = round(total_cost_and_vat, 2)
    costs_display = costs_display + \
        f"<tr><td ><b>Total Costs and VAT:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{
            total_cost_and_vat}</td></tr>"
    costs_display = costs_display + f"<tr><td>&nbsp;</td><td></td></tr>"

    desc_and_cost_display = desc_and_cost_display + costs_display

    total_pink_slips = 0
    if invoice.disbs_ids.exists():
        pink_slips_display = "<tr><td colspan='2'><b>Add Disbursement</b></td><tr>"
        for slip in invoice.disbs_ids.all():
            date = slip.date.strftime('%d/%m/%Y')
            pink_slips_display = pink_slips_display + \
                f"<tr><td>{
                    slip.description} - {date}</td><td style='text-align: center;"
            if total_pink_slips == 0:
                pink_slips_display = pink_slips_display + \
                    "border-top: solid; border-top-width: thin;'"
            else:
                pink_slips_display = pink_slips_display + "'"

            total_pink_slips = total_pink_slips + slip.amount
            pink_slips_display = pink_slips_display + \
                f">£{slip.amount}</td></tr>"
        pink_slips_display = pink_slips_display + \
            f"<tr><td><b>Total Disbursements:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{
                total_pink_slips}</td></tr>"

        pink_slips_display = pink_slips_display + f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        pink_slips_display = ''

    total_blue_slips = 0
    if invoice.moa_ids.exists():
        blue_slips_display = "<tr><td colspan='2'><b>Less Monies Received</b></td></tr>"
        for slip in invoice.moa_ids.all():

            if isinstance(slip.amount_invoiced, str):
                amount_invoiced = json.loads(slip.amount_invoiced)
            elif isinstance(slip.amount_invoiced, (bytes, bytearray)):
                amount_invoiced = json.loads(
                    slip.amount_invoiced.decode('utf-8'))
            elif isinstance(slip.amount_invoiced, dict):
                amount_invoiced = slip.amount_invoiced
            else:
                raise ValueError("Unsupported type for slip.amount_invoiced")
            date = slip.date.strftime('%d/%m/%Y')
            amt = amount_invoiced[f"{invoice.id}"]['amt_invoiced']

            blue_slips_display = blue_slips_display + f"<tr><td>Remittance {
                date} - balance of monies remaining on account</td><td style='text-align: center;"
            if total_blue_slips == 0:
                blue_slips_display = blue_slips_display + \
                    f"border-top: solid; border-top-width: thin;'"
            else:
                blue_slips_display = blue_slips_display + f"'"
            total_blue_slips = total_blue_slips + Decimal(amt)
            blue_slips_display = blue_slips_display + f" >£{amt}</td></tr>"
        blue_slips_display = blue_slips_display + \
            f"<tr><td><b>Total Monies Received:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{
                round(total_blue_slips, 2)}</td></tr>"
        blue_slips_display = blue_slips_display + f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        blue_slips_display = ''

    total_green_slips = 0
    if invoice.green_slip_ids.exists():
        green_slips_display = "<tr><td colspan='2'><b>Inter Matter(s) Transfers</b></td></tr>"
        for slip in invoice.green_slip_ids.all():
            date = slip.date.strftime('%d/%m/%Y')
            if slip.file_number_from.file_number == invoice.file_number.file_number:

                green_slips_display = green_slips_display + \
                    f"<tr><td>Transfer to {
                        slip.file_number_to} - {date}</td><td style='text-align: center;"
                if total_green_slips == 0:
                    green_slips_display = green_slips_display + \
                        "border-top: solid; border-top-width: thin;'"
                else:
                    green_slips_display = green_slips_display + "'"
                total_green_slips = total_green_slips - slip.amount
                green_slips_display = green_slips_display + \
                    f">£{slip.amount}</td></tr>"
            else:
                amount_invoiced = json.loads(slip.amount_invoiced_to)

                date = slip.date.strftime('%d/%m/%Y')
                amt = amount_invoiced[f"{invoice.id}"]['amt_invoiced']
                total_green_slips = total_green_slips + Decimal(amt)
                green_slips_display = green_slips_display + f"<tr><td>Transfer from {
                    slip.file_number_from} - {date}</td><td style='text-align: center;"
                if total_green_slips == 0:
                    green_slips_display = green_slips_display + \
                        "border-top: solid; border-top-width: thin;'"
                else:
                    green_slips_display = green_slips_display + "'"
                green_slips_display = green_slips_display + \
                    f">£{amt}</td></tr>"
        green_slips_display = green_slips_display + \
            f"<tr><td ><b>Total Green Slips:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{
                total_green_slips}</td></tr>"
        green_slips_display = green_slips_display + f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        green_slips_display = ""

    balance = (total_cost_and_vat + total_pink_slips) - \
        total_blue_slips - total_green_slips
    balance = round(balance, 2)
    if balance >= 0:
        total_due_display = f"<tr class='mt-5'><td><b>Total Due:</b></td><td style='text-align: center; border-top: solid; border-top-width: thin; border-bottom: solid;  border-bottom-style:double;'>£{balance}</td></tr>"
        bank_details = f"""
                <tr>
                        <td>&nbsp;</td>
                        <td></td>

                        </tr>
                    <tr>
                    <td style=" font-size: 12px"><b>Account Name:</b> ANP Solicitors Limited; <b>Sort Code:</b> 20-70-93; <b>Account No:</b> 13065049;  <b>Ref:</b>{invoice.file_number.file_number}<td>
                    </tr>
                 """
    else:
        balance = balance * -1
        bank_details = ""
        total_due_display = f"<tr><td><b>Balance Remaining on Account</b>&nbsp;</td><td style='text-align: center; border-top: solid; border-top-width: thin; border-bottom: solid;  border-bottom-style:double;'>£{
            balance}</td></tr>"

    if invoice.state == "D":
        state = """
                <div>
                    <h1 class="position-fixed top-50 start-50 translate-middle z-n1 text-secondary opacity-50 text-center strong" style="font-size: 1200%;">
                        DRAFT
                    <h1>
                </div>
                """
    else:
        state = ""

    footer = """
            ANP Solicitors is a trading name of ANP Solicitors Limited<br>
            Registered in England and Wales – Company No: 6948759 | Registered office at 290 Kiln Road, Benfleet, Essex SS7 1QT<br>
            T: 01702 556688 | F: 01702 556696 | E: info@anpsolicitors.com | www.anpsolicitors.com<br>
            This firm is authorised and regulated by the Solicitors Regulatory Authority<br>
            A list of directors is open to inspection at the office<br>
            VAT No. 977 542 767 | SRA No. 515388<br>
            """
    style = """
            @page :first {
                    size: A4;
                    margin-top: 0mm;
                    margin-bottom: 4px; 
                    margin-left: 40px;
                    margin-right: 40px;
            }
            @page {
                    size: A4; 
                    margin-top: 20px;
                    margin-bottom: 4px; 
                    margin-left: 40px;
                    margin-right: 40px;
            }
            .logoDiv{
                position: absolute;
                top: 15px;
                left: 40px;
                right: 40px;
                z-index: 1000;
                width: auto;
                text-align: right;
                margin: 0;
                padding: 0;
            }
            .overflow-auto {
                padding-top: 0;
            }
            /* Ensure table respects page margins */
            table {
                margin-top: 0;
            }
            @media print {
                /* Logo only appears on first page - absolute positioning */
                .logoDiv {
                    position: absolute;
                    top: 15px;
                    left: 40px;
                    right: 40px;
                    width: auto;
                    text-align: right;
                }
                /* First page: no top margin, content can start at top */
                @page :first {
                    size: A4;
                    margin-top: 0mm;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
                }
                /* Subsequent pages: small top margin for spacing */
                @page {
                    size: A4;
                    margin-top: 20px;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
                }
            }
            .logoDiv img {
                width: 180px;
                height: auto;
                margin: 0;
                padding: 0;
                display: inline-block;
            }
            
            """

    html = render_to_string('download_templates/invoice.html', {'invoice_number': invoice.invoice_number,
                                                                'style': style,
                                                                'state': mark_safe(state),
                                                                'file_details_display': mark_safe(file_details_display),
                                                                'desc_and_cost_display': mark_safe(desc_and_cost_display),
                                                                'pink_slips_display': mark_safe(pink_slips_display),
                                                                'blue_slips_display': mark_safe(blue_slips_display),
                                                                'green_slips_display': mark_safe(green_slips_display),
                                                                'total_due_display': mark_safe(total_due_display),
                                                                'bank_details': mark_safe(bank_details),
                                                                'footer': mark_safe(footer)})

    pdf_file = HTML(
        string=html, base_url=request.build_absolute_uri()).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Invoice {invoice.invoice_number} - {
        invoice.file_number.client1.name} ({invoice.file_number.matter_description}).pdf"'
    return response


@login_required
def download_credited_invoice(request, id):
    invoice = get_object_or_404(
        Invoices.objects.select_related(
            'file_number',
            'file_number__client1',
            'file_number__client2',
        ).prefetch_related(
            'disbs_ids',
            'moa_ids',
            'green_slip_ids',
            'cash_allocated_slips',
            'credit_notes',
        ),
        id=id,
    )

    def parse_json_dict(value):
        if isinstance(value, str):
            return json.loads(value) if value else {}
        if isinstance(value, (bytes, bytearray)):
            return json.loads(value.decode('utf-8'))
        if isinstance(value, dict):
            return value
        if value in (None, ''):
            return {}
        raise ValueError("Unsupported JSON value type")

    file_details_display = f"<tr><td><b>Our Ref:</b>{invoice.file_number.file_number}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td><b>Invoice No:</b>{invoice.invoice_number}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td><b>Date:</b>{invoice.date.strftime('%d/%m/%Y')}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td>&nbsp;</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td><b>Private & Confidential</b></td><td></td></tr>"
    file_details_display = file_details_display + \
        f"""<tr><td class='d-flex flex-row'>
        <div class="me-4 " >{invoice.file_number.client1.name}<br>
        {invoice.file_number.client1.address_line1}<br>
        {invoice.file_number.client1.address_line2}<br>
        {invoice.file_number.client1.county}, {invoice.file_number.client1.postcode}
        </div>"""
    if invoice.file_number.client2:
        file_details_display = file_details_display + f"""
            <div class="border-start ps-4">{invoice.file_number.client2.name}<br>
            {invoice.file_number.client2.address_line1}<br>
            {invoice.file_number.client2.address_line2}<br>
            {invoice.file_number.client2.county}, {invoice.file_number.client2.postcode}
            </div>"""
    file_details_display = file_details_display + """ </td><td></td></tr>"""
    if invoice.payable_by == 'Client':
        payable_by = "&nbsp;"
    else:
        payable_by = invoice.payable_by
        payable_by = f"<tr><td><b>Payable by: </b>{payable_by}</td><td></td></tr>"

    file_details_display = file_details_display + payable_by
    if invoice.by_email == True and invoice.by_post == True:
        send_via = f'<b>By post and email to: </b>{invoice.file_number.client1.email}'
    elif invoice.by_email == True:
        send_via = f'<b>By email to: </b>{invoice.file_number.client1.email}'
    elif invoice.by_post == True:
        send_via = f'<b>By post</b>'
    else:
        send_via = ""

    file_details_display = file_details_display + \
        f"<tr><td><b>Re: </b>{invoice.file_number.matter_description}</td><td></td></tr>"
    file_details_display = file_details_display + \
        f"<tr><td colspan='2' style='text-align: right;'>{send_via}</td></tr>"

    desc_and_cost_display = "<tr><td>&nbsp;</td><td></td></tr>"
    desc_and_cost_display = desc_and_cost_display + \
        f"<tr><td style='text-align: justify; text-justify: inter-word;' colspan='2'>{invoice.description}</td></tr>"
    desc_and_cost_display = desc_and_cost_display + \
        "<tr><td>&nbsp;</td><td></td></tr>"

    costs = ast.literal_eval(invoice.our_costs) if type(
        invoice.our_costs) != type([]) else invoice.our_costs
    our_costs_desc = ast.literal_eval(invoice.our_costs_desc) if type(
        invoice.our_costs_desc) != type([]) else invoice.our_costs_desc
    costs_display = ""
    for i in range(len(costs)):
        costs_display = costs_display + \
            f"<tr><td><b>{our_costs_desc[i]}</b>:</td><td style='text-align: center;"
        if i == 0:
            costs_display = costs_display + "border-top: solid; border-top-width: thin;'"
        else:
            costs_display = costs_display + "'"
        costs_display = costs_display + \
            f">£{round(Decimal(costs[i]), 2)}</td></tr>"
    _, vat_inv, total_cost_and_vat = calculate_invoice_total_with_vat(invoice)
    costs_display = costs_display + \
        f"<tr><td >Add VAT @{CURRENT_VAT_RATE_PERCENT}%:</td><td style='text-align: center; border-top: solid; border-top-width: thin;'>£{round(vat_inv, 2)}</td></tr>"
    total_cost_and_vat = round(total_cost_and_vat, 2)
    costs_display = costs_display + \
        f"<tr><td ><b>Total Costs and VAT:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{total_cost_and_vat}</td></tr>"
    costs_display = costs_display + f"<tr><td>&nbsp;</td><td></td></tr>"
    desc_and_cost_display = desc_and_cost_display + costs_display

    total_pink_slips = Decimal('0')
    if invoice.disbs_ids.exists():
        pink_slips_display = "<tr><td colspan='2'><b>Add Disbursement</b></td><tr>"
        for slip in invoice.disbs_ids.all():
            date = slip.date.strftime('%d/%m/%Y')
            pink_slips_display = pink_slips_display + \
                f"<tr><td>{slip.description} - {date}</td><td style='text-align: center;"
            if total_pink_slips == 0:
                pink_slips_display = pink_slips_display + \
                    "border-top: solid; border-top-width: thin;'"
            else:
                pink_slips_display = pink_slips_display + "'"

            total_pink_slips = total_pink_slips + slip.amount
            pink_slips_display = pink_slips_display + \
                f">£{slip.amount}</td></tr>"
        pink_slips_display = pink_slips_display + \
            f"<tr><td><b>Total Disbursements:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{total_pink_slips}</td></tr>"
        pink_slips_display = pink_slips_display + f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        pink_slips_display = ''

    total_blue_slips = Decimal('0')
    if invoice.moa_ids.exists():
        blue_slips_display = "<tr><td colspan='2'><b>Less Monies Received</b></td></tr>"
        for slip in invoice.moa_ids.all():
            amount_invoiced = parse_json_dict(slip.amount_invoiced)
            date = slip.date.strftime('%d/%m/%Y')
            amt = Decimal(
                str(amount_invoiced[f"{invoice.id}"]['amt_invoiced']))

            blue_slips_display = blue_slips_display + \
                f"<tr><td>Remittance {date} - balance of monies remaining on account</td><td style='text-align: center;"
            if total_blue_slips == 0:
                blue_slips_display = blue_slips_display + \
                    f"border-top: solid; border-top-width: thin;'"
            else:
                blue_slips_display = blue_slips_display + f"'"
            total_blue_slips = total_blue_slips + amt
            blue_slips_display = blue_slips_display + f" >£{amt}</td></tr>"
        blue_slips_display = blue_slips_display + \
            f"<tr><td><b>Total Monies Received:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{round(total_blue_slips, 2)}</td></tr>"
        blue_slips_display = blue_slips_display + f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        blue_slips_display = ''

    total_green_slips = Decimal('0')
    if invoice.green_slip_ids.exists():
        green_slips_display = "<tr><td colspan='2'><b>Inter Matter(s) Transfers</b></td></tr>"
        for slip in invoice.green_slip_ids.all():
            date = slip.date.strftime('%d/%m/%Y')
            if slip.file_number_from.file_number == invoice.file_number.file_number:
                green_slips_display = green_slips_display + \
                    f"<tr><td>Transfer to {slip.file_number_to} - {date}</td><td style='text-align: center;"
                if total_green_slips == 0:
                    green_slips_display = green_slips_display + \
                        "border-top: solid; border-top-width: thin;'"
                else:
                    green_slips_display = green_slips_display + "'"
                total_green_slips = total_green_slips - slip.amount
                green_slips_display = green_slips_display + \
                    f">£{slip.amount}</td></tr>"
            else:
                amount_invoiced = parse_json_dict(slip.amount_invoiced_to)
                amt = Decimal(
                    str(amount_invoiced[f"{invoice.id}"]['amt_invoiced']))
                total_green_slips = total_green_slips + amt
                green_slips_display = green_slips_display + \
                    f"<tr><td>Transfer from {slip.file_number_from} - {date}</td><td style='text-align: center;"
                if total_green_slips == 0:
                    green_slips_display = green_slips_display + \
                        "border-top: solid; border-top-width: thin;'"
                else:
                    green_slips_display = green_slips_display + "'"
                green_slips_display = green_slips_display + \
                    f">£{amt}</td></tr>"
        green_slips_display = green_slips_display + \
            f"<tr><td ><b>Total Green Slips:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{total_green_slips}</td></tr>"
        green_slips_display = green_slips_display + f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        green_slips_display = ""

    total_cash_allocated_slips = Decimal('0')
    if invoice.cash_allocated_slips.exists():
        cash_allocated_slips_display = "<tr><td colspan='2'><b>Less Monies Received After Invoice Creation</b></td></tr>"
        for slip in invoice.cash_allocated_slips.all():
            amount_allocated = parse_json_dict(slip.amount_allocated)
            invoice_id_str = f"{invoice.id}"
            if invoice_id_str not in amount_allocated:
                continue
            date = slip.date.strftime('%d/%m/%Y')
            amt = Decimal(str(amount_allocated[invoice_id_str]))
            cash_allocated_slips_display = cash_allocated_slips_display + \
                f"<tr><td>Payment from {slip.pmt_person} - {date}</td><td style='text-align: center;"
            if total_cash_allocated_slips == 0:
                cash_allocated_slips_display = cash_allocated_slips_display + \
                    "border-top: solid; border-top-width: thin;'"
            else:
                cash_allocated_slips_display = cash_allocated_slips_display + "'"
            total_cash_allocated_slips = total_cash_allocated_slips + amt
            cash_allocated_slips_display = cash_allocated_slips_display + \
                f">£{amt}</td></tr>"
        cash_allocated_slips_display = cash_allocated_slips_display + \
            f"<tr><td><b>Total Post-Invoice Monies Received:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{round(total_cash_allocated_slips, 2)}</td></tr>"
        cash_allocated_slips_display = cash_allocated_slips_display + \
            f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        cash_allocated_slips_display = ""

    approved_credit_notes = invoice.credit_notes.filter(
        status='F').order_by('date', 'id')
    approved_credit_total = Decimal('0')
    if approved_credit_notes.exists():
        credit_notes_display = "<tr><td colspan='2'><b>Less Approved Credit Notes</b></td></tr>"
        for note in approved_credit_notes:
            approved_credit_total = approved_credit_total + note.amount
            credit_notes_display = credit_notes_display + \
                f"<tr><td>Credit Note CN-{note.id} - {note.date.strftime('%d/%m/%Y')}</td><td style='text-align: center;"
            if approved_credit_total == note.amount:
                credit_notes_display = credit_notes_display + \
                    "border-top: solid; border-top-width: thin;'"
            else:
                credit_notes_display = credit_notes_display + "'"
            credit_notes_display = credit_notes_display + \
                f">£{note.amount}</td></tr>"
        credit_notes_display = credit_notes_display + \
            f"<tr><td><b>Total Approved Credit Notes:</b></td><td style='text-align: center; border-bottom: solid; border-bottom-width: thin; border-top: solid; border-top-width: thin;'>£{round(approved_credit_total, 2)}</td></tr>"
        credit_notes_display = credit_notes_display + \
            f"<tr><td>&nbsp;</td><td></td></tr>"
    else:
        credit_notes_display = ""

    balance = (total_cost_and_vat + total_pink_slips) - \
        total_blue_slips - total_green_slips - \
        total_cash_allocated_slips - approved_credit_total
    balance = round(balance, 2)
    if balance >= 0:
        total_due_display = f"<tr class='mt-5'><td><b>Total Due:</b></td><td style='text-align: center; border-top: solid; border-top-width: thin; border-bottom: solid;  border-bottom-style:double;'>£{balance}</td></tr>"
        bank_details = f"""
                <tr>
                        <td>&nbsp;</td>
                        <td></td>

                        </tr>
                    <tr>
                    <td style=" font-size: 12px"><b>Account Name:</b> ANP Solicitors Limited; <b>Sort Code:</b> 20-70-93; <b>Account No:</b> 13065049;  <b>Ref:</b>{invoice.file_number.file_number}<td>
                    </tr>
                 """
    else:
        balance = balance * -1
        bank_details = ""
        total_due_display = f"<tr><td><b>Balance Remaining on Account</b>&nbsp;</td><td style='text-align: center; border-top: solid; border-top-width: thin; border-bottom: solid;  border-bottom-style:double;'>£{balance}</td></tr>"

    if invoice.state == "D":
        state = """
                <div>
                    <h1 class="position-fixed top-50 start-50 translate-middle z-n1 text-secondary opacity-50 text-center strong" style="font-size: 1200%;">
                        DRAFT
                    <h1>
                </div>
                """
    else:
        state = ""

    footer = """
            ANP Solicitors is a trading name of ANP Solicitors Limited<br>
            Registered in England and Wales - Company No: 6948759 | Registered office at 290 Kiln Road, Benfleet, Essex SS7 1QT<br>
            T: 01702 556688 | F: 01702 556696 | E: info@anpsolicitors.com | www.anpsolicitors.com<br>
            This firm is authorised and regulated by the Solicitors Regulation Authority<br>
            A list of directors is open to inspection at the office<br>
            VAT No. 977 542 767 | SRA No. 515388<br>
            """
    style = """
            @page :first {
                    size: A4;
                    margin-top: 0mm;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
            }
            @page {
                    size: A4;
                    margin-top: 20px;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
            }
            .logoDiv{
                position: absolute;
                top: 15px;
                left: 40px;
                right: 40px;
                z-index: 1000;
                width: auto;
                text-align: right;
                margin: 0;
                padding: 0;
            }
            .docTitle {
                text-align: center;
                font-size: 28px;
                font-weight: bold;
                margin-top: 4px;
                margin-bottom: 8px;
            }
            .overflow-auto {
                padding-top: 0;
            }
            table {
                margin-top: 0;
            }
            @media print {
                .logoDiv {
                    position: absolute;
                    top: 15px;
                    left: 40px;
                    right: 40px;
                    width: auto;
                    text-align: right;
                }
                @page :first {
                    size: A4;
                    margin-top: 0mm;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
                }
                @page {
                    size: A4;
                    margin-top: 20px;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
                }
            }
            .logoDiv img {
                width: 180px;
                height: auto;
                margin: 0;
                padding: 0;
                display: inline-block;
            }
            """

    html = render_to_string('download_templates/credited_invoice.html', {
        'invoice_number': invoice.invoice_number,
        'style': style,
        'state': mark_safe(state),
        'file_details_display': mark_safe(file_details_display),
        'desc_and_cost_display': mark_safe(desc_and_cost_display),
        'pink_slips_display': mark_safe(pink_slips_display),
        'blue_slips_display': mark_safe(blue_slips_display),
        'green_slips_display': mark_safe(green_slips_display),
        'cash_allocated_slips_display': mark_safe(cash_allocated_slips_display),
        'credit_notes_display': mark_safe(credit_notes_display),
        'total_due_display': mark_safe(total_due_display),
        'bank_details': mark_safe(bank_details),
        'footer': mark_safe(footer),
    })

    pdf_file = HTML(
        string=html, base_url=request.build_absolute_uri()).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response[
        'Content-Disposition'] = f'attachment; filename="Credited Invoice {invoice.invoice_number} - {invoice.file_number.client1.name} ({invoice.file_number.matter_description}).pdf"'
    return response


@login_required
def download_credit_note(request, id):
    credit_note = get_object_or_404(CreditNote, id=id)
    file_obj = credit_note.file_number
    invoice = credit_note.invoice
    net_amount, vat_amount, gross_amount = calculate_credit_note_breakdown(
        credit_note.amount)
    status_display = dict(CreditNote.STATUSES).get(
        credit_note.status, credit_note.status)

    client_address_block = f"""
        <div class="me-4">{file_obj.client1.name}<br>
        {file_obj.client1.address_line1}<br>
        {file_obj.client1.address_line2}<br>
        {file_obj.client1.county}, {file_obj.client1.postcode}
        </div>
    """
    if file_obj.client2:
        client_address_block = client_address_block + f"""
            <div class="border-start ps-4">{file_obj.client2.name}<br>
            {file_obj.client2.address_line1}<br>
            {file_obj.client2.address_line2}<br>
            {file_obj.client2.county}, {file_obj.client2.postcode}
            </div>
        """

    file_details_display = f"""
        <tr><td><b>Our Ref:</b> {file_obj.file_number}</td><td></td></tr>
        <tr><td><b>Credit Note No:</b> CN-{credit_note.id}</td><td></td></tr>
        <tr><td><b>Date:</b> {credit_note.date.strftime('%d/%m/%Y')}</td><td></td></tr>
        <tr><td>&nbsp;</td><td></td></tr>
        <tr><td><b>Private & Confidential</b></td><td></td></tr>
        <tr><td class='d-flex flex-row'>{client_address_block}</td><td></td></tr>
        <tr><td><b>Re:</b> {file_obj.matter_description}</td><td></td></tr>
        <tr><td>&nbsp;</td><td></td></tr>
    """

    amount_summary_display = f"""
        <div>
            <table style='width: 100%; border-collapse: collapse; table-layout: fixed;'>
                <tr>
                    <td style='padding: 4px 0; text-align: left;'><b>Credit Amount Excl. VAT:</b></td>
                    <td style='padding: 4px 0; text-align: right; width: 220px;'>
                        <span style='display: inline-block; border-top: solid 1px; padding-top: 2px;'>£{net_amount}</span>
                    </td>
                </tr>
                <tr>
                    <td style='padding: 4px 0; text-align: left;'><b>VAT @{CURRENT_VAT_RATE_PERCENT}% (included):</b></td>
                    <td style='padding: 4px 0; text-align: right;'>£{vat_amount}</td>
                </tr>
                <tr>
                    <td style='padding: 4px 0; text-align: left;'><b>Total Credit (Incl. VAT):</b></td>
                    <td style='padding: 4px 0; text-align: right;'>
                        <span style='display: inline-block; border-top: solid 1px; border-bottom: solid 1px; padding: 2px 0;'>£{gross_amount}</span>
                    </td>
                </tr>
            </table>
        </div>
    """
    desc_and_cost_display = f"""
        <tr><td colspan='2'><b>CREDIT NOTE FOR THE PARTIAL/FULL AMOUNT</b></td></tr>
        <tr><td colspan='2' style='text-align: justify; text-justify: inter-word;'><b>Reason:</b> {credit_note.reason}</td></tr>
        <tr><td>&nbsp;</td><td></td></tr>
        <tr><td colspan='2'>{amount_summary_display}</td></tr>
    """

    meta_line = f"Status: {status_display}"
    if credit_note.approved_by:
        approved_on_display = credit_note.approved_on.strftime(
            '%d/%m/%Y %H:%M') if credit_note.approved_on else ''
        meta_line = f"{meta_line} | Approved: {credit_note.approved_by}"
        if approved_on_display:
            meta_line = f"{meta_line} ({approved_on_display})"

    if credit_note.status == "P":
        state = """
                <div>
                    <h1 class="position-fixed top-50 start-50 translate-middle z-n1 text-secondary opacity-50 text-center strong" style="font-size: 600%;">
                        PENDING APPROVAL
                    <h1>
                </div>
                """
    else:
        state = ""

    footer = """
            ANP Solicitors is a trading name of ANP Solicitors Limited<br>
            Registered in England and Wales - Company No: 6948759 | Registered office at 290 Kiln Road, Benfleet, Essex SS7 1QT<br>
            T: 01702 556688 | F: 01702 556696 | E: info@anpsolicitors.com | www.anpsolicitors.com<br>
            This firm is authorised and regulated by the Solicitors Regulation Authority<br>
            A list of directors is open to inspection at the office<br>
            VAT No. 977 542 767 | SRA No. 515388<br>
            """
    style = """
            @page :first {
                    size: A4;
                    margin-top: 0mm;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
            }
            @page {
                    size: A4;
                    margin-top: 20px;
                    margin-bottom: 4px;
                    margin-left: 40px;
                    margin-right: 40px;
            }
            .logoDiv{
                position: absolute;
                top: 15px;
                right: 40px;
                z-index: 1000;
                width: 75px;
                height: 50px;
                margin: 0;
                padding: 0;
            }
            img {
                width: 75px;
                height: 50px;
                margin: 0;
                padding: 0;
                display: block;
            }
            .docTitle {
                text-align: center;
                font-size: 32px;
                font-weight: bold;
                margin-top: 4px;
                margin-bottom: 12px;
            }
            .creditTable {
                width: 100%;
                table-layout: fixed;
            }
            """

    html = render_to_string('download_templates/credit_note.html', {
        'credit_note_number': credit_note.id,
        'style': style,
        'state': mark_safe(state),
        'file_details_display': mark_safe(file_details_display),
        'desc_and_cost_display': mark_safe(desc_and_cost_display),
        'meta_line': meta_line,
        'footer': mark_safe(footer),
    })
    pdf_file = HTML(
        string=html, base_url=request.build_absolute_uri()).write_pdf()
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Credit_Note_CN-{credit_note.id}_{file_obj.file_number}.pdf"'
    return response


def get_all_financials(file_number):
    file = get_object_or_404(WIP, file_number=file_number)

    slips = PmtsSlips.objects.filter(file_number=file.id)
    green_slips = LedgerAccountTransfers.objects.filter(
        Q(file_number_from=file.id) | Q(file_number_to=file.id))
    invoices = Invoices.objects.filter(file_number=file.id)
    credit_notes = CreditNote.objects.filter(
        file_number=file.id,
        status='F'
    ).select_related('invoice')
    all_objects = []
    for slip in slips:
        if slip.is_money_out == True:
            type_obj = 'money_out'
            desc = f"Payment to {slip.pmt_person} - {slip.description}"
        else:
            type_obj = 'money_in'
            desc = f"Payment from {slip.pmt_person} - {slip.description}"

        obj = {'date': slip.date.strftime('%d/%m/%Y'),
               'desc': desc,
               'type': type_obj,
               'ledger': slip.ledger_account,
               'amount': slip.amount
               }
        all_objects.append(obj)

    for slip in green_slips:
        if slip.file_number_from == slip.file_number_to:
            type_obj = 'client_to_office_tfr'
            desc = f"Client to Office Transfer"
        elif slip.file_number_from.file_number == file_number:
            type_obj = 'money_out'
            desc = f"Transfer to {slip.file_number_to}"
        else:
            type_obj = 'money_in'
            desc = f"Transfer from {slip.file_number_from}"

        obj = {'date': slip.date.strftime('%d/%m/%Y'),
               'desc': desc,
               'type': type_obj,
               'ledger': slip.from_ledger_account,
               'amount': slip.amount
               }
        all_objects.append(obj)

    for invoice in invoices:
        if invoice.state == 'F':
            desc = f"ANP Invoice {invoice.invoice_number}"
        else:
            desc = f"DRAFT ANP Invoice"

        type_obj = "money_out"

        _, _, total_cost_invoice = calculate_invoice_total_with_vat(invoice)
        obj = {'date': invoice.date.strftime('%d/%m/%Y'),
               'desc': desc,
               'type': type_obj,
               'ledger': 'O',
               'amount': total_cost_invoice
               }
        all_objects.append(obj)

    for credit_note in credit_notes:
        invoice_number = credit_note.invoice.invoice_number or 'Draft'
        obj = {
            'date': credit_note.date.strftime('%d/%m/%Y'),
            'desc': f'Credit Note for ANP Invoice {invoice_number}',
            'type': 'money_in',
            'ledger': 'O',
            'amount': credit_note.amount
        }
        all_objects.append(obj)

    def sort_rows(rows):
        def get_sort_key(row):
            # Handling empty dates
            date_str = row['date'] if row['date'] else '01/01/0001'

            date_time = datetime.strptime(date_str, '%d/%m/%Y')
            return date_time

        sorted_rows = sorted(rows, key=get_sort_key)
        return sorted_rows

    sorted_rows = sort_rows(all_objects)
    return sorted_rows


@login_required
def download_statement_account(request, file_number):
    file = WIP.objects.filter(file_number=file_number).first()

    """
    Date | Desc | Money in | Money Out | Balance

    if pink slip, invoice or green_slip (file_from) then money out

    """
    sorted_rows = get_all_financials(file_number)
    now_time = datetime.now().strftime("%d_%m_%Y, %H_%M")
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="({
        file_number}) Statement of Account {now_time}.csv"'

    writer = csv.writer(response)

    writer.writerow(
        ['', f'Client Name: {file.client1.name} Matter:{file.matter_description}[{file.file_number}]'])
    writer.writerow(['', f'Statement of Account'])

    writer.writerow([])
    writer.writerow(
        ['Date', 'Description', 'Money In', 'Money Out', 'Balance'])
    balance = 0
    client_to_office_tfr_rows = 0
    for row in sorted_rows:
        if row['type'] == 'client_to_office_tfr':
            client_to_office_tfr_rows = client_to_office_tfr_rows + 1
            continue
        if row['type'] == 'money_out':
            balance = balance - row['amount']
        else:
            balance = balance + row['amount']
        writer.writerow([row['date'], row['desc'], row['amount'] if row['type'] ==
                        'money_in' else '', row['amount'] if row['type'] == 'money_out' else '', balance])
    writer.writerow([])
    final_cell = (len(sorted_rows)-client_to_office_tfr_rows) + 4
    writer.writerow(
        ['', 'Total', f'=sum(c5:c{final_cell})', f'=sum(d5:d{final_cell})'])
    return response


@login_required
def generate_ledgers_report(request, file_number):
    file = get_object_or_404(WIP, file_number=file_number)
    html_content = f"""
    <html>
        <head>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">

            <style>
                @page {{
                    size: landscape;
                }}
                .balance {{
                    font-weight: bold;
                }}
            </style>
        </head>
        <body style="font-family: Times New Roman, Times, serif">
            <h2>Ledger</h2>
            <h2>Client Name: {file.client1.name} Matter: {file.matter_description} [{file.file_number}]</h2>
            <table class='table table-striped'>
                <thead>
                    <tr>
                        <th>Date</th>
                        <th>Description</th>
                        <th colspan="2">Office Account</th>
                        <th colspan="2">Client Account</th>
                    </tr>
                    <tr>
                        <th></th>
                        <th></th>
                        <th>Amount</th>
                        <th>Balance</th>
                        <th>Amount</th>
                        <th>Balance</th>
                    </tr>
                </thead>
                <tbody>
    """
    office_balance = 0
    client_balance = 0
    all_objects = get_all_financials(file_number)

    for row in all_objects:

        if row['type'] == 'client_to_office_tfr':
            client_balance -= row['amount']
            client_amount = '-' + str(row['amount'])
            office_amount = ''
            html_content += f"""
                <tr>
                    <td>{row['date']}</td>
                    <td>{row['desc']}</td>
                    <td>{office_amount}</td>
                    <td class="balance">{office_balance}</td>
                    <td>{client_amount}</td>
                    <td class="balance">{client_balance}</td>
                </tr>
             """

            office_balance += row['amount']
            client_amount = ''
            office_amount = str(row['amount'])

            html_content += f"""
                <tr>
                    <td>{row['date']}</td>
                    <td>{row['desc']}</td>
                    <td>{office_amount}</td>
                    <td class="balance">{office_balance}</td>
                    <td>{client_amount}</td>
                    <td class="balance">{client_balance}</td>
                </tr>
            """
            continue

        if row['ledger'] == 'C':
            if row['type'] == 'money_out':
                client_balance -= row['amount']
                client_amount = '-' + str(row['amount'])
                office_amount = ''
            else:
                client_balance += row['amount']
                client_amount = str(row['amount'])
                office_amount = ''
        else:
            if row['type'] == 'money_out':
                office_balance -= row['amount']
                client_amount = ''
                office_amount = '-' + str(row['amount'])
            else:
                office_balance += row['amount']
                client_amount = ''
                office_amount = str(row['amount'])

        html_content += f"""
                <tr>
                    <td>{row['date']}</td>
                    <td>{row['desc']}</td>
                    <td>{office_amount}</td>
                    <td class="balance">{office_balance}</td>
                    <td>{client_amount}</td>
                    <td class="balance">{client_balance}</td>
                </tr>
        """

    html_content += f"""
                <tr class="bg-body-tertiary">
                    <td></td>
                    <td class="balance">Balance</td>
                    <td colspan='2' class="balance">{office_balance}</td>
                    <td colspan='2' class="balance">{client_balance}</td>
                </tr>
        """

    html_content += """
                </tbody>
            </table>
        </body>
    </html>
    """

    now_time = datetime.now().strftime("%d_%m_%Y_%H_%M")
    file_name = f"Ledger_Printout_{file_number}_{now_time}.pdf"

    pdf_file = HTML(string=html_content,
                    base_url=request.build_absolute_uri()).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')

    response['Content-Disposition'] = f'attachment; filename="{file_name}'
    return response


@login_required
def edit_invoice(request, id):
    if request.method == 'POST':

        invoice = Invoices.objects.filter(id=id).first()

        serializer = InvoicesSerializer(invoice)
        prev_serialized_data = serializer.to_dict()
        if invoice.state == 'D':
            if request.POST['state'] == 'F':
                largest_invoice_number = Invoices.objects.aggregate(Max('invoice_number'))[
                    'invoice_number__max']
                invoice.invoice_number = largest_invoice_number + 1
                invoice.state = 'F'

        invoice.payable_by = request.POST['payable_by']
        if 'by_email' in request.POST:
            invoice.by_email = True

        if 'by_post' in request.POST:
            invoice.by_post = True

        desc = request.POST['description']
        invoice.description = desc
        date = request.POST['date']
        invoice.date = date
        our_costs_desc = request.POST.getlist('our_costs_desc[]')
        our_costs = request.POST.getlist('our_costs[]')

        our_costs_desc_filtered = [
            desc for desc in our_costs_desc if desc != '']

        our_costs_filtered = [cost for cost in our_costs if cost != '']

        invoice.our_costs_desc = json.dumps(our_costs_desc_filtered)
        invoice.our_costs = json.dumps(our_costs_filtered)

        total_costs = 0
        for cost in our_costs_filtered:

            total_costs = total_costs + Decimal(cost)
        vat_mode = request.POST.get('vat_mode', 'auto')
        if vat_mode != 'manual':
            vat_mode = 'auto'
        auto_vat_amount = round(total_costs * CURRENT_VAT_RATE, 2)
        if vat_mode == 'manual':
            manual_vat_raw = request.POST.get('manual_vat', '').strip()
            try:
                vat_amount = round(Decimal(manual_vat_raw), 2)
                if vat_amount < 0:
                    raise ValueError
            except Exception:
                messages.error(
                    request, 'Manual VAT must be a valid non-negative amount.')
                return redirect('edit_invoice', id=id)
        else:
            vat_amount = auto_vat_amount
        total_costs_and_vat = vat_amount + total_costs
        invoice.vat = vat_amount
        invoice.vat_calculation_mode = vat_mode

        prev_pink_slip_ids = list(
            invoice.disbs_ids.values_list('id', flat=True))
        for id in prev_pink_slip_ids:
            slip = PmtsSlips.objects.filter(id=id).first()
            slip.balance_left = slip.amount
            slip.amount_invoiced = json.dumps({})
            slip.save()

        prev_moa_ids = list(invoice.moa_ids.values_list('id', flat=True))

        for id in prev_moa_ids:
            slip = PmtsSlips.objects.filter(id=id).first()
            if isinstance(slip.amount_invoiced, str):
                amount_invoiced = json.loads(slip.amount_invoiced)
            elif isinstance(slip.amount_invoiced, (bytes, bytearray)):
                amount_invoiced = json.loads(
                    slip.amount_invoiced.decode('utf-8'))
            elif isinstance(slip.amount_invoiced, dict):
                amount_invoiced = slip.amount_invoiced
            else:
                raise ValueError("Unsupported type for slip.amount_invoiced")
            inv_data = amount_invoiced.get(str(str(invoice.id)), {})
            amount_inv = inv_data.get('amt_invoiced', 0)
            balance_left = inv_data.get('balance_left', 0)

            difference = round(Decimal(amount_inv) - Decimal(balance_left), 2)

            amount_invoiced.pop(str(invoice.id), None)
            slip.amount_invoiced = json.dumps(amount_invoiced)
            slip.balance_left += difference
            slip.save()

        prev_green_slip_ids = list(
            invoice.green_slip_ids.values_list('id', flat=True))

        for id in prev_green_slip_ids:
            slip = LedgerAccountTransfers.objects.filter(id=id).first()
            if slip.file_number_to == invoice.file_number:
                amount_invoiced = json.loads(slip.amount_invoiced_to)
                inv_data = amount_invoiced.get(str(str(invoice.id)), {})
                amount_inv = inv_data.get('amt_invoiced', 0)
                balance_left = inv_data.get('balance_left', 0)

                difference = round(Decimal(amount_inv) -
                                   Decimal(balance_left), 2)

                amount_invoiced.pop(str(invoice.id), None)
                slip.amount_invoiced_to = json.dumps(amount_invoiced)
                slip.balance_left_to += difference
                slip.save()
            else:
                slip.balance_left_from = slip.amount
                slip.amount_invoiced_from = json.dumps({})

        disbs_ids = [int(id_str)
                     for id_str in request.POST.getlist('pink_slips[]')]
        total_disbs = 0
        for id in disbs_ids:
            obj = PmtsSlips.objects.filter(id=id).first()
            total_disbs = total_disbs + obj.amount
            obj.amount_invoiced = json.dumps(str(obj.amount))
            obj.balance_left = 0
            obj.save()

        total_costs_and_disbs = total_costs_and_vat + total_disbs
        temp_costs = total_costs_and_disbs

        moa_ids = [int(id_str)
                   for id_str in request.POST.getlist('blue_slips[]')]
        for id in moa_ids:
            obj = PmtsSlips.objects.filter(id=id).first()
            temp_costs = temp_costs - obj.balance_left
            invoice_slip_obj = {str(invoice.id): {'amt_invoiced': str(
                obj.balance_left), 'balance_left': ''}}
            if temp_costs < 0:
                obj.balance_left = obj.balance_left - total_costs_and_disbs
            elif temp_costs >= 0:
                obj.balance_left = 0

            invoice_slip_obj[str(invoice.id)]['balance_left'] = str(
                obj.balance_left)
            prev_amount_invoiced_to_obj = json.loads(
                obj.amount_invoiced) if obj.amount_invoiced != {} else {}
            prev_amount_invoiced_to_obj.update(invoice_slip_obj)
            obj.amount_invoiced = json.dumps(prev_amount_invoiced_to_obj)
            total_costs_and_disbs = temp_costs
            obj.save()

        green_slips_ids = [int(id_str)
                           for id_str in request.POST.getlist('green_slips[]')]
        for id in green_slips_ids:
            obj = LedgerAccountTransfers.objects.filter(id=id).first()
            if obj.file_number_from == invoice.file_number:
                obj.amount_invoiced_from = json.dumps(str(obj.amount))
                obj.balance_left_from = 0
            else:
                temp_costs = temp_costs - obj.balance_left_to
                invoice_slip_obj = {str(invoice.id): {'amt_invoiced': str(
                    obj.balance_left_to), 'balance_left': ''}}
                if temp_costs < 0:
                    obj.balance_left_to = obj.balance_left_to - total_costs_and_disbs
                elif temp_costs >= 0:
                    obj.balance_left_to = 0
                invoice_slip_obj[str(invoice.id)]['balance_left'] = str(
                    obj.balance_left_to)
                prev_amount_invoiced_to_obj = json.loads(
                    obj.amount_invoiced_to) if obj.amount_invoiced_to != {} else {}
                prev_amount_invoiced_to_obj.update(invoice_slip_obj)
                obj.amount_invoiced_to = json.dumps(
                    prev_amount_invoiced_to_obj)
                total_costs_and_disbs = temp_costs
            obj.save()

        if temp_costs <= 0:
            invoice.total_due_left = 0
        else:
            invoice.total_due_left = temp_costs

        invoice.disbs_ids.set(disbs_ids)
        invoice.moa_ids.set(moa_ids)
        invoice.green_slip_ids.set(green_slips_ids)

        invoice.save()
        serializer = InvoicesSerializer(invoice)
        after_serialized_data = serializer.to_dict()
        create_modification(
            request.user,
            invoice,
            json.dumps({'prev': prev_serialized_data,
                        'after': after_serialized_data})
        )

        return redirect('finance_view', file_number=invoice.file_number)
    else:
        invoice = Invoices.objects.filter(id=id).first()

        our_costs = invoice.our_costs

        costs = ast.literal_eval(our_costs) if type(
            our_costs) != type([]) else our_costs

        our_costs_desc_pre = invoice.our_costs_desc
        our_costs_desc = ast.literal_eval(our_costs_desc_pre) if type(
            our_costs_desc_pre) != type([]) else our_costs_desc_pre
        our_costs_rows = []
        for i in range(len(costs)):

            our_costs_display = f"""
            <div class="grid grid-cols-1 md:grid-cols-12 gap-4 mt-2">
                <div class="col-span-5">
                    <input type="text" class="form-input" id="our_costs_description" placeholder="Costs Description"
                    name="our_costs_desc[]" value="{our_costs_desc[i]}">
                </div>
                <div class="col-span-5">
                    <input required type="number" step="0.01" class="form-input" name="our_costs[]" id="our_costs"
                    placeholder="£0.00" value={round(Decimal(costs[i]), 2)}>
                </div>
                <div class="col-span-2 flex items-center">
                    <span type='button' class='btn btn-danger px-4' onclick="removeField(this);" >-</span>
                </div>
            </div>
            """
            our_costs_rows.append(mark_safe(our_costs_display))

        slips = PmtsSlips.objects.filter(
            file_number=invoice.file_number).order_by('date')

        green_slips_objs = LedgerAccountTransfers.objects.filter(
            Q(file_number_from=invoice.file_number.id) | Q(
                file_number_to=invoice.file_number.id)
        ).exclude(
            file_number_from=F('file_number_to')
        ).order_by('date')

        disbs_ids = list(invoice.disbs_ids.values_list('id', flat=True))

        # Get IDs of moa_ids
        moa_ids = list(invoice.moa_ids.values_list('id', flat=True))

        # Get IDs of cash_allocated_slips
        cash_allocated_slips_ids = list(
            invoice.cash_allocated_slips.values_list('id', flat=True))

        # Get IDs of green_slip_ids
        green_slip_ids = list(
            invoice.green_slip_ids.values_list('id', flat=True))
        pink_slips = []
        blue_slips = []
        green_slips = []

        for slip in slips:
            if slip.id in disbs_ids or slip.id in moa_ids:
                checked = 'checked'
            else:
                checked = ''
            if slip.ledger_account == 'C':
                ledger_acc = 'Client'
            else:
                ledger_acc = 'Office'
            if (slip.is_money_out == True and slip.balance_left > 0) or slip.id in disbs_ids:

                slip_display = f"""
                 <div class="form-check">
                    <input class="form-check-input" name="pink_slips[]" type="checkbox" value="{slip.id}" {checked}>
                    <label class="form-check-label" data-toggle="tooltip" data-bs-title="Description: '{slip.description}'">
                    Payment to&nbsp;<b >'{slip.pmt_person}'</b>&nbsp;of £{slip.amount}
                    from {ledger_acc} Ledger on {slip.date.strftime('%d/%m/%Y')}</label>
                </div>
                """
                pink_slips.append(mark_safe(slip_display))
            elif slip.id in moa_ids or slip.balance_left > 0:
                if isinstance(slip.amount_invoiced, str):
                    amount_invoiced = json.loads(slip.amount_invoiced)
                elif isinstance(slip.amount_invoiced, (bytes, bytearray)):
                    amount_invoiced = json.loads(
                        slip.amount_invoiced.decode('utf-8'))
                elif isinstance(slip.amount_invoiced, dict):
                    amount_invoiced = slip.amount_invoiced
                else:
                    raise ValueError(
                        "Unsupported type for slip.amount_invoiced")

                if slip.id in moa_ids:
                    amt = amount_invoiced[f"{invoice.id}"]['amt_invoiced']
                else:
                    amt = slip.balance_left
                slip_display = f"""
                 <div class="form-check">
                    <input class="form-check-input" name="blue_slips[]" type="checkbox" value="{slip.id}" {checked}>
                    <label class="form-check-label" data-toggle="tooltip" data-bs-title="Description: '{slip.description}'">
                    Payment from&nbsp;<b >'{slip.pmt_person}'</b>&nbsp;of £{amt} from (£{slip.amount})
                    from {ledger_acc} Ledger on {slip.date.strftime('%d/%m/%Y')}</label>
                </div>
                """
                blue_slips.append(mark_safe(slip_display))

        for slip in green_slips_objs:

            if slip.id in green_slip_ids:
                checked = "checked"
            else:
                checked = ""

            if slip.file_number_from == invoice.file_number:
                if slip.balance_left_from > 0 or slip.id in green_slip_ids:
                    slip_display = f"""
                                    <div class="form-check">
                                        <input class="form-check-input" name="green_slips[]" type="checkbox" value="{slip.id}" {checked}>
                                        <label class="form-check-label" data-toggle="tooltip" data-bs-title="Description: '{slip.description}'">
                                            Transfer to&nbsp;<b >{slip.file_number_to}</b> of £{slip.amount} on {slip.date.strftime('%d/%m/%Y')}
                                        </label>
                                    </div>
                                    """
                    green_slips.append(mark_safe(slip_display))
            if slip.file_number_to == invoice.file_number:
                if slip.id in green_slip_ids or slip.balance_left_to > 0:

                    amount_invoiced = json.loads(slip.amount_invoiced_to) if slip.amount_invoiced_to != {
                    } else slip.amount_invoiced_to
                    if slip.id in green_slip_ids:
                        amt = amount_invoiced[f"{invoice.id}"]['amt_invoiced']
                    else:
                        amt = slip.balance_left_to
                    slip_display = f"""
                                    <div class="form-check">
                                        <input class="form-check-input" name="green_slips[]" type="checkbox" value="{slip.id}" {checked}>
                                        <label class="form-check-label" data-toggle="tooltip" data-bs-title="Description: '{slip.description}'">
                                        Transfer from&nbsp;<b >{slip.file_number_from}</b> of £{amt} (from £{slip.amount}) {slip.date.strftime('%d/%m/%Y')}</label>
                                    </div>
                                    """
                    green_slips.append(mark_safe(slip_display))

        context = {
            'invoice_id': invoice.id,
            'file_number': invoice.file_number.file_number,
            'state': invoice.state,
            'date': invoice.date.strftime('%Y-%m-%d'),
            'payable_by': invoice.payable_by,
            'by_email': invoice.by_email,
            'by_post': invoice.by_post,
            'description': invoice.description,
            'our_costs_rows': our_costs_rows,
            'vat_mode': invoice.vat_calculation_mode or 'auto',
            'manual_vat': round(Decimal(str(invoice.vat or 0)), 2),
            'current_vat_rate_percent': CURRENT_VAT_RATE_PERCENT,
            'pink_slips': pink_slips,
            'blue_slips': blue_slips,
            'green_slips': green_slips,
            'green_slip_ids': green_slip_ids

        }
        return render(request, 'edit_invoice.html', context)


@login_required
def download_estate_accounts(request, file_number):
    file = WIP.objects.filter(file_number=file_number).first()
    slips = PmtsSlips.objects.filter(file_number=file.id)
    green_slips = LedgerAccountTransfers.objects.filter(
        Q(file_number_from=file.id) | Q(file_number_to=file.id))
    invoices = Invoices.objects.filter(file_number=file.id)
    credit_notes = CreditNote.objects.filter(
        file_number=file.id, status='F').select_related('invoice')
    money_in_objects = []
    money_out_objects = []
    num_money_in_objs = 0
    for slip in slips:
        if slip.is_money_out == True:

            desc = f"Payment to {slip.pmt_person} - {slip.description}"
            obj = {'date': slip.date.strftime('%d/%m/%Y'),
                   'desc': desc,
                   'amount': slip.amount
                   }
            money_out_objects.append(obj)
        else:

            desc = f"Payment from {slip.pmt_person} - {slip.description}"

            obj = {'date': slip.date.strftime('%d/%m/%Y'),
                   'desc': desc,
                   'amount': slip.amount
                   }
            money_in_objects.append(obj)

    for slip in green_slips:
        if slip.file_number_from.file_number == file_number:

            desc = f"Transfer to {slip.file_number_to}"
            obj = {'date': slip.date.strftime('%d/%m/%Y'),
                   'desc': desc,
                   'amount': slip.amount
                   }
            money_out_objects.append(obj)

        else:

            desc = f"Transfer from {slip.file_number_from}"
            obj = {'date': slip.date.strftime('%d/%m/%Y'),
                   'desc': desc,
                   'amount': slip.amount
                   }
            money_in_objects.append(obj)

    for invoice in invoices:
        if invoice.state == 'F':
            desc = f"ANP Invoice {invoice.invoice_number}"
        else:
            desc = f"DRAFT ANP Invoice"

        _, _, total_cost_invoice = calculate_invoice_total_with_vat(invoice)
        obj = {'date': invoice.date.strftime('%d/%m/%Y'),
               'desc': desc,
               'amount': total_cost_invoice
               }
        money_out_objects.append(obj)

    for credit_note in credit_notes:
        invoice_number = credit_note.invoice.invoice_number or "Draft"
        obj = {'date': credit_note.date.strftime('%d/%m/%Y'),
               'desc': f"Credit Note for ANP Invoice {invoice_number}",
               'amount': credit_note.amount
               }
        money_in_objects.append(obj)

    def sort_rows(rows):
        def get_sort_key(row):
            # Handling empty dates
            date_str = row['date'] if row['date'] else '01/01/0001'

            date_time = datetime.strptime(date_str, '%d/%m/%Y')
            return date_time

        sorted_rows = sorted(rows, key=get_sort_key)
        return sorted_rows

    sorted_money_in_objs = sort_rows(money_in_objects)
    sorted_money_out_objs = sort_rows(money_out_objects)

    doc = Document()

    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = 'Times New Roman'
            style.font.color.rgb = RGBColor(0, 0, 0)

    '''Page 1'''

    # Add heading on the first page
    doc.add_paragraph('\n\n\n\n')

    heading = doc.add_paragraph("Document Heading\n\n\n")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.font.size = Pt(46)

    doc.add_paragraph('')

    heading = doc.add_paragraph("Estate Account\n\n\n\n\n")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(32)

    heading = doc.add_paragraph(
        "Prepared by:\nANP Solicitors\n290 Kiln Road Benfleet\nEssex SS7 1QT")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(12)
    # Add page break
    doc.add_page_break()

    '''Page 2 Assets (money in)'''

    heading = doc.add_paragraph(
        "Person Decesased\nDate of Death\nEstate Account")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(16)

    heading = doc.add_paragraph("Assets")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(12)

    table = doc.add_table(rows=len(money_in_objects)+1,
                          cols=3, style='Light Shading')
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    i = 0
    j = 0

    for row in table.rows:
        if i == 0:
            row.cells[0].text = "Date"
            row.cells[1].text = "Description"
            row.cells[2].text = "Amount"
            i = 1
        elif j < len(money_in_objects):
            row.cells[0].text = money_in_objects[j]['date']
            row.cells[1].text = money_in_objects[j]['desc']
            row.cells[2].text = '£' + str(money_in_objects[j]['amount'])
            j = j + 1

    heading = doc.add_paragraph(
        "\nGross Value of Estate Carried Forward\t\t\t£1000.00")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(12)

    doc.add_page_break()

    '''Page 3'''

    table = doc.add_table(rows=len(money_out_objects)+1,
                          cols=3, style='Light Shading')
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    i = 0
    j = 0

    for row in table.rows:
        if i == 0:
            row.cells[0].text = "Date"
            row.cells[1].text = "Description"
            row.cells[2].text = "Amount"
            i = 1
        elif j < len(money_out_objects):
            row.cells[0].text = money_out_objects[j]['date']
            row.cells[1].text = money_out_objects[j]['desc']
            row.cells[2].text = '£' + str(money_out_objects[j]['amount'])
            j = j + 1

    heading = doc.add_paragraph(
        "Person Decesased\nDate of Death\nEstate Account")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(16)

    heading = doc.add_paragraph("Assets")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(12)

    heading = doc.add_paragraph(
        "\nGross Value of Estate Carried Forward\t\t\t£1000.00")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(12)
    # Create an in-memory stream to store the document
    output = io.BytesIO()

    # Save the document to the in-memory stream
    doc.save(output)

    # Create a response and set appropriate content type and headers
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="example.docx"'

    # Set the content of the response to the content of the in-memory stream
    response.write(output.getvalue())

    return response


@login_required
def unallocated_emails(request):

    unallocated_emails_obj = MatterEmails.objects.filter(
        file_number=None).order_by('time').only()
    i = 0
    rows = []

    def file_number_options():
        files = WIP.objects.all().order_by('file_number')
        select = f"""<input list="file_options"  type="text" name="FileNumber[]" class="form-input"
                 placeholder="XYZ0010001" pattern="^(XXXXXXXXXX|[A-Z]{{3}}\\d{{7}})$" 
               title="Must be either XXXXXXXXXX or 3 uppercase letters followed by 7 digits">
                    <datalist id="file_options">
                    <option value=''></option>
                    <option value="XXXXXXXXXX">To be deleted</option>
                """
        options = [
            f'<option value="{file.file_number}">{file.file_number}</option>' for file in files]
        select += ''.join(options)

        select = select + """</datalist>
                """
        return select
    files_options = file_number_options()
    for email in unallocated_emails_obj:
        i = i + 1
        receiver = json.loads(email.receiver)
        sender = json.loads(email.sender)

        row = f"""<tr class="email-row even:bg-white odd:bg-gray-50 border-b text-gray-900 px-2" data-email="{receiver[0]['emailAddress']['address']} {sender['emailAddress']['address']}">
                        <td class='td'>{i}</td>
                        <td class='td'>{email.time.strftime('%d-%m-%Y <br> %H:%M %p')}</td>
                        <td>
                            <b>From:</b> {sender['emailAddress']['name']} ({sender['emailAddress']['address']})<br>
                            <b>To:</b> {receiver[0]['emailAddress']['name']} ({receiver[0]['emailAddress']['address']})
                        </td>
                        <td class='td'>{email.subject}</td>
                        <td class='td'><a class="link" target="_blank" href="{email.link}">See Email</a></td>
                        <td>
                            <input class="hidden" name="email_ids[]" value={email.id}></input>
                            {files_options}
                        </td>
                    </tr>
        """
        rows.append(mark_safe(row))

    context = {
        'emails': rows
    }

    return render(request, 'unallocated_emails.html', context=context)


@login_required
def allocate_emails(request):
    i = 0
    file_numbers = request.POST.getlist('FileNumber[]')
    email_ids = request.POST.getlist('email_ids[]')
    j = 0
    error_count = 0

    for file_number in file_numbers:

        if file_number != '':
            file_number = file_number
            email_id = email_ids[i]
            email = MatterEmails.objects.filter(id=email_id).first()

            # Check if file exists
            file = WIP.objects.filter(file_number=file_number).first()
            if not file:
                error_count += 1
                messages.error(
                    request, f'File with file number {file_number} does not exist. Please choose a correct file number.')
                i += 1
                continue

            email.file_number = file
            j += 1
            email.fee_earner = file.fee_earner if file.fee_earner is not None else None
            email.save()

        i += 1

    if j > 0:
        messages.success(request, f'Successfully allocated {j} emails')

    if error_count > 0:
        messages.error(
            request, f'{error_count} file(s) were not found. Please check the file numbers.')

    return redirect('unallocated_emails')


def get_transfers_context():
    all_transfers = LedgerAccountTransfers.objects.filter(
        is_cashier_co_transfer=True,
        file_number_from=F('file_number_to'),
    ).filter(
        Q(from_ledger_account='C', to_ledger_account='O') |
        Q(from_ledger_account='O', to_ledger_account='C')
    ).select_related(
        'file_number_from',
        'created_by',
        'bank_transfer_done_by'
    ).order_by('-date', 'file_number_from__file_number', 'id')

    grouped = []
    current_group = None
    for t in all_transfers:
        if current_group is None or current_group['date'] != t.date:
            current_group = {
                'date': t.date,
                'co_rows': [],
                'oc_rows': [],
                'total': Decimal('0.00'),
                'co_total': Decimal('0.00'),
                'oc_total': Decimal('0.00'),
                'co_pending_count': 0,
                'oc_pending_count': 0,
            }
            grouped.append(current_group)

        is_co = t.from_ledger_account == 'C'
        if is_co:
            current_group['co_rows'].append(t)
            current_group['co_total'] += t.amount
            if not t.is_bank_transfer_done:
                current_group['co_pending_count'] += 1
        else:
            current_group['oc_rows'].append(t)
            current_group['oc_total'] += t.amount
            if not t.is_bank_transfer_done:
                current_group['oc_pending_count'] += 1
        current_group['total'] += t.amount

    return {
        'active_files': WIP.objects.order_by('file_number'),
        'grouped_transfers': grouped,
    }


@login_required
def download_cashier_data(request):
    if request.method == 'POST':
        cashier_action = request.POST.get('cashier_action')

        if cashier_action == 'export_client_to_office_transfers':
            start_date_str = request.POST.get('start_date', '').strip()
            end_date_str = request.POST.get('end_date', '').strip()

            if not start_date_str or not end_date_str:
                messages.error(
                    request, 'Start date and end date are required for export.')
                return redirect('download_cashier_data')

            try:
                start_date = datetime.strptime(
                    start_date_str, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Export date range is invalid.')
                return redirect('download_cashier_data')

            if end_date < start_date:
                messages.error(
                    request, 'End date cannot be before start date.')
                return redirect('download_cashier_data')

            transfers = LedgerAccountTransfers.objects.filter(
                is_cashier_co_transfer=True,
                from_ledger_account='C',
                to_ledger_account='O',
                file_number_from=F('file_number_to'),
                date__range=(start_date, end_date)
            ).select_related(
                'file_number_from',
                'created_by',
                'bank_transfer_done_by'
            ).order_by('date', 'file_number_from__file_number', 'id')

            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = (
                f'attachment; filename="CO_TFR_{start_date.strftime("%Y%m%d")}_{end_date.strftime("%Y%m%d")}.csv"'
            )

            writer = csv.writer(response)
            writer.writerow([
                'Date', 'File', 'Amount', 'Description', 'Created By',
                'Bank Status', 'Bank Done On', 'Bank Done By'
            ])

            if not transfers.exists():
                writer.writerow(
                    ['No C-O TFR rows found in selected date range.'])
                return response

            current_date = None
            group_total = Decimal('0.00')
            grand_total = Decimal('0.00')
            for transfer in transfers:
                if current_date is not None and current_date != transfer.date:
                    writer.writerow([
                        '', f'SUBTOTAL {current_date.strftime("%d/%m/%Y")}',
                        f'{group_total:.2f}', '', '', '', '', ''
                    ])
                    writer.writerow([])
                    group_total = Decimal('0.00')

                current_date = transfer.date
                group_total += transfer.amount
                grand_total += transfer.amount

                writer.writerow([
                    transfer.date.strftime('%d/%m/%Y'),
                    transfer.file_number_from.file_number if transfer.file_number_from else '-',
                    f'{transfer.amount:.2f}',
                    transfer.description,
                    str(transfer.created_by) if transfer.created_by else '-',
                    'Done' if transfer.is_bank_transfer_done else 'Pending',
                    transfer.bank_transfer_done_on.strftime(
                        '%d/%m/%Y') if transfer.bank_transfer_done_on else '-',
                    str(transfer.bank_transfer_done_by) if transfer.bank_transfer_done_by else '-'
                ])

            if current_date is not None:
                writer.writerow([
                    '', f'SUBTOTAL {current_date.strftime("%d/%m/%Y")}',
                    f'{group_total:.2f}', '', '', '', '', ''
                ])
            writer.writerow([])
            writer.writerow(
                ['', 'GRAND TOTAL', f'{grand_total:.2f}', '', '', '', '', ''])
            return response

        if cashier_action == 'add_client_to_office_transfers':
            row_dates = request.POST.getlist('row_date[]')
            row_file_ids = request.POST.getlist('row_file_id[]')
            row_amounts = request.POST.getlist('row_amount[]')
            row_descriptions = request.POST.getlist('row_description[]')

            max_rows = max(
                len(row_dates),
                len(row_file_ids),
                len(row_amounts),
                len(row_descriptions),
                0
            )

            created_count = 0
            errors = []
            for i in range(max_rows):
                row_date = row_dates[i].strip() if i < len(row_dates) else ''
                row_file_id = row_file_ids[i].strip(
                ) if i < len(row_file_ids) else ''
                row_amount = row_amounts[i].strip(
                ) if i < len(row_amounts) else ''
                row_description = row_descriptions[i].strip(
                ) if i < len(row_descriptions) else ''

                has_core_values = any([row_file_id, row_amount])
                has_custom_description = row_description not in ['', 'C-O TFR']
                if not has_core_values and not has_custom_description:
                    continue

                if not row_date or not row_file_id or not row_amount:
                    errors.append(
                        f'Row {i + 1}: date, file and amount are required.')
                    continue

                if len(row_description) > 100:
                    errors.append(
                        f'Row {i + 1}: description must be 100 characters or fewer.')
                    continue

                try:
                    transfer_date = datetime.strptime(
                        row_date, '%Y-%m-%d').date()
                except ValueError:
                    errors.append(f'Row {i + 1}: date is invalid.')
                    continue

                try:
                    amount = Decimal(row_amount)
                except Exception:
                    errors.append(f'Row {i + 1}: amount is invalid.')
                    continue

                if amount <= 0:
                    errors.append(
                        f'Row {i + 1}: amount must be greater than zero.')
                    continue

                file_obj = WIP.objects.filter(id=row_file_id).first()
                if not file_obj:
                    errors.append(f'Row {i + 1}: selected file was not found.')
                    continue

                LedgerAccountTransfers.objects.create(
                    file_number_from=file_obj,
                    file_number_to=file_obj,
                    from_ledger_account='C',
                    to_ledger_account='O',
                    amount=amount,
                    date=transfer_date,
                    description=row_description or 'C-O TFR',
                    amount_invoiced_from={},
                    balance_left_from=amount,
                    amount_invoiced_to={},
                    balance_left_to=amount,
                    is_cashier_co_transfer=True,
                    created_by=request.user
                )
                created_count += 1

            if created_count > 0:
                messages.success(
                    request, f'Added {created_count} client to office transfer row(s).')
            if errors:
                messages.error(
                    request, 'Some rows were not added: ' + ' | '.join(errors[:5]))
            if created_count == 0 and not errors:
                messages.error(request, 'No rows were provided to add.')
            return redirect('download_cashier_data')

        if cashier_action == 'mark_group_done':
            group_date_str = request.POST.get('group_date', '').strip()
            done_date_str = request.POST.get(
                'bank_transfer_done_on', '').strip()

            if not group_date_str:
                messages.error(request, 'Group date is required.')
                return redirect('download_cashier_data')

            if not done_date_str:
                messages.error(request, 'Bank done date is required.')
                return redirect('download_cashier_data')

            try:
                group_date = datetime.strptime(
                    group_date_str, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Group date is invalid.')
                return redirect('download_cashier_data')

            try:
                done_date = datetime.strptime(done_date_str, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Bank done date is invalid.')
                return redirect('download_cashier_data')

            pending_qs = LedgerAccountTransfers.objects.filter(
                is_cashier_co_transfer=True,
                from_ledger_account='C',
                to_ledger_account='O',
                file_number_from=F('file_number_to'),
                date=group_date,
                is_bank_transfer_done=False
            )
            updated_count = pending_qs.update(
                is_bank_transfer_done=True,
                bank_transfer_done_on=done_date,
                bank_transfer_done_by=request.user
            )

            if updated_count == 0:
                messages.error(
                    request, 'No pending transfers found in this date group.')
            else:
                messages.success(
                    request, f'Marked {updated_count} transfer(s) as done for {group_date.strftime("%d/%m/%Y")}.')
            return redirect('download_cashier_data')

        if cashier_action == 'export_office_to_client_transfers':
            start_date_str = request.POST.get('start_date', '').strip()
            end_date_str = request.POST.get('end_date', '').strip()

            if not start_date_str or not end_date_str:
                messages.error(
                    request, 'Start date and end date are required for export.')
                return redirect('download_cashier_data')

            try:
                start_date = datetime.strptime(
                    start_date_str, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Export date range is invalid.')
                return redirect('download_cashier_data')

            if end_date < start_date:
                messages.error(
                    request, 'End date cannot be before start date.')
                return redirect('download_cashier_data')

            transfers = LedgerAccountTransfers.objects.filter(
                is_cashier_co_transfer=True,
                from_ledger_account='O',
                to_ledger_account='C',
                file_number_from=F('file_number_to'),
                date__range=(start_date, end_date)
            ).select_related(
                'file_number_from',
                'created_by',
                'bank_transfer_done_by'
            ).order_by('date', 'file_number_from__file_number', 'id')

            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = (
                f'attachment; filename="OC_TFR_{start_date.strftime("%Y%m%d")}_{end_date.strftime("%Y%m%d")}.csv"'
            )

            writer = csv.writer(response)
            writer.writerow([
                'Date', 'File', 'Amount', 'Description', 'Created By',
                'Bank Status', 'Bank Done On', 'Bank Done By'
            ])

            if not transfers.exists():
                writer.writerow(
                    ['No O-C TFR rows found in selected date range.'])
                return response

            current_date = None
            group_total = Decimal('0.00')
            grand_total = Decimal('0.00')
            for transfer in transfers:
                if current_date is not None and current_date != transfer.date:
                    writer.writerow([
                        '', f'SUBTOTAL {current_date.strftime("%d/%m/%Y")}',
                        f'{group_total:.2f}', '', '', '', '', ''
                    ])
                    writer.writerow([])
                    group_total = Decimal('0.00')

                current_date = transfer.date
                group_total += transfer.amount
                grand_total += transfer.amount

                writer.writerow([
                    transfer.date.strftime('%d/%m/%Y'),
                    transfer.file_number_from.file_number if transfer.file_number_from else '-',
                    f'{transfer.amount:.2f}',
                    transfer.description,
                    str(transfer.created_by) if transfer.created_by else '-',
                    'Done' if transfer.is_bank_transfer_done else 'Pending',
                    transfer.bank_transfer_done_on.strftime(
                        '%d/%m/%Y') if transfer.bank_transfer_done_on else '-',
                    str(transfer.bank_transfer_done_by) if transfer.bank_transfer_done_by else '-'
                ])

            if current_date is not None:
                writer.writerow([
                    '', f'SUBTOTAL {current_date.strftime("%d/%m/%Y")}',
                    f'{group_total:.2f}', '', '', '', '', ''
                ])
            writer.writerow([])
            writer.writerow(
                ['', 'GRAND TOTAL', f'{grand_total:.2f}', '', '', '', '', ''])
            return response

        if cashier_action == 'add_office_to_client_transfers':
            row_dates = request.POST.getlist('oc_row_date[]')
            row_file_ids = request.POST.getlist('oc_row_file_id[]')
            row_amounts = request.POST.getlist('oc_row_amount[]')
            row_descriptions = request.POST.getlist('oc_row_description[]')

            max_rows = max(
                len(row_dates),
                len(row_file_ids),
                len(row_amounts),
                len(row_descriptions),
                0
            )

            created_count = 0
            errors = []
            for i in range(max_rows):
                row_date = row_dates[i].strip() if i < len(row_dates) else ''
                row_file_id = row_file_ids[i].strip(
                ) if i < len(row_file_ids) else ''
                row_amount = row_amounts[i].strip(
                ) if i < len(row_amounts) else ''
                row_description = row_descriptions[i].strip(
                ) if i < len(row_descriptions) else ''

                has_core_values = any([row_file_id, row_amount])
                has_custom_description = row_description not in ['', 'O-C TFR']
                if not has_core_values and not has_custom_description:
                    continue

                if not row_date or not row_file_id or not row_amount:
                    errors.append(
                        f'Row {i + 1}: date, file and amount are required.')
                    continue

                if len(row_description) > 100:
                    errors.append(
                        f'Row {i + 1}: description must be 100 characters or fewer.')
                    continue

                try:
                    transfer_date = datetime.strptime(
                        row_date, '%Y-%m-%d').date()
                except ValueError:
                    errors.append(f'Row {i + 1}: date is invalid.')
                    continue

                try:
                    amount = Decimal(row_amount)
                except Exception:
                    errors.append(f'Row {i + 1}: amount is invalid.')
                    continue

                if amount <= 0:
                    errors.append(
                        f'Row {i + 1}: amount must be greater than zero.')
                    continue

                file_obj = WIP.objects.filter(id=row_file_id).first()
                if not file_obj:
                    errors.append(f'Row {i + 1}: selected file was not found.')
                    continue

                LedgerAccountTransfers.objects.create(
                    file_number_from=file_obj,
                    file_number_to=file_obj,
                    from_ledger_account='O',
                    to_ledger_account='C',
                    amount=amount,
                    date=transfer_date,
                    description=row_description or 'O-C TFR',
                    amount_invoiced_from={},
                    balance_left_from=amount,
                    amount_invoiced_to={},
                    balance_left_to=amount,
                    is_cashier_co_transfer=True,
                    created_by=request.user
                )
                created_count += 1

            if created_count > 0:
                messages.success(
                    request, f'Added {created_count} office to client transfer row(s).')
            if errors:
                messages.error(
                    request, 'Some rows were not added: ' + ' | '.join(errors[:5]))
            if created_count == 0 and not errors:
                messages.error(request, 'No rows were provided to add.')
            return redirect('download_cashier_data')

        if cashier_action == 'mark_oc_group_done':
            group_date_str = request.POST.get('group_date', '').strip()
            done_date_str = request.POST.get(
                'bank_transfer_done_on', '').strip()

            if not group_date_str:
                messages.error(request, 'Group date is required.')
                return redirect('download_cashier_data')

            if not done_date_str:
                messages.error(request, 'Bank done date is required.')
                return redirect('download_cashier_data')

            try:
                group_date = datetime.strptime(
                    group_date_str, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Group date is invalid.')
                return redirect('download_cashier_data')

            try:
                done_date = datetime.strptime(done_date_str, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Bank done date is invalid.')
                return redirect('download_cashier_data')

            pending_qs = LedgerAccountTransfers.objects.filter(
                is_cashier_co_transfer=True,
                from_ledger_account='O',
                to_ledger_account='C',
                file_number_from=F('file_number_to'),
                date=group_date,
                is_bank_transfer_done=False
            )
            updated_count = pending_qs.update(
                is_bank_transfer_done=True,
                bank_transfer_done_on=done_date,
                bank_transfer_done_by=request.user
            )

            if updated_count == 0:
                messages.error(
                    request, 'No pending transfers found in this date group.')
            else:
                messages.success(
                    request, f'Marked {updated_count} transfer(s) as done for {group_date.strftime("%d/%m/%Y")}.')
            return redirect('download_cashier_data')

        start_date_str = request.POST['start_date']

        start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
        start_datetime = datetime.combine(start_date.date(), time.min)
        get_pending_slips = request.POST['type_of_report'] == 'Pending Slips'

        if get_pending_slips:
            end_datetime = datetime.combine(datetime.today().date(), time.max)
            title = "Pending Slips"
            check_headings = ""
            check_body = ""
        else:
            end_date_str = request.POST['end_date']
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d')
            end_datetime = datetime.combine(end_date.date(), time.max)

            title = "Audit Slips"
            check_headings = """
                            <th>Checked Millenium</th>
                            <th>Checked Bank A/c</th>
                        """
            check_body = """
                            <td></td>
                            <td></td>
                        """
        start_date_only = start_datetime.date()
        end_date_only = end_datetime.date()
        invoices = Invoices.objects.filter(
            Q(date__range=(start_date_only, end_date_only)) & Q(state='F')).order_by('invoice_number')
        credit_notes = CreditNote.objects.filter(
            date__range=(start_date_only, end_date_only)
        ).select_related('invoice', 'file_number', 'created_by', 'approved_by').order_by('date')
        slips = PmtsSlips.objects.filter(
            timestamp__range=(start_datetime, end_datetime)).order_by('date')
        green_slips = LedgerAccountTransfers.objects.filter(
            timestamp__range=(start_datetime, end_datetime)).order_by('date')

        invoice_display_table = ''
        if get_pending_slips:
            invoice_display_table = invoice_display_table + """<h2 class='mt-3'>Invoices</h2><table class='table table-striped'>
                <thead>
                    <th>
                        File Number
                    </th>
                    <th>
                        Invoice Number
                    </th>
                    <th>
                        Date
                    </th>
                    <th>
                        Our Costs
                    </th>
                    <th>
                        VAT
                    </th>
                    <th>
                        Total Costs and VAT
                    </th>
                    <th>
                        Disbursements
                    </th>
                </thead>
            """
            for invoice in invoices:
                total_cost_invoice, vat, total_cost_and_vat = calculate_invoice_total_with_vat(
                    invoice)
                slip_display = ''
                for slip in invoice.disbs_ids.all():
                    slip_display = slip_display + \
                        f"({slip.date.strftime('%d/%m/%Y')}, £{slip.amount})"

                invoice_display_table = invoice_display_table + f"""
                <tr>
                    <td>{invoice.file_number.file_number}</td>
                    <td>{invoice.invoice_number}</td>
                    <td>{invoice.date.strftime('%d/%m/%Y')}</td>
                    <td>£{total_cost_invoice}</td>
                    <td>£{vat}</td>
                    <td>£{total_cost_and_vat}</td>
                    <td>{slip_display}</td>
                </tr>
                """
            invoice_display_table = invoice_display_table + "</table>"

        status_labels = dict(CreditNote.STATUSES)
        credit_notes_table = """
            <table class='table table-striped'>
                <thead>
                    <th>Date</th>
                    <th>File Number</th>
                    <th>Credit Amount Excl. VAT</th>
                    <th>VAT</th>
                    <th>Total Credit Incl. VAT</th>
                    <th>Status</th>
                    <th>Reason</th>
                    <th>Created By</th>
                    <th>Approved By</th>
                    <th>Approved On</th>
                </thead>
                <tbody>
        """
        for credit_note in credit_notes:
            approved_on = credit_note.approved_on.strftime(
                '%d/%m/%Y %H:%M') if credit_note.approved_on else '-'
            net_amount, vat_amount, gross_amount = calculate_credit_note_breakdown(
                credit_note.amount)
            credit_notes_table = credit_notes_table + f"""
            <tr>
                <td>{credit_note.date.strftime('%d/%m/%Y')}</td>
                <td>{credit_note.file_number.file_number}</td>
                <td>£{net_amount}</td>
                <td>£{vat_amount}</td>
                <td>£{gross_amount}</td>
                <td>{status_labels.get(credit_note.status, credit_note.status)}</td>
                <td>{credit_note.reason}</td>
                <td>{credit_note.created_by or '-'}</td>
                <td>{credit_note.approved_by or '-'}</td>
                <td>{approved_on}</td>
            </tr>
            """
        credit_notes_table = credit_notes_table + "</tbody></table>"

        common_slips_header = f"""
                <thead>
                    <th>
                        Date
                    </th>
                    <th>
                        File Number
                    </th>
                    <th>
                        Mode of Payment
                    </th>
                    <th>
                        Ledger Account
                    </th>
                    <th>
                        Amount
                    </th>
                    <th>
                        Payment to or from
                    </th>
                    <th>
                        Description
                    </th>
                    {check_headings}

                </thead>
            """

        client_blue_slips_table = f"""<table class='table table-striped'>{
            common_slips_header}"""
        client_pink_slips_table = f"""<table class='table table-striped'>{
            common_slips_header}"""

        office_blue_slips_table = f"""<table class='table table-striped'>{
            common_slips_header}"""
        office_pink_slips_table = f"""<table class='table table-striped'>{
            common_slips_header}"""

        for slip in slips:

            if (slip.is_money_out == False and slip.ledger_account == 'C'):
                client_blue_slips_table = client_blue_slips_table + f"""
                <tr>
                    <td>{slip.date.strftime('%d/%m/%Y')}</td>
                    <td>{slip.file_number}</td>
                    <td>{slip.mode_of_pmt}</td>
                    <td>Client</td>
                    <td>£{slip.amount}</td>
                    <td>{slip.pmt_person}</td>
                    <td>{slip.description}</td>
                    {check_body}
                </tr>
                """
            elif (slip.is_money_out == False and slip.ledger_account == 'O'):
                office_blue_slips_table = office_blue_slips_table + f"""
                <tr>
                    <td>{slip.date.strftime('%d/%m/%Y')}</td>
                    <td>{slip.file_number}</td>
                    <td>{slip.mode_of_pmt}</td>
                    <td>Office</td>
                    <td>£{slip.amount}</td>
                    <td>{slip.pmt_person}</td>
                    <td>{slip.description}</td>
                    {check_body}
                </tr>
                """
            elif (slip.is_money_out == True and slip.ledger_account == 'C'):
                client_pink_slips_table = client_pink_slips_table + f"""
                <tr>
                    <td>{slip.date.strftime('%d/%m/%Y')}</td>
                    <td>{slip.file_number}</td>
                    <td>{slip.mode_of_pmt}</td>
                    <td>Client</td>
                    <td>£{slip.amount}</td>
                    <td>{slip.pmt_person}</td>
                    <td>{slip.description}</td>
                    {check_body}
                </tr>
                """
            else:
                office_pink_slips_table = office_pink_slips_table + f"""
                <tr>
                    <td>{slip.date.strftime('%d/%m/%Y')}</td>
                    <td>{slip.file_number}</td>
                    <td>{slip.mode_of_pmt}</td>
                    <td>Office</td>
                    <td>£{slip.amount}</td>
                    <td>{slip.pmt_person}</td>
                    <td>{slip.description}</td>
                    {check_body}
                </tr>
                """

        client_blue_slips_table = client_blue_slips_table + "</table>"
        client_pink_slips_table = client_pink_slips_table + "</table>"
        office_blue_slips_table = office_blue_slips_table + "</table>"
        office_pink_slips_table = office_pink_slips_table + "</table>"

        green_slips_table = f"""
            <table class='table table-striped'>
                <thead>
                    <th>
                        Date
                    </th>
                    <th>
                        File Number From
                    </th>
                    <th>
                        File Number To
                    </th>
                    <th>
                        Ledger Account From
                    </th>
                    <th>
                        Ledget Account To
                    </th>
                    <th>
                        Description
                    </th>
                    <th>
                        Amount
                    </th>
                </thead>
                <tbody>
        """

        for slip in green_slips:
            green_slips_table = green_slips_table + f"""
            <tr>
                <td>{slip.date.strftime("%d/%m/%Y")}</td>
                <td>{slip.file_number_from.file_number if slip.file_number_from else '-'}</td>
                <td>{slip.file_number_to.file_number if slip.file_number_to else '-'}</td>
                <td>{slip.from_ledger_account}</td>
                <td>{slip.to_ledger_account}</td>
                <td>{slip.description}</td>
                <td>£{slip.amount}</td>
            </tr>
            """
        green_slips_table = green_slips_table + f"""</tbody></table>"""
        page_style = '@page { size: landscape; }'
        html = f"""
                <html>
                    <head>
                        <style>{page_style}</style>
                        <meta charset="UTF-8">
                        <meta name="viewport" content="width=device-width, initial-scale=1.0">
                        <title>{title}</title>
                        <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-T3c6CoIi6uLrA9TneNEoa7RxnatzjcDSCmG1MXxSR1GAsXEV/Dwwykc2MPK8M2HN" crossorigin="anonymous"/>

                    </head>
                    <body style="font-family: Times New Roman, Times, serif">
                        <h1 class="text-center">{title}</h1>
                        {invoice_display_table}
                        <h2 class='mt-3'>Credit Notes</h2>
                        {credit_notes_table}
                        <h2 class='mt-3'>Client Blue Slips</h2>
                        {client_blue_slips_table}
                        <h2 class='mt-3'>Client Pink Slips</h2>
                        {client_pink_slips_table}
                        <h2 class='mt-3'>Office Blue Slips</h2>
                        {office_blue_slips_table}
                        <h2 class='mt-3'>Office Pink Slips</h2>
                        {office_pink_slips_table}
                        <h2 class='mt-3'>Green Slips </h2>
                        {green_slips_table}
                        <div>
                            <br>
                            <p><b>Date:</b> {datetime.today().date().strftime('%d/%m/%Y')}</p>
                            <br>
                            <p><b>Signed:</b>...........................................</p>

                        </div>
                    </body>
                </html>
                """
        pdf_file = HTML(string=html).write_pdf()

        return HttpResponse(pdf_file, content_type='application/pdf')

    context = get_transfers_context()
    return render(request, 'cashier_data.html', context)


@login_required
def download_file_logs(request, file_number):
    logs, log_meta = enrich_file_logs(get_file_logs(file_number))
    downloaded_at = timezone.localtime(
        timezone.now()).strftime('%d/%m/%Y %H:%M')
    context = {
        'file_number': file_number,
        'logs': logs,
        'log_meta': log_meta,
        'downloaded_by': request.user.get_full_name() or request.user.username,
        'downloaded_at': downloaded_at,
    }
    html = render_to_string(
        'download_templates/file_logs.html',
        context,
        request=request,
    )
    pdf_file = HTML(
        string=html, base_url=request.build_absolute_uri('/')).write_pdf()
    response = HttpResponse(pdf_file, content_type='application/pdf')
    filename_ts = timezone.now().strftime('%Y%m%d_%H%M')
    filename = f'activity_log_{file_number}_{request.user.username}_{filename_ts}.pdf'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def download_frontsheet(request, file_number):
    file = WIP.objects.get(file_number=file_number)
    page_style = '@page { margin-top: 24pt; margin-bottom:0; font-size:7pt !important; size: A4;}'
    title = f"Frontsheet - {file_number}"

    def format_date(value):
        return value.strftime('%d/%m/%Y') if value else ''

    def client_check_label(client):
        return 'UK Business Check' if client and client.is_business else 'AML Check'

    client1_check_label = client_check_label(file.client1)
    if file.client2:
        client2_name = file.client2.name
        client2_address = f'{file.client2.address_line1},{
            file.client2.address_line2},<br>{file.client2.county}, {file.client2.postcode}'
        client2_contact_number = file.client2.contact_number
        client2_email = file.client2.email
        client2_dob = format_date(file.client2.dob)
        client2_id_verified = 'Yes' if file.client2.id_verified else 'No'
        client2_date_of_last_aml = format_date(file.client2.date_of_last_aml)
        client2_terms_signed = 'Yes' if file.client2.terms_of_engagement_signed else 'No'
        client2_ncba_signed = 'Yes' if file.client2.ncba_signed else 'No'
        client2_pep_signed = 'Yes' if file.client2.pep_signed else 'No'
        client2_sof_signed = 'Yes' if file.client2.source_of_funds_signed else 'No'
        client2_check_label = client_check_label(file.client2)
    else:
        client2_name = ''
        client2_address = ''
        client2_contact_number = ''
        client2_email = ''
        client2_dob = ''
        client2_id_verified = ''
        client2_date_of_last_aml = ''
        client2_terms_signed = ''
        client2_ncba_signed = ''
        client2_pep_signed = ''
        client2_sof_signed = ''
        client2_check_label = ''

    if file.authorised_party1:
        ap1_name = file.authorised_party1.name
        ap1_addr = f'{file.authorised_party1.address_line1}, {file.authorised_party1.address_line2},<br>{
            file.authorised_party1.county}, {file.authorised_party1.postcode}'
        ap1_email = file.authorised_party1.email
        ap1_contact_number = file.authorised_party1.contact_number
        ap1_date_id_check = format_date(
            file.authorised_party1.date_of_id_check)
        ap1_date_aml_check = format_date(
            file.authorised_party1.date_of_last_aml)
        ap1_relationship = file.authorised_party1.relationship_to_client
    else:
        ap1_name = ''
        ap1_addr = ''
        ap1_email = ''
        ap1_contact_number = ''
        ap1_date_id_check = ''
        ap1_date_aml_check = ''
        ap1_relationship = ''

    if file.authorised_party2:
        ap2_name = file.authorised_party2.name
        ap2_addr = f'{file.authorised_party2.address_line1}, {file.authorised_party2.address_line2},<br>{
            file.authorised_party2.county}, {file.authorised_party2.postcode}'
        ap2_email = file.authorised_party2.email
        ap2_contact_number = file.authorised_party2.contact_number
        ap2_date_id_check = format_date(
            file.authorised_party2.date_of_id_check)
        ap2_date_aml_check = format_date(
            file.authorised_party2.date_of_last_aml)
        ap2_relationship = file.authorised_party2.relationship_to_client
    else:
        ap2_name = ''
        ap2_addr = ''
        ap2_email = ''
        ap2_contact_number = ''
        ap2_date_id_check = ''
        ap2_date_aml_check = ''
        ap2_relationship = ''

    if file.other_side:
        other_side_name = file.other_side.name
        other_side_address = f'{file.other_side.name}, {file.other_side.address_line1},<br>{
            file.other_side.address_line2}, {file.other_side.postcode}'
        other_side_mobile = file.other_side.contact_number
        other_side_email = file.other_side.email
        other_side_solicitors = file.other_side.solicitors
        other_side_solicitors_email = file.other_side.solicitors_email
    else:
        other_side_name = ''
        other_side_address = ''
        other_side_mobile = ''
        other_side_email = ''
        other_side_solicitors = ''
        other_side_solicitors_email = ''

    html = f"""
    <html>
        <head>
            <style>{page_style}</style>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=0.8">
            <title>{title}</title>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-T3c6CoIi6uLrA9TneNEoa7RxnatzjcDSCmG1MXxSR1GAsXEV/Dwwykc2MPK8M2HN" crossorigin="anonymous"/>

        </head>
        <body style="font-family: Times New Roman, Times, serif; font-size:9pt !important;" class='vh-100'>
            <table class="table table-bordered table-sm" border="1" style="margin-left:50pt;  border-color: black;" cellpadding="2" cellspacing="0" >
                <tbody>
                <tr>
                    <td style="font-weight: 900; font-size:12px;" ><h3>FILE NUMBER</h3></td>
                    <td style="font-weight: 900 !important; font-size:16px; !important; text-align:center !important;" colspan='2'>{file.file_number}</td>

                </tr>
                <tr>
                    <td>Z DRIVE LOCATION</td>
                    <td class='text-center' colspan='2'>{file.zdrive_location or ''}</td>
                </tr>
                <tr>
                    <td style="background-color:grey;"  colspan='3'></td>
                </tr>
                <tr>
                    <td class='' colspan='3'><b>CLIENT DETAILS</b></td>
                </tr>
                <tr>
                    <td></td>
                    <td class='text-center'><b>CLIENT 1</b></td>
                    <td class='text-center'><b>CLIENT 2</b></td>
                </tr>
                <tr>
                    <td>NAME</td>
                    <td class='text-center'>{file.client1.name}</td>
                    <td class='text-center'>{client2_name}</td>
                </tr>
                <tr>
                    <td>ADDRESS</td>
                    <td class='text-center'>{file.client1.address_line1}, {file.client1.address_line2},<br>{file.client1.county}, {file.client1.postcode}</td>
                    <td class='text-center'>{client2_address}</td>
                </tr>
                <tr>
                    <td>CONTACT NO.</td>
                    <td class='text-center'>{file.client1.contact_number}</td>
                    <td class='text-center'>{client2_contact_number}</td>
                </tr>
                <tr>
                    <td>EMAIL</td>
                    <td class='text-center'>{file.client1.email}</td>
                    <td class='text-center'>{client2_email}</td>
                </tr>
                <tr >
                    <td>DATE OF BIRTH</td>
                    <td class='text-center'>{format_date(file.client1.dob)}</td>
                    <td class='text-center'>{client2_dob}</td>
                </tr>
                <tr>
                    <td>CHECK TYPE</td>
                    <td class='text-center'>{client1_check_label}</td>
                    <td class='text-center'>{client2_check_label}</td>
                </tr>
                <tr >
                    <td>DATE OF LAST CHECK</td>
                    <td class='text-center'>{format_date(file.client1.date_of_last_aml)}</td>
                    <td class='text-center'>{client2_date_of_last_aml}</td>
                </tr>
                <tr>
                    <td>ID VERIFIED</td>
                    <td class='text-center'>{'Yes' if file.client1.id_verified else 'No'}</td>
                    <td class='text-center'>{client2_id_verified}</td>
                </tr>
                <tr>
                    <td>SIGNED TERMS OF ENGAGEMENT</td>
                    <td class='text-center'>{'Yes' if file.client1.terms_of_engagement_signed else 'No'}</td>
                    <td class='text-center'>{client2_terms_signed}</td>
                </tr>
                <tr>
                    <td>SIGNED NCBA</td>
                    <td class='text-center'>{'Yes' if file.client1.ncba_signed else 'No'}</td>
                    <td class='text-center'>{client2_ncba_signed}</td>
                </tr>
                <tr>
                    <td>SIGNED PEP</td>
                    <td class='text-center'>{'Yes' if file.client1.pep_signed else 'No'}</td>
                    <td class='text-center'>{client2_pep_signed}</td>
                </tr>
                <tr>
                    <td>SIGNED SOF</td>
                    <td class='text-center'>{'Yes' if file.client1.source_of_funds_signed else 'No'}</td>
                    <td class='text-center'>{client2_sof_signed}</td>
                </tr>
                <tr>
                    <td style="background-color:grey;"  colspan='3'></td>
                </tr>
                <tr>
                    <td>CLIENT CARE LETTER SENT</td>
                    <td class='text-center' colspan='2'>{file.date_of_client_care_sent.strftime('%d/%m/%Y') if file.date_of_client_care_sent else ''}</td>
                </tr>
                <tr>
                    <td>FUNDING</td>
                    <td class='text-center' colspan='2'>{'Private Funding' if file.funding == 'PF' else file.funding}</td>
                </tr>
                <tr>
                    <td style="background-color:grey;"  colspan='3'></td>
                </tr>
                <tr>
                    <td class='' colspan='3'>AUTHORISED PARTIES</td>
                </tr>
                <tr>
                    <td></td>
                    <td class='text-center'><b>AUTHORISED PARTY 1</b></td>
                    <td class='text-center'><b>AUTHORISED PARTY 2</b></td>
                </tr>
                <tr>
                    <td>NAME</td>
                    <td class='text-center'>{ap1_name}</td>
                    <td class='text-center'>{ap2_name}</td>
                </tr>
                <tr>
                    <td>RELATIONSHIP</td>
                    <td class='text-center'>{ap1_relationship}</td>
                    <td class='text-center'>{ap2_relationship}</td>
                </tr>
                <tr>
                    <td>EMAIL</td>
                    <td class='text-center'>{ap1_email}</td>
                    <td class='text-center'>{ap2_email}</td>
                </tr>
                <tr>
                    <td>ADDRESS</td>
                    <td class='text-center'>{ap1_addr}</td>
                    <td class='text-center'>{ap2_addr}</td>
                </tr>
                <tr>
                    <td>CONTACT NUMBER</td>
                    <td class='text-center'>{ap1_contact_number}</td>
                    <td class='text-center'>{ap2_contact_number}</td>
                </tr>
                <tr>
                    <td>DATE OF ID CHECK</td>
                    <td class='text-center'>{ap1_date_id_check}</td>
                    <td class='text-center'>{ap2_date_id_check}</td>
                </tr>
                <tr>
                    <td>DATE OF LAST AML CHECK</td>
                    <td class='text-center'>{ap1_date_aml_check}</td>
                    <td class='text-center'>{ap2_date_aml_check}</td>
                </tr>
                <tr>
                    <td style="background-color:grey;"  colspan='3'></td>
                </tr>
                <tr>
                    <td colspan='3'><b>OTHER SIDE'S DETAILS</b></td>
                </tr>
                <tr>
                    <td>NAME</td>
                    <td class='text-center' colspan='2'>{other_side_name}</td>
                </tr>
                <tr>
                    <td>ADDRESS</td>
                    <td class='text-center' colspan='2'>{other_side_address}</td>
                </tr>
                <tr>
                    <td>MOBILE</td>
                    <td class='text-center' colspan='2'>{other_side_mobile}</td>
                </tr>
                <tr>
                    <td>EMAIL</td>
                    <td class='text-center' colspan='2'>{other_side_email}</td>
                </tr>
                <tr>
                    <td>SOLICITORS</td>
                    <td class='text-center' colspan='2'>{other_side_solicitors}</td>
                </tr>
                <tr>
                    <td>SOLICITORS - EMAIL</td>
                    <td class='text-center' colspan='2'>{other_side_solicitors_email}</td>
                </tr>
                <tr>
                    <td style="background-color:grey;"  colspan='3'></td>
                </tr>
                <tr>
                    <td class='' colspan='3'><b>KEY INFORMATION</b></td>
                </tr>
                <tr>
                    <td class='' colspan='3'>{file.key_information}</td>
                </tr>

                </tbody>
            </table>
        </body>
    </html>
     """

    pdf_file = HTML(string=html).write_pdf()
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="frontsheet_{
        file_number}_{request.user}_{datetime.now().strftime('%d/%m/%Y %I:%M %p.pdf"')}'
    return response


@login_required
def download_risk_assessment(request, id):
    risk_assessment = get_object_or_404(RiskAssessment, pk=id)
    html_string = render_to_string(
        'download_templates/risk_assessment.html', {"obj": risk_assessment})

    pdf_file = HTML(string=html_string).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="risk_assessment_{risk_assessment.matter.file_number}_{id}.pdf"'
    return response


@login_required
def add_ongoing_monitoring(request, file_number):
    try:
        matter = WIP.objects.get(file_number=file_number)
    except WIP.DoesNotExist:
        messages.error(
            request, 'Matter with the given file number does not exist.')
        return redirect('index')
    if request.method == 'POST':
        post_data = request.POST.copy()
        post_data['created_by'] = request.user
        post_data['file_number'] = matter.id
        form = OngoingMonitoringForm(post_data)

        if form.is_valid():
            ongoing_monitoring = form.save()
            log_created(
                request.user,
                ongoing_monitoring,
                f'Ongoing monitoring for {ongoing_monitoring.file_number.file_number}',
            )
            messages.success(
                request, 'Ongoing Monitoring successfully recorded.')
            return redirect('home', ongoing_monitoring.file_number.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = OngoingMonitoringForm()
        return render(request, 'ongoing_monitoring.html', {'form': form, 'file_number': file_number, 'title': 'Add'})


@login_required
def policies_display(request):
    latest_version_subquery = PolicyVersion.objects.filter(
        policy=OuterRef('pk')
    ).order_by('-version_number').values('pk')[:1]

    policies = Policy.objects.annotate(
        latest_version_id=Subquery(latest_version_subquery),
        is_read=Exists(
            PoliciesRead.objects.filter(
                policy=OuterRef('pk'),
                policy_version_id=OuterRef('latest_version_id'),
                read_by=request.user
            )
        )
    ).order_by('description')

    policies_read = PoliciesRead.objects.filter(
        read_by=request.user).order_by('-timestamp')
    any_unread = policies.filter(is_read=False).exists()
    context = {
        'policies': policies,
        'policies_read': policies_read,
        'all_read': not any_unread

    }

    return render(request, 'policies/policies_home.html', context)


@login_required
def policy_read(request, policy_id):
    policy = get_object_or_404(Policy, pk=policy_id)

    latest_version = policy.versions.order_by('-version_number').first()

    if not latest_version:
        messages.error(
            request, f"No versions available for policy '{policy.description}'.")
        return redirect('policies_display')

    try:
        PoliciesRead.objects.create(
            policy=policy, policy_version=latest_version, read_by=request.user)
        messages.success(
            request, f'Successfully marked the latest version of "{policy.description}" as read.')
    except Exception as e:
        messages.error(request, f"An error occurred: {str(e)}")

    return redirect('policies_display')


@login_required
def add_policy(request):
    if not request.user.is_manager:
        messages.error(request, "You do not have permission to add policies.")
        return redirect('policies_display')

    if request.method == 'POST':
        form = PolicyForm(request.POST)
        if form.is_valid():
            try:
                policy = form.save()
                PolicyVersion.objects.create(
                    policy=policy,
                    content=form.cleaned_data['content'],
                    version_number=1,
                    changes_by=request.user,
                    timestamp=timezone.now()
                )
                messages.success(request, "Policy added successfully.")
                return redirect('policies_display')
            except Exception as e:
                messages.error(
                    request, f"An error occurred while adding the policy: {str(e)}")
        else:
            messages.error(
                request, "There were errors in the form. Please correct them and try again.")
    else:
        form = PolicyForm()

    return render(request, 'policies/add_policy.html', {'form': form})


@login_required
def edit_policy(request, policy_id):
    policy = get_object_or_404(Policy, id=policy_id)

    if not request.user.is_manager:
        messages.error(request, "You do not have permission to edit policies.")
        return redirect('policies_display')

    if request.method == 'POST':
        form = PolicyForm(request.POST, instance=policy)
        if form.is_valid():
            try:

                new_content = form.cleaned_data['content']
                latest_version = policy.versions.order_by(
                    '-version_number').first()

                if latest_version is None or latest_version.content != new_content:
                    # Determine new version number
                    version_number = latest_version.version_number + 1 if latest_version else 1

                    # Create a new PolicyVersion with the updated content
                    PolicyVersion.objects.create(
                        policy=policy,
                        content=new_content,
                        version_number=version_number,
                        changes_by=request.user,
                        timestamp=timezone.now()
                    )

                form.save()
                messages.success(request, "Policy edited successfully.")
                return redirect('policies_display')
            except Exception as e:
                messages.error(
                    request, f"An error occurred while editing the policy: {str(e)}")
        else:
            messages.error(
                request, "There were errors in the form. Please correct them and try again.")
    else:
        form = PolicyForm(instance=policy)

    return render(request, 'policies/edit_policy.html', {'form': form, 'policy': policy})


@login_required
def download_policy_pdf(request, policy_version_id):
    #
    policy_version = get_object_or_404(PolicyVersion, id=policy_version_id)

    try:
        html_string = render_to_string('download_templates/policy_pdf.html', {
            'policy_version': policy_version
        })

        # Create the PDF using WeasyPrint
        html = HTML(string=html_string)
        pdf = html.write_pdf()

        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Policy_{policy_version.policy.id}_Version_{policy_version.version_number}.pdf"'

        return response

    except Exception as e:
        print(e)
        messages.error(
            request, f"An error occurred while generating the PDF: {str(e)}")
        return redirect('policies_display')


@login_required
def invoices_list(request):
    # Get start and end dates from GET parameters
    start_date = request.GET.get('start_date')

    # Calculate current financial year if start and end are not provided
    current_date = timezone.now()
    if not start_date:
        if current_date.month >= 11:  # November to October financial year
            start_date = datetime(current_date.year, 11, 1)
            end_date = datetime(current_date.year + 1, 10, 31)
        else:
            start_date = datetime(current_date.year - 1, 11, 1)
            end_date = datetime(current_date.year, 10, 31)
    else:
        # Parse dates from GET parameters
        try:
            start_date = datetime.strptime(start_date, '%Y-%m')
            end_date = start_date.replace(
                year=start_date.year + 1, month=10, day=31)

        except ValueError:
            messages.error(
                request, 'Invalid date format. Please use YYYY-MM-DD.')
            return redirect('index')

    # Filter invoices by finalized state and date range
    invoices = Invoices.objects.filter(
        state='F', date__range=[start_date, end_date])

    if request.user.is_manager:
        approved_credit_totals = get_approved_credit_note_totals(
            [invoice.id for invoice in invoices]
        )
        for invoice in invoices:
            total_cost_invoice, vat_inv, total_cost_and_vat = calculate_invoice_total_with_vat(
                invoice)
            effective_due = get_effective_invoice_due(
                invoice, approved_credit_totals.get(invoice.id, Decimal('0')))

            # Update invoice attributes dynamically
            invoice.our_costs = total_cost_invoice
            invoice.vat = round(vat_inv, 2)
            invoice.total_cost_and_vat = total_cost_and_vat
            invoice.total_due_left = effective_due

        # Render template with invoices and date range
        return render(request, 'invoices_list.html', {
            'invoices': invoices,
            'start_date': start_date,
            'end_date': end_date,
        })
    else:
        messages.error(
            request, 'You do not have the right level of permissions.')
        return redirect('index')


@login_required
def download_invoices(request):

    if request.method == 'POST' and request.user.is_manager:

        start_date_str = request.POST['start']
        end_date_str = request.POST['end']
        start_date = datetime.strptime(start_date_str, '%d/%m/%Y')
        end_date = datetime.strptime(end_date_str, '%d/%m/%Y')

        invoices = Invoices.objects.filter(
            state='F', date__range=[start_date, end_date])

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="Invoices_from_{start_date_str}_to_{end_date_str}.csv"'

        writer = csv.writer(response)

        writer.writerow(
            ['', f'Invoices from {start_date_str} to {end_date_str}'])

        writer.writerow([])
        writer.writerow(['Invoice Number', 'Matter Type', 'Date',
                        'File Number', 'Our Costs', 'VAT', 'Total Costs and VAT'])

        for invoice in invoices:
            total_cost_invoice, vat_inv, total_cost_and_vat = calculate_invoice_total_with_vat(
                invoice)

            writer.writerow([f'{invoice.invoice_number}', f'{invoice.file_number.matter_type.type}', f'{invoice.date.strftime("%d/%m/%Y")}',
                            f'{invoice.file_number.file_number}', f'{total_cost_invoice}', f'{vat_inv}', f'{total_cost_and_vat}'])

        return response


@login_required
def edit_risk_assessment(request, id):
    try:
        risk_assesssment = get_object_or_404(RiskAssessment, pk=id)
    except Exception as e:
        messages.error(request, f'Risk Assessment not found. {str(e)}')
        return redirect('index')
    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(risk_assesssment)
        form = RiskAssessmentForm(request.POST, instance=risk_assesssment)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()

            for field in changed_fields:
                changes[field]['new_value'] = str(
                    getattr(risk_assesssment, field))

            if changes:
                create_modification(
                    user=request.user,
                    modified_obj=risk_assesssment,
                    changes=changes
                )
            messages.success(request, 'Successfully updated Risk Assessment.')
            return redirect('home', risk_assesssment.matter.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = RiskAssessmentForm(instance=risk_assesssment)

    return render(request, 'risk_assessment.html', {'form': form, 'file_number': risk_assesssment.matter.file_number, 'title': 'Edit'})


@login_required
def edit_ongoing_monitoring(request, id):
    try:
        ongoing_monitoring = get_object_or_404(OngoingMonitoring, pk=id)
    except Exception as e:
        messages.error(request, f'Ongoing Monitoring not found. {str(e)}')
        return redirect('index')
    if request.method == 'POST':
        duplicate_obj = copy.deepcopy(ongoing_monitoring)
        post_copy = request.POST.copy()
        post_copy['file_number'] = ongoing_monitoring.file_number.id
        post_copy['created_by'] = ongoing_monitoring.created_by
        form = OngoingMonitoringForm(post_copy, instance=ongoing_monitoring)
        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                changes[field] = {
                    'old_value': str(getattr(duplicate_obj, field)),
                    'new_value': None
                }
            form.save()

            for field in changed_fields:
                changes[field]['new_value'] = str(
                    getattr(ongoing_monitoring, field))

            if changes:
                create_modification(
                    user=request.user,
                    modified_obj=ongoing_monitoring,
                    changes=changes
                )
            messages.success(
                request, 'Successfully updated Ongoing Monitoring.')
            return redirect('home', ongoing_monitoring.file_number.file_number)
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = OngoingMonitoringForm(instance=ongoing_monitoring)

    return render(request, 'ongoing_monitoring.html', {'form': form, 'file_number': ongoing_monitoring.file_number.file_number, 'title': 'Edit'})


@login_required
def download_ongoing_monitoring(request, id):
    obj = get_object_or_404(OngoingMonitoring, pk=id)
    html_string = render_to_string(
        'download_templates/ongoing_monitoring.html', {"obj": obj})

    pdf_file = HTML(string=html_string).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="ongoing_monitoring_{obj.file_number.file_number}_{id}.pdf"'
    return response


@login_required
def onboarding_documents_display(request):
    if request.method == "POST":
        file = request.POST['file']
        html_string = render_to_string(file)

        pdf_file = HTML(string=html_string).write_pdf()

        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{file}_document.pdf"'
        return response
    else:
        return render(request, 'onboarding_documents.html')


@login_required
def download_document(request):

    file_name = 'pep_questionnaire.doc'
    file_dir = 'files'
    file_path = None

    # Check each directory in STATICFILES_DIRS
    for static_dir in settings.STATICFILES_DIRS:
        potential_path = os.path.join(static_dir, file_dir, file_name)
        if os.path.exists(potential_path):
            file_path = potential_path
            break

    if file_path and os.path.exists(file_path):
        response = FileResponse(open(
            file_path, 'rb'), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        return response
    else:
        raise Http404("File does not exist")


def free30mins(request):
    free30_mins_form = Free30MinsForm()
    free30_mins_attendees_form = Free30MinsAttendeesForm()

    if request.method == 'POST':
        try:
            number_of_attendees = int(request.POST['number_of_attendees'])
            attendee_ids = []

            for i in range(number_of_attendees):
                index = f'{i}_' if i > 0 else ''
                name = request.POST[f'{index}name']
                address_line1 = request.POST[f'{index}address_line1']
                address_line2 = request.POST[f'{index}address_line2']
                county = request.POST[f'{index}county']
                postcode = request.POST[f'{index}postcode']
                email = request.POST[f'{index}email']
                contact_number = request.POST[f'{index}contact_number']
                created_by = request.user

                attendee = Free30MinsAttendees.objects.create(
                    name=name,
                    address_line1=address_line1,
                    address_line2=address_line2,
                    county=county,
                    postcode=postcode,
                    email=email,
                    contact_number=contact_number,
                    created_by=created_by,
                )
                attendee.save()
                attendee_ids.append(attendee.id)

            matter_type_id = request.POST['matter_type']
            matter_type = MatterType.objects.filter(
                pk=matter_type_id).first() if matter_type_id != '' else None
            notes = request.POST['notes']

            date = request.POST['date']
            start_time = request.POST['start_time']
            finish_time = request.POST['finish_time']
            fee_earner_id = request.POST['fee_earner']
            fee_earner = CustomUser.objects.filter(pk=fee_earner_id).first()
            created_by = request.user

            free_30_mins = Free30Mins.objects.create(
                matter_type=matter_type,
                notes=notes,
                date=date,
                start_time=start_time,
                finish_time=finish_time,
                fee_earner=fee_earner,
                created_by=created_by
            )

            # Assign attendees to the meeting
            free_30_mins.attendees.set(attendee_ids)

            free_30_mins.save()
            messages.success(request, "Meeting successfully created.")
            return redirect('free30mins')

        except Exception as e:
            messages.error(request, f"An error occurred: {e}")

    search_query = request.GET.get('q', '').strip()
    allowed_page_sizes = [10, 25, 50, 100]
    try:
        page_size = int(request.GET.get('per_page', 25))
    except (TypeError, ValueError):
        page_size = 25
    if page_size not in allowed_page_sizes:
        page_size = 25

    free_30mins_meetings = (
        Free30Mins.objects
        .select_related('matter_type', 'created_by')
        .prefetch_related('attendees')
        .order_by('-date', '-start_time', '-id')
    )

    if search_query:
        free_30mins_meetings = free_30mins_meetings.filter(
            Q(attendees__name__icontains=search_query)
            | Q(attendees__email__icontains=search_query)
            | Q(matter_type__type__icontains=search_query)
            | Q(created_by__username__icontains=search_query)
            | Q(created_by__first_name__icontains=search_query)
            | Q(created_by__last_name__icontains=search_query)
        ).distinct()

    paginator = Paginator(free_30mins_meetings, page_size)
    meetings_page = paginator.get_page(request.GET.get('page'))

    return render(request, 'free_30mins.html', {
        'free30_mins_form': free30_mins_form,
        'free30_mins_attendees_form': free30_mins_attendees_form,
        'meetings': meetings_page,
        'page_obj': meetings_page,
        'paginator': paginator,
        'page_range': paginator.get_elided_page_range(meetings_page.number),
        'search_query': search_query,
        'page_size': page_size,
        'allowed_page_sizes': allowed_page_sizes,
    })


def download_free30mins(request, id):
    obj = Free30Mins.objects.filter(id=id).first()

    obj_dict = {
        'id': obj.id,
        # Adjust attribute if needed
        'matter_type': obj.matter_type.type if obj.matter_type else None,
        # Assuming QuillField stores HTML content
        'notes': mark_safe(obj.notes.html),
        'date': obj.date,
        'start_time': obj.start_time,
        'finish_time': obj.finish_time,
        # Adjust attribute if needed
        'attendees': [{
            'name': attendee.name,
            'address_line1': attendee.address_line1,
            'address_line2': attendee.address_line2,
            'county': attendee.county,
            'postcode': attendee.postcode,
            'email': attendee.email,
            'contact_number': attendee.contact_number,
        } for attendee in obj.attendees.all()],
        'fee_earner': obj.fee_earner.username if obj.fee_earner else None,
        'created_by': obj.created_by.username if obj.created_by else None,
        'timestamp': obj.timestamp,
    }

    html_string = render_to_string(
        'download_templates/free_30mins.html', obj_dict)
    pdf_file = HTML(string=html_string).write_pdf()

    return HttpResponse(pdf_file, content_type='application/pdf')


def edit_free30mins(request, id):
    instance = get_object_or_404(Free30Mins, pk=id)
    if request.method == 'POST':

        duplicate_obj = copy.deepcopy(instance)
        form = Free30MinsForm(
            request.POST, instance=instance)

        if form.is_valid():
            changed_fields = form.changed_data
            changes = {}
            for field in changed_fields:
                if field == 'content':
                    changes[field] = {
                        'old_value': duplicate_obj.content.html,
                        'new_value': None
                    }
                else:
                    changes[field] = {
                        'old_value': str(getattr(duplicate_obj, field)),
                        'new_value': None
                    }
            form.save()

            for field in changed_fields:
                if field == 'content':
                    changes[field]['new_value'] = instance.content.html
                else:
                    changes[field]['new_value'] = str(
                        getattr(instance, field))

            create_modification(
                user=request.user,
                modified_obj=instance,
                changes=changes
            )
            messages.success(request, 'Successfully updated Free 30 Mins.')
            return redirect('free30mins')
        else:
            error_message = 'Form is not valid. Please correct the errors:'
            for field, errors in form.errors.items():
                error_message += f'\n{field}: {", ".join(errors)}'
            messages.error(request, error_message)
    else:
        form = Free30MinsForm(instance=instance)

    return render(request, 'edit_models.html', {'form': form, 'title': 'Free 30 Mins'})


@login_required
def undertakings(request):
    # If the request is a POST (form submission), process the form data
    if request.method == 'POST':
        request_post_copy = request.POST.copy()

        file = WIP.objects.filter(
            file_number=request_post_copy.get('file_number')).first().id
        request_post_copy['file_number'] = file
        form = UndertakingForm(request_post_copy, request.FILES)
        if form.is_valid():
            undertaking = form.save(commit=False)

            # Automatically assign fields that are not user inputs
            undertaking.created_by = request.user
            undertaking.date_discharged = None
            undertaking.discharged_proof = None
            undertaking.discharged_by = None

            # Save the undertaking object
            try:
                undertaking.save()
            except SharePointClientError as exc:
                logger.error(
                    'SharePoint upload failed creating undertaking: %s', exc)
                messages.error(
                    request,
                    'Could not upload the file to SharePoint (access denied). '
                    'Please contact your administrator.',
                )
                form = UndertakingForm(request_post_copy, request.FILES)
            else:
                messages.success(
                    request, "Undertaking has been successfully created.")
                return redirect('undertakings')
        else:
            # Iterate through form errors and display them
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")

    else:
        # If the request is a GET, load a blank form
        form = UndertakingForm()

    # Fetch all undertakings to display in the template
    undertakings = Undertaking.objects.all()

    undertakings_pending = undertakings.filter(date_discharged=None).count()

    return render(request, 'undertakings.html', {'form': form, 'undertakings': undertakings, 'undertakings_pending': undertakings_pending})


UNDERTAKING_FILE_FIELDS = frozenset({'document_given_on', 'discharged_proof'})


@login_required
def undertaking_file_download(request, pk, field):
    if field not in UNDERTAKING_FILE_FIELDS:
        raise Http404('Invalid file field')

    undertaking = get_object_or_404(Undertaking, pk=pk)
    file_field = getattr(undertaking, field, None)
    if not file_field:
        raise Http404('File not found')

    filename = os.path.basename(file_field.name)
    response = FileResponse(
        file_field.open('rb'),
        content_type='application/octet-stream',
    )
    response['Content-Disposition'] = f'inline; filename="{filename}"'
    return response


@login_required
def edit_undertaking(request, id):
    # Get the specific undertaking object
    undertaking = get_object_or_404(Undertaking, pk=id)

    # Get list of users and WIPs for the dropdowns
    users = CustomUser.objects.all().order_by('username')
    wips = WIP.objects.all().order_by('file_number')

    if request.method == 'POST':
        try:
            # Get the form data from the request
            undertaking.file_number_id = request.POST.get('file_number')
            undertaking.date_given = request.POST.get('date_given')
            undertaking.given_to = request.POST.get('given_to')
            undertaking.description = request.POST.get('description')
            undertaking.given_by_id = request.POST.get('given_by')

            # Handle file upload for document_given_on
            if request.FILES.get('document_given_on'):
                if undertaking.document_given_on:
                    undertaking.document_given_on.delete(save=False)
                undertaking.document_given_on = request.FILES['document_given_on']

            undertaking.date_discharged = request.POST.get('date_discharged')
            undertaking.discharged_by_id = request.POST.get('discharged_by')

            # Handle file upload for discharged_proof
            if request.FILES.get('discharged_proof'):
                if undertaking.discharged_proof:
                    undertaking.discharged_proof.delete(save=False)
                undertaking.discharged_proof = request.FILES['discharged_proof']

            # Save the updated undertaking
            undertaking.save()

            # Display success message
            messages.success(request, 'Undertaking updated successfully.')

            # Redirect to the list view or wherever appropriate
            return redirect('undertakings')

        except Exception as e:
            # If there was an error, log the exception (if needed) and display an error message
            messages.error(request, f'Error updating undertaking: {str(e)}')

    # Render the template with the undertaking object and other context
    context = {
        'undertaking': undertaking,
        'users': users,
        'wips': wips,
    }
    return render(request, 'forms/edit_undertaking.html', context)


@login_required
def management_reports(request):
    users = CustomUser.objects.filter(is_active=True).order_by('username')

    twelve_months_ago = timezone.now() - relativedelta(months=11)
    unique_aml_checks_due = get_aml_checks_due_from_wips(
        WIP.objects.all(), twelve_months_ago, sort_by='name')

    risk_assessments_due = get_risk_assessments_due_queryset(WIP.objects.all())
    cpds = CPDTrainingLog.objects.all()

    return render(request, 'management_reports.html', {
        'users': users,
        'aml_checks_due': unique_aml_checks_due,
        'risk_assessments_due': risk_assessments_due,
        'cpds': cpds
    })


@login_required
def export_user_tasks_pdf(request):
    """Export selected user's tasks as a PDF for printing and sending"""
    user_id = request.GET.get('user_id')

    if not user_id:
        return JsonResponse({'error': 'User ID is required'}, status=400)

    try:
        user = CustomUser.objects.get(id=user_id)
    except CustomUser.DoesNotExist:
        return JsonResponse({'error': 'User not found'}, status=404)

    # Get user's tasks
    tasks = NextWork.objects.filter(person=user).select_related(
        'file_number', 'created_by').order_by('status', 'urgency', 'date')

    # Separate by status
    to_do_tasks = tasks.filter(status='to_do')
    in_progress_tasks = tasks.filter(status='in_progress')
    completed_tasks = tasks.filter(status='completed')

    # Render HTML template for PDF
    html_string = render_to_string('download_templates/user_tasks_export.html', {
        'user': user,
        'to_do_tasks': to_do_tasks,
        'in_progress_tasks': in_progress_tasks,
        'completed_tasks': completed_tasks,
        'export_date': timezone.now(),
        'total_tasks': tasks.count(),
    })

    # Generate PDF
    pdf_file = HTML(string=html_string).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="tasks_{user.first_name}_{user.last_name}_{timezone.now().strftime("%Y%m%d")}.pdf"'

    return response


@login_required
def load_management_tasks(request):
    """AJAX endpoint to load tasks for management kanban board"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method'})

    try:
        data = json.loads(request.body)
        count = data.get('count', 5)
        filter_user_ids = data.get('filter_user_ids', [])

        # Base query
        tasks_query = NextWork.objects.select_related(
            'file_number', 'person', 'created_by')

        # Apply user filter
        if filter_user_ids:
            tasks_query = tasks_query.filter(person_id__in=filter_user_ids)

        # Get tasks by status
        tasks_data = {}
        total_counts = {}

        for status in ['to_do', 'in_progress', 'completed']:
            if status == 'completed':
                # For completed, only show this week's tasks
                one_week_ago = timezone.now() - timedelta(days=7)
                status_tasks = tasks_query.filter(
                    status=status,
                    timestamp__gte=one_week_ago
                ).order_by('-timestamp')
            else:
                # For to_do and in_progress, sort by urgency and due date
                status_tasks = tasks_query.filter(status=status).order_by(
                    '-urgency', 'date', '-timestamp'
                )

            total_counts[status] = status_tasks.count()

            # Limit tasks if count is specified
            if count != 'all':
                status_tasks = status_tasks[:int(count)]

            # Serialize tasks
            tasks_data[status] = []
            for task in status_tasks:
                tasks_data[status].append({
                    'id': task.id,
                    'file_number': task.file_number.file_number if task.file_number else 'N/A',
                    'task': task.task,
                    'date': task.date.isoformat() if task.date else None,
                    'urgency': task.urgency,
                    'status': task.status,
                    'assigned_to': f"{task.person.first_name} {task.person.last_name}" if task.person else 'Unassigned',
                    'created_by': f"{task.created_by.first_name} {task.created_by.last_name}" if task.created_by else 'Unknown',
                    'timestamp': task.timestamp.isoformat() if task.timestamp else None,
                })

        return JsonResponse({
            'success': True,
            'tasks': tasks_data,
            'total_counts': total_counts
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


def calculate_minutes_for_date(user, date):
    billed_minutes = 0
    non_billed_minutes = 0

    # Free 30-Minute Meetings (Non-Billed)
    free_30_mins = Free30Mins.objects.filter(fee_earner=user, date=date)
    for meeting in free_30_mins:
        duration = (datetime.combine(date, meeting.finish_time) -
                    # Convert to minutes
                    datetime.combine(date, meeting.start_time)).seconds / 60
        non_billed_minutes += duration

    # Matter Emails (Billed)
    matter_emails = MatterEmails.objects.filter(
        fee_earner=user, time__date=date)
    for email in matter_emails:
        billed_minutes += (email.units or 0) * 6  # Assuming 1 unit = 6 minutes

    # Attendance Notes
    attendance_notes = MatterAttendanceNotes.objects.filter(
        person_attended=user, date=date)
    for note in attendance_notes:
        duration = (datetime.combine(date, note.finish_time) -
                    # Convert to minutes
                    datetime.combine(date, note.start_time)).seconds / 60
        if note.is_charged:
            billed_minutes += duration
        else:
            non_billed_minutes += duration

    # Total minutes in a day
    total_minutes = 7.5 * 60  # 7.5 hours in minutes
    missing_minutes = max(
        0, total_minutes - (billed_minutes + non_billed_minutes))

    return {
        "billed_minutes": round(billed_minutes),
        "non_billed_minutes": round(non_billed_minutes),
        "missing_minutes": round(missing_minutes),
    }


def calculate_weekly_report(user, start_date):
    weekly_report = []
    for i in range(7):  # Iterate through the week
        date = start_date + timedelta(days=i)
        # Skip Saturdays and Sundays
        if date.weekday() in [5, 6]:
            continue
        daily_data = calculate_minutes_for_date(user, date)
        weekly_report.append({
            "date": date.strftime("%a, %d/%m/%Y "),
            **daily_data,
        })
    return weekly_report


@login_required
def weekly_report_view(request):
    user_id = request.GET.get("user")
    week_start = request.GET.get("week_start")

    if not user_id or not week_start:
        return JsonResponse({"error": "Invalid parameters"}, status=400)

    try:
        user = CustomUser.objects.get(id=user_id)
        week_start_date = datetime.strptime(week_start, "%Y-%m-%d").date()
        data = calculate_weekly_report(user, week_start_date)
        return JsonResponse(data, safe=False)
    except CustomUser.DoesNotExist:
        return JsonResponse({"error": "User not found"}, status=404)


@login_required
def policies_read_per_user(request):
    user_id = request.GET.get("user_id")

    if not user_id:
        return JsonResponse({"error": "User ID is required"}, status=400)

    try:
        user = CustomUser.objects.get(id=user_id)
    except CustomUser.DoesNotExist:
        return JsonResponse({"error": "User not found"}, status=404)

    # Total number of policies
    total_policies = Policy.objects.count()

    # Get the latest version of each policy
    latest_versions = PolicyVersion.objects.annotate(
        max_version=Max('policy__versions__version_number')
    ).filter(version_number=F('max_version'))

    # IDs of latest policy versions
    latest_version_ids = latest_versions.values_list('id', flat=True)

    # Policies read by the user (matching the latest versions)
    read_policy_ids = PoliciesRead.objects.filter(
        read_by=user,
        policy_version_id__in=latest_version_ids
    ).values_list('policy_id', flat=True).distinct()

    # Count of latest versions read
    latest_versions_read_count = read_policy_ids.count()

    # Get unread policies
    unread_policies = Policy.objects.filter(
        ~Q(id__in=read_policy_ids)).order_by('description')

    # Prepare unread policies descriptions
    unread_policies_descriptions = list(
        unread_policies.values('id', 'description'))

    # Response
    return JsonResponse({
        "user": user.username,
        "total_policies": total_policies,
        "latest_versions_read": latest_versions_read_count,
        "policies_unread": total_policies - latest_versions_read_count,
        "unread_policies": unread_policies_descriptions
    })


@login_required
def download_aml_checks_due(request):
    twelve_months_ago = timezone.now() - relativedelta(months=11)
    unique_aml_checks_due = get_aml_checks_due_from_wips(
        WIP.objects.all(), twelve_months_ago, sort_by='name')

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="aml_checks_due_{timezone.now()}.csv"'

    writer = csv.writer(response)

    writer.writerow(['AML Checks Due'])
    writer.writerow(['Name', 'Type', 'Date of Last AML Check'])

    for check in unique_aml_checks_due:
        writer.writerow([check['entity_name'], check['entity_type'],
                         check['date_of_last_aml']])

    return response


@login_required
def download_user_risk_assessments_due(request):
    user = CustomUser.objects.get(username=request.user)
    risk_scope = request.GET.get('risk_scope', 'associated')
    risk_scope_wips, validated_risk_scope = get_dashboard_risk_scope_wips(
        user, risk_scope
    )
    risk_assessments_due = get_risk_assessments_due_queryset(
        risk_scope_wips
    )

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = (
        f'attachment; filename="risk_assessments_due_{user.username}_{validated_risk_scope}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.csv"'
    )
    writer = csv.writer(response)

    writer.writerow(['Risk Assessments Due'])
    if validated_risk_scope == 'all_active':
        writer.writerow(['Scope', 'All Open and To Be Closed files'])
    else:
        writer.writerow(['Scope', 'My associated files'])
    writer.writerow([
        'File Number',
        'Matter Description',
        'Client 1',
        'Client 2',
        'Last Risk Assessment Date',
        'Last Ongoing Monitoring Date',
        'Reason Due'
    ])

    for assessment in risk_assessments_due:
        reason_due = 'No risk assessment completed'
        if assessment.latest_assessment_date:
            reason_due = 'No ongoing monitoring in the last year'

        writer.writerow([
            assessment.file_number,
            assessment.matter_description,
            assessment.client1.name if assessment.client1 else '',
            assessment.client2.name if assessment.client2 else '',
            assessment.latest_assessment_date,
            assessment.latest_monitoring_date,
            reason_due
        ])

    return response


@login_required
def download_user_key_documents_due(request):
    user = CustomUser.objects.get(username=request.user)
    key_doc_scope = request.GET.get('key_doc_scope', 'associated')
    key_doc_scope_wips, validated_key_doc_scope = get_dashboard_key_document_scope_wips(
        user, key_doc_scope
    )
    missing_alerts = get_missing_key_document_alerts(key_doc_scope_wips)
    expiry_alerts = get_key_document_expiry_alerts(key_doc_scope_wips)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = (
        f'attachment; filename="key_documents_{user.username}_{validated_key_doc_scope}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.csv"'
    )
    writer = csv.writer(response)

    writer.writerow(['Key Document Issues'])
    if validated_key_doc_scope == 'all_active':
        writer.writerow(['Scope', 'All Open and To Be Closed files'])
    else:
        writer.writerow(['Scope', 'My associated files'])
    writer.writerow([
        'Issue Type',
        'File Number(s)',
        'Client',
        'Document Category',
        'Document Type',
        'Document Reference',
        'Expiry Date',
        'Status'
    ])

    for alert in missing_alerts:
        writer.writerow([
            'Missing',
            alert['file_number'],
            alert['client_name'],
            alert['document_category'],
            '',
            '',
            '',
            'Missing'
        ])

    for alert in expiry_alerts:
        writer.writerow([
            'Expiry',
            ', '.join(alert['file_numbers']),
            alert['client_name'],
            alert['document_category'],
            alert['document_type'],
            alert['document_reference'],
            alert['expiry_date'],
            'Expired' if alert['status'] == 'expired' else 'Due soon'
        ])

    return response


@login_required
def download_risk_assessments_due(request):
    risk_assessments_due = get_risk_assessments_due_queryset(WIP.objects.all())

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="risk_assessments_due_{timezone.now()}.csv"'

    writer = csv.writer(response)

    writer.writerow(['Risk Assessments Due'])
    writer.writerow([
        'File Number',
        'Matter Description',
        'Client 1',
        'Client 2',
        'Last Risk Assessment Date',
        'Last Ongoing Monitoring Date',
        'Reason Due'
    ])

    for assessment in risk_assessments_due:
        reason_due = 'No risk assessment completed'
        if assessment.latest_assessment_date:
            reason_due = 'No ongoing monitoring in the last year'

        writer.writerow([
            assessment.file_number,
            assessment.matter_description,
            assessment.client1.name if assessment.client1 else '',
            assessment.client2.name if assessment.client2 else '',
            assessment.latest_assessment_date,
            assessment.latest_monitoring_date,
            reason_due
        ])

    return response


@login_required
def add_memo(request):
    if request.method == 'POST':
        form = MemoForm(request.POST)
        if form.is_valid():
            memo = form.save(commit=False)
            memo.created_by = request.user
            memo.save()
            messages.success(request, 'Memo successfully added.')

        else:
            messages.error(
                request, 'There was an error adding the memo. Please check the form and try again.')

    return redirect('profile_page')


@login_required
def edit_memo(request, memo_id):
    memo = get_object_or_404(Memo, id=memo_id)
    if request.method == 'POST':
        form = MemoForm(request.POST, instance=memo)
        if form.is_valid():
            form.save()
            messages.success(request, 'Memo successfully updated.')
            return redirect('profile_page')
        else:
            messages.error(
                request, 'There was an error updating the memo. Please check the form and try again.')
    else:
        form = MemoForm(instance=memo)
    return render(request, 'edit_memo.html', {'form': form})


@login_required
def delete_memo(request, memo_id):
    memo = get_object_or_404(Memo, id=memo_id)
    if request.method == 'POST':
        memo.delete()
        messages.success(request, 'Memo successfully deleted.')
        return redirect('profile_page')
    else:
        messages.warning(request, 'Are you sure you want to delete this memo?')
    return render(request, 'confirm_delete.html', {'memo': memo})


@login_required
def read_memo(request, memo_id):
    memo = get_object_or_404(Memo, id=memo_id)
    PoliciesRead.objects.get_or_create(memo=memo, read_by=request.user)
    messages.success(request, 'Memo marked as read.')
    return redirect('profile_page')


# Bundle Views
def _bundle_pdf_progress_key(user_id, bundle_id):
    return f'bundle_pdf_progress:{user_id}:{bundle_id}'


def _set_bundle_pdf_progress(user_id, bundle_id, **payload):
    cache.set(
        _bundle_pdf_progress_key(user_id, bundle_id),
        {
            'updated_at': timezone.now().isoformat(),
            **payload,
        },
        900,
    )


def _clear_bundle_pdf_progress(user_id, bundle_id):
    cache.delete(_bundle_pdf_progress_key(user_id, bundle_id))


def _generate_bundle_pdf_job(bundle_id, user_id):
    from django.db import close_old_connections

    close_old_connections()
    try:
        bundle = Bundle.objects.get(pk=bundle_id, created_by_id=user_id)
        user = CustomUser.objects.filter(pk=user_id).first()

        def progress_callback(percent, message):
            _set_bundle_pdf_progress(
                user_id,
                bundle_id,
                status='running',
                percent=percent,
                message=message,
            )

        success, error, _regenerated = _ensure_bundle_final_pdf(
            bundle,
            user=user,
            progress_callback=progress_callback,
        )
        if success:
            _set_bundle_pdf_progress(
                user_id,
                bundle_id,
                status='ready',
                percent=100,
                message='PDF ready',
                ready=True,
            )
        else:
            _set_bundle_pdf_progress(
                user_id,
                bundle_id,
                status='error',
                percent=0,
                message=error or 'PDF generation failed',
                error=error or 'PDF generation failed',
            )
    except Exception as exc:
        logger.exception(
            'Background PDF generation failed for bundle %s: %s', bundle_id, exc)
        _set_bundle_pdf_progress(
            user_id,
            bundle_id,
            status='error',
            percent=0,
            message=str(exc),
            error=str(exc),
        )
    finally:
        close_old_connections()


def _touch_bundle(bundle):
    """Mark bundle content as changed so cached PDF is regenerated on download."""
    now = timezone.now()
    Bundle.objects.filter(pk=bundle.pk).update(updated_at=now)
    bundle.updated_at = now


def _bundle_pdf_is_current(bundle, verify_file=True):
    if not bundle.pdf_is_current():
        return False
    if not verify_file:
        return True
    if not bundle.final_pdf:
        return False
    try:
        if default_storage.exists(bundle.final_pdf.name):
            return True
        # SharePoint can lag between upload and exists(); opening verifies readiness.
        with bundle.final_pdf.open('rb'):
            pass
        return True
    except Exception:
        return False


def _ensure_bundle_final_pdf(bundle, user=None, progress_callback=None):
    """Generate the final PDF when missing or out of date. Returns (success, error_message, regenerated)."""
    if _bundle_pdf_is_current(bundle):
        if progress_callback:
            progress_callback(100, 'PDF is ready')
        return True, None, False

    cache_obj = None
    try:
        if progress_callback:
            progress_callback(5, 'Collecting documents...')
        cache_obj = BundleTempCache(bundle)
        documents_info = _collect_bundle_documents(bundle, cache=cache_obj)
        if not documents_info:
            return False, 'Cannot generate bundle: no valid documents with pages.', False

        if progress_callback:
            progress_callback(
                15,
                f'Building index for {len(documents_info)} document{"s" if len(documents_info) != 1 else ""}...',
            )
        pdf_content = _generate_bundle_pdf(
            bundle,
            cache=cache_obj,
            progress_callback=progress_callback,
            documents_info=documents_info,
        )

        if progress_callback:
            progress_callback(95, 'Saving PDF...')

        if bundle.final_pdf:
            try:
                default_storage.delete(bundle.final_pdf.name)
            except Exception:
                logger.warning(
                    'Could not delete previous final PDF for bundle %s', bundle.id)

        bundle.final_pdf.save(f'{bundle.uuid}.pdf', ContentFile(pdf_content))
        bundle.refresh_from_db(fields=['updated_at'])
        now = timezone.now()
        if bundle.updated_at and bundle.updated_at > now:
            now = bundle.updated_at
        Bundle.objects.filter(pk=bundle.pk).update(
            final_pdf=bundle.final_pdf.name,
            pdf_generated_at=now,
        )
        bundle.pdf_generated_at = now
        if user:
            log_bundle_event(
                user,
                bundle,
                'Bundle PDF generated',
                document_count=len(documents_info),
            )
        if progress_callback:
            progress_callback(100, 'PDF ready')
        return True, None, True
    except Exception as e:
        return False, str(e), False
    finally:
        if cache_obj is not None:
            cache_obj.cleanup()


@login_required
def bundle_list(request, file_number=None):
    """Redirect legacy bundle list URLs to the home page or dashboard."""
    if file_number:
        return redirect('home', file_number=file_number)
    return redirect('user_dashboard')


@login_required
def bundle_create(request, file_number=None):
    """Create a new bundle and open the editor."""
    if request.method == 'POST':
        bundle_name = request.POST.get('bundle_name', '').strip() or 'PDF'
        file_number_str = request.POST.get('file_number', file_number)

        wip_file = None
        if file_number_str:
            try:
                wip_file = WIP.objects.get(file_number=file_number_str)
            except WIP.DoesNotExist:
                if request.headers.get('Content-Type') == 'application/json' or 'application/json' in request.headers.get('Accept', ''):
                    return JsonResponse({'error': f'File {file_number_str} not found.'}, status=400)
                messages.error(request, f'File {file_number_str} not found.')
                return redirect('bundle_create')

        bundle = Bundle.objects.create(
            name=bundle_name,
            file_number=wip_file,
            created_by=request.user
        )
        log_bundle_event(
            request.user,
            bundle,
            'Bundle created',
            name=bundle_name,
            file_number=wip_file.file_number if wip_file else '',
        )
        BundleSection.objects.create(
            bundle=bundle,
            heading='Section 1',
            order=1,
        )

        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.content_type.startswith('multipart/form-data'):
            return JsonResponse({
                'success': True,
                'bundle_id': bundle.id,
                'message': f'Bundle "{bundle_name}" created successfully.'
            })

        return redirect(f"{reverse('bundle_edit', kwargs={'bundle_id': bundle.id})}?new=1")

    wip_file = None
    if file_number:
        wip_file = get_object_or_404(WIP, file_number=file_number)

    bundle = Bundle.objects.create(
        name='PDF',
        file_number=wip_file,
        created_by=request.user,
    )
    log_bundle_event(
        request.user,
        bundle,
        'Bundle created',
        name=bundle.name,
        file_number=wip_file.file_number if wip_file else '',
    )
    BundleSection.objects.create(
        bundle=bundle,
        heading='Section 1',
        order=1,
    )
    return redirect(f"{reverse('bundle_edit', kwargs={'bundle_id': bundle.id})}?new=1")


@login_required
def bundle_update(request, bundle_id):
    """Update bundle metadata."""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    if request.method == 'POST':
        bundle_name = request.POST.get('bundle_name', '').strip()
        if not bundle_name:
            return JsonResponse({'error': 'Bundle name is required'}, status=400)

        old_name = bundle.name
        name_changed = bundle_name != old_name

        file_number_str = request.POST.get('file_number')
        file_number_changed = False
        linked_file_number = ''
        if file_number_str is not None and not bundle.file_number:
            file_number_str = file_number_str.strip()
            if file_number_str:
                try:
                    wip_file = WIP.objects.get(file_number=file_number_str)
                except WIP.DoesNotExist:
                    return JsonResponse({'error': f'File {file_number_str} not found.'}, status=400)
                bundle.file_number = wip_file
                file_number_changed = True
                linked_file_number = wip_file.file_number

        if not name_changed and not file_number_changed:
            return JsonResponse({
                'success': True,
                'bundle': {
                    'id': bundle.id,
                    'name': bundle.name,
                    'file_number': bundle.file_number.file_number if bundle.file_number else '',
                }
            })

        update_fields = []
        if name_changed:
            bundle.name = bundle_name
            update_fields.append('name')
        if file_number_changed:
            update_fields.append('file_number')
        bundle.save(update_fields=update_fields)
        if name_changed:
            log_bundle_event(
                request.user,
                bundle,
                'Bundle renamed',
                name={'old_value': old_name, 'new_value': bundle_name},
            )
        if file_number_changed:
            log_bundle_event(
                request.user,
                bundle,
                'Bundle linked to matter',
                file_number=linked_file_number,
            )
        return JsonResponse({
            'success': True,
            'bundle': {
                'id': bundle.id,
                'name': bundle.name,
                'file_number': bundle.file_number.file_number if bundle.file_number else '',
            }
        })

    return JsonResponse({'error': 'Invalid request'}, status=400)


def _default_court_parties():
    return [
        {'side': 'claimant', 'name': '', 'role': 'Claimant 1'},
        {'side': 'defendant', 'name': '', 'role': 'Defendant 1'},
    ]


def _normalise_court_parties(raw_parties):
    parties = []
    claimant_count = 0
    defendant_count = 0
    if not isinstance(raw_parties, list):
        return _default_court_parties()

    for entry in raw_parties:
        if not isinstance(entry, dict):
            continue
        side = entry.get('side')
        if side not in {'claimant', 'defendant'}:
            continue
        name = str(entry.get('name', '')).strip()
        role = str(entry.get('role', '')).strip()
        if side == 'claimant':
            claimant_count += 1
            if not role:
                role = f'Claimant {claimant_count}'
        else:
            defendant_count += 1
            if not role:
                role = f'Defendant {defendant_count}'
        parties.append({'side': side, 'name': name, 'role': role})

    if not parties:
        return _default_court_parties()
    if not any(party['side'] == 'claimant' for party in parties):
        parties.insert(
            0, {'side': 'claimant', 'name': '', 'role': 'Claimant 1'})
    if not any(party['side'] == 'defendant' for party in parties):
        parties.append(
            {'side': 'defendant', 'name': '', 'role': 'Defendant 1'})
    return parties


def _bundle_court_settings_dict(bundle):
    parties = _normalise_court_parties(bundle.court_parties)
    return {
        'is_court_bundle': bundle.is_court_bundle,
        'court_name': bundle.court_name or '',
        'case_number_type': bundle.case_number_type or Bundle.CASE_NUMBER_CLAIM,
        'case_number': bundle.case_number or '',
        'index_title': bundle.index_title or 'Index to the Bundle',
        'hearing_line': bundle.hearing_line or '',
        'conference_line': bundle.conference_line or '',
        'parties': parties,
    }


@login_required
def bundle_court_update(request, bundle_id):
    """Update court bundle heading settings."""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request'}, status=400)

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON payload'}, status=400)

    is_court_bundle = bool(payload.get('is_court_bundle'))
    court_name = str(payload.get('court_name', '')).strip()
    case_number_type = payload.get(
        'case_number_type', Bundle.CASE_NUMBER_CLAIM)
    if case_number_type not in {Bundle.CASE_NUMBER_CLAIM, Bundle.CASE_NUMBER_CASE}:
        return JsonResponse({'error': 'Invalid case number type'}, status=400)

    case_number = str(payload.get('case_number', '')).strip()
    index_title = str(payload.get('index_title', '')
                      ).strip() or 'Index to the Bundle'
    hearing_line = str(payload.get('hearing_line', '')).strip()
    conference_line = str(payload.get('conference_line', '')).strip()
    parties = _normalise_court_parties(payload.get('parties', []))

    bundle.is_court_bundle = is_court_bundle
    bundle.court_name = court_name
    bundle.case_number_type = case_number_type
    bundle.case_number = case_number
    bundle.index_title = index_title
    bundle.hearing_line = hearing_line
    bundle.conference_line = conference_line
    bundle.court_parties = parties
    bundle.save(update_fields=[
        'is_court_bundle',
        'court_name',
        'case_number_type',
        'case_number',
        'index_title',
        'hearing_line',
        'conference_line',
        'court_parties',
    ])
    _touch_bundle(bundle)
    log_bundle_event(request.user, bundle, 'Court bundle settings updated')

    return JsonResponse({
        'success': True,
        'court': _bundle_court_settings_dict(bundle),
    })


@login_required
def bundle_edit(request, bundle_id):
    """Edit bundle - manage sections and documents"""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    sections = list(bundle.sections.prefetch_related(
        'documents').order_by('order'))
    for section in sections:
        section.bundle_documents = section.ordered_documents()

    context = {
        'bundle': bundle,
        'sections': sections,
        'court_bundle_json': json.dumps(_bundle_court_settings_dict(bundle)),
        'is_new': request.GET.get('new') == '1',
    }
    if not bundle.file_number:
        context['files'] = WIP.objects.filter(
            file_status__status='Open').order_by('file_number')
    return render(request, 'bundle_edit.html', context)


@login_required
def bundle_section_add(request, bundle_id):
    """Add a new section to the bundle"""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    if request.method == 'POST':
        _touch_bundle(bundle)

        heading = request.POST.get('heading')

        if not heading:
            return JsonResponse({'error': 'Section heading is required'}, status=400)

        # Get the next order number
        last_section = bundle.sections.order_by('-order').first()
        next_order = (last_section.order + 1) if last_section else 1

        section = BundleSection.objects.create(
            bundle=bundle,
            heading=heading,
            order=next_order
        )
        log_bundle_event(
            request.user,
            bundle,
            'Section added',
            section=heading,
        )

        return JsonResponse({
            'success': True,
            'section_id': section.id,
            'section': {
                'id': section.id,
                'heading': section.heading,
                'order': section.order,
                'date_sort': section.date_sort,
            }
        })

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_section_delete(request, section_id):
    """Delete a bundle section"""
    section = get_object_or_404(
        BundleSection, id=section_id, bundle__created_by=request.user)

    if request.method == 'POST':
        _touch_bundle(section.bundle)
        bundle = section.bundle
        section_heading = section.heading
        section.delete()
        log_bundle_event(
            request.user,
            bundle,
            'Section deleted',
            section=section_heading,
        )
        return JsonResponse({'success': True})

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_section_update(request, section_id):
    """Update a bundle section heading or date sort."""
    section = get_object_or_404(
        BundleSection, id=section_id, bundle__created_by=request.user)

    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request'}, status=400)

    _touch_bundle(section.bundle)

    if 'date_sort' in request.POST:
        date_sort = request.POST.get('date_sort', '')
        valid_sorts = {choice[0] for choice in BundleSection.DATE_SORT_CHOICES}
        if date_sort not in valid_sorts:
            return JsonResponse({'error': 'Invalid sort order'}, status=400)

        if date_sort != section.date_sort:
            old_sort = section.date_sort
            section.date_sort = date_sort
            section.save(update_fields=['date_sort'])
            log_bundle_event(
                request.user,
                section.bundle,
                'Section sort order changed',
                section=section.heading,
                date_sort={'old_value': old_sort, 'new_value': date_sort},
            )

        return JsonResponse({
            'success': True,
            'section': {
                'id': section.id,
                'date_sort': section.date_sort,
                'document_ids': _section_ordered_document_ids(section),
            }
        })

    heading = request.POST.get('heading', '').strip()
    if not heading:
        return JsonResponse({'error': 'Section heading is required'}, status=400)

    old_heading = section.heading
    section.heading = heading
    section.save(update_fields=['heading'])
    if old_heading != heading:
        log_bundle_event(
            request.user,
            section.bundle,
            'Section renamed',
            section={'old_value': old_heading, 'new_value': heading},
        )
    return JsonResponse({
        'success': True,
        'section': {
            'id': section.id,
            'heading': section.heading,
        }
    })


@login_required
def bundle_section_reorder(request, bundle_id):
    """Reorder sections in the bundle"""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    if request.method == 'POST':
        _touch_bundle(bundle)

        try:
            section_orders = _parse_order_ids(
                request.POST.getlist('section_orders[]'))
        except ValueError:
            return JsonResponse({'error': 'Invalid section order'}, status=400)
        existing_section_ids = list(bundle.sections.order_by(
            'order').values_list('id', flat=True))

        if set(section_orders) != set(existing_section_ids) or len(section_orders) != len(existing_section_ids):
            return JsonResponse({'error': 'Invalid section order'}, status=400)

        _update_order_safely(BundleSection.objects.filter(
            bundle=bundle), section_orders)
        log_bundle_event(request.user, bundle, 'Sections reordered')

        return JsonResponse({'success': True})

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_document_upload(request, section_id):
    """Upload a document to a section"""
    section = get_object_or_404(
        BundleSection, id=section_id, bundle__created_by=request.user)

    if request.method == 'POST':
        _touch_bundle(section.bundle)

        files = request.FILES.getlist('files[]')
        descriptions = request.POST.getlist('descriptions[]')
        dates = request.POST.getlist('dates[]')

        if not files:
            return JsonResponse({'error': 'No files uploaded'}, status=400)

        uploaded_docs = []

        for i, file in enumerate(files):
            if not file.name.lower().endswith('.pdf'):
                return JsonResponse(
                    {'error': f'Only PDF files are allowed: {file.name}'},
                    status=400,
                )

            description = descriptions[i] if i < len(descriptions) else ''
            date_str = dates[i] if i < len(dates) and dates[i] else None

            parsed_description, parsed_date = parse_bundle_filename(file.name)
            if not description.strip():
                description = parsed_description

            if not description:
                return JsonResponse({'error': f'Description required for file {file.name}'}, status=400)

            # Parse date
            doc_date = None
            if date_str:
                try:
                    doc_date = datetime.strptime(date_str, '%Y-%m-%d').date()
                except ValueError:
                    return JsonResponse({'error': f'Invalid date format for {file.name}'}, status=400)
            elif parsed_date:
                doc_date = parsed_date

            # Get next order
            last_doc = section.documents.order_by('-order').first()
            next_order = (last_doc.order + 1) if last_doc else 1

            # Create document
            document = BundleDocument.objects.create(
                section=section,
                file=file,
                description=description,
                date=doc_date,
                order=next_order
            )

            uploaded_docs.append({
                'id': document.id,
                'description': document.description,
                'date': document.date.strftime('%Y-%m-%d') if document.date else '',
                'filename': document.file.name,
                'order': document.order
            })
            log_bundle_event(
                request.user,
                section.bundle,
                'Document uploaded',
                section=section.heading,
                document=description,
            )

        return JsonResponse({
            'success': True,
            'documents': uploaded_docs,
            'section': {
                'id': section.id,
                'date_sort': section.date_sort,
                'document_ids': _section_ordered_document_ids(section),
            },
        })

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_document_file(request, document_id):
    """Serve a bundle document PDF for in-app page previews."""
    document = get_object_or_404(
        BundleDocument, id=document_id, section__bundle__created_by=request.user)

    if not document.file:
        raise Http404('Document file not found')

    filename = os.path.basename(document.file.name)
    response = FileResponse(
        document.file.open('rb'),
        content_type='application/pdf',
    )
    response['Content-Disposition'] = f'inline; filename="{filename}"'
    return response


@login_required
def bundle_document_update(request, document_id):
    """Update a bundle document's description and date."""
    document = get_object_or_404(
        BundleDocument, id=document_id, section__bundle__created_by=request.user)

    if request.method == 'POST':
        _touch_bundle(document.section.bundle)

        description = request.POST.get('description', '').strip()
        if not description:
            return JsonResponse({'error': 'Document name is required'}, status=400)

        date_str = request.POST.get('date', '')
        doc_date = None
        if date_str:
            try:
                doc_date = datetime.strptime(date_str, '%Y-%m-%d').date()
            except ValueError:
                return JsonResponse({'error': 'Invalid date format'}, status=400)

        old_description = document.description
        old_date = document.date
        document.description = description
        document.date = doc_date
        document.save(update_fields=['description', 'date'])
        log_bundle_event(
            request.user,
            document.section.bundle,
            'Document updated',
            section=document.section.heading,
            description={'old_value': old_description,
                         'new_value': description},
            date={'old_value': old_date or '', 'new_value': doc_date or ''},
        )
        response = {
            'success': True,
            'document': {
                'id': document.id,
                'description': document.description,
                'date': document.date.strftime('%Y-%m-%d') if document.date else '',
            }
        }
        if document.section.date_sort != BundleSection.DATE_SORT_MANUAL:
            response['section'] = {
                'id': document.section.id,
                'document_ids': _section_ordered_document_ids(document.section),
            }
        return JsonResponse(response)

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_document_delete(request, document_id):
    """Delete a bundle document"""
    document = get_object_or_404(
        BundleDocument, id=document_id, section__bundle__created_by=request.user)

    if request.method == 'POST':
        _touch_bundle(document.section.bundle)

        # Delete the file
        if document.file:
            try:
                default_storage.delete(document.file.name)
            except:
                pass  # File might not exist

        section_id = document.section_id
        bundle = document.section.bundle
        doc_description = document.description
        section_heading = document.section.heading
        document.delete()
        log_bundle_event(
            request.user,
            bundle,
            'Document deleted',
            section=section_heading,
            document=doc_description,
        )
        remaining_documents = BundleDocument.objects.filter(
            section_id=section_id).count()
        return JsonResponse({
            'success': True,
            'section_id': section_id,
            'document_count': remaining_documents,
        })

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_document_pages_update(request, document_id):
    """Get or update the selected page order for a bundle document."""
    document = get_object_or_404(
        BundleDocument, id=document_id, section__bundle__created_by=request.user)

    if request.method == 'GET':
        page_choices = _get_bundle_document_page_choices(document)
        if not page_choices:
            return JsonResponse({'error': 'Could not read document pages'}, status=400)

        included_pages = [
            page['number'] for page in page_choices if page['included']
        ]
        return JsonResponse({
            'success': True,
            'page_count': document.page_count,
            'page_choices': page_choices,
            'page_summary': _format_page_order_summary(
                included_pages, document.page_count),
        })

    if request.method == 'POST':
        _touch_bundle(document.section.bundle)

        try:
            page_order = _parse_order_ids(request.POST.getlist('page_order[]'))
        except ValueError:
            return JsonResponse({'error': 'Invalid page order'}, status=400)

        page_count = _get_pdf_page_count(document)
        if (
            not page_order
            or len(page_order) != len(set(page_order))
            or any(page_number < 1 or page_number > page_count for page_number in page_order)
        ):
            return JsonResponse({'error': 'Invalid page order'}, status=400)

        old_page_order = document.page_order
        document.page_order = page_order if page_order != list(
            range(1, page_count + 1)) else None
        document.save(update_fields=['page_order'])
        log_bundle_event(
            request.user,
            document.section.bundle,
            'Document pages updated',
            section=document.section.heading,
            document=document.description,
            page_order={'old_value': old_page_order or 'all',
                        'new_value': page_order},
        )

        return JsonResponse({
            'success': True,
            'page_order': page_order,
            'page_summary': _format_page_order_summary(page_order, page_count),
        })

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_document_reorder(request, section_id):
    """Reorder documents within a section"""
    section = get_object_or_404(
        BundleSection, id=section_id, bundle__created_by=request.user)

    if request.method == 'POST':
        _touch_bundle(section.bundle)

        try:
            document_orders = _parse_order_ids(
                request.POST.getlist('document_orders[]'))
        except ValueError:
            return JsonResponse({'error': 'Invalid document order'}, status=400)
        existing_document_ids = list(section.documents.order_by(
            'order').values_list('id', flat=True))

        if set(document_orders) != set(existing_document_ids) or len(document_orders) != len(existing_document_ids):
            return JsonResponse({'error': 'Invalid document order'}, status=400)

        _update_order_safely(BundleDocument.objects.filter(
            section=section), document_orders)
        log_bundle_event(
            request.user,
            section.bundle,
            'Documents reordered',
            section=section.heading,
        )

        return JsonResponse({'success': True})

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_generate(request, bundle_id):
    """Generate the final PDF bundle"""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    if request.method == 'POST':
        success, error, _regenerated = _ensure_bundle_final_pdf(
            bundle, user=request.user)
        if not success:
            messages.error(request, error)
            return redirect('bundle_edit', bundle_id=bundle.id)

        messages.success(request, 'Bundle generated successfully!')
        return redirect('bundle_edit', bundle_id=bundle.id)

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
def bundle_view(request, bundle_id):
    """View a finalized bundle"""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)
    sections = bundle.sections.all().order_by('order')

    # Calculate total documents
    total_docs = sum(section.documents.count() for section in sections)

    context = {
        'bundle': bundle,
        'sections': sections,
        'total_docs': total_docs,
    }
    return render(request, 'bundle_view.html', context)


@login_required
def bundle_pdf_prepare(request, bundle_id):
    """Start background PDF generation when the cached bundle PDF is stale."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request'}, status=400)

    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    if _bundle_pdf_is_current(bundle):
        _clear_bundle_pdf_progress(request.user.id, bundle_id)
        return JsonResponse({
            'ready': True,
            'percent': 100,
            'message': 'PDF is ready',
        })

    progress = cache.get(_bundle_pdf_progress_key(request.user.id, bundle_id))
    if progress and progress.get('status') == 'running':
        return JsonResponse({'started': True, **progress})

    _set_bundle_pdf_progress(
        request.user.id,
        bundle_id,
        status='running',
        percent=0,
        message='Starting PDF generation...',
    )
    thread = threading.Thread(
        target=_generate_bundle_pdf_job,
        args=(bundle.id, request.user.id),
        daemon=True,
    )
    thread.start()
    return JsonResponse({
        'started': True,
        'percent': 0,
        'message': 'Starting PDF generation...',
        'status': 'running',
    })


@login_required
def bundle_pdf_status(request, bundle_id):
    """Poll PDF generation progress for a bundle."""
    get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    progress = cache.get(_bundle_pdf_progress_key(request.user.id, bundle_id))
    if progress:
        if progress.get('status') == 'error':
            return JsonResponse(progress, status=400)
        if progress.get('status') == 'ready':
            return JsonResponse({**progress, 'ready': True})
        return JsonResponse(progress)

    bundle = Bundle.objects.get(pk=bundle_id)
    if _bundle_pdf_is_current(bundle):
        return JsonResponse({
            'ready': True,
            'status': 'ready',
            'percent': 100,
            'message': 'PDF is ready',
        })

    return JsonResponse({
        'status': 'pending',
        'percent': 0,
        'message': 'Waiting to start...',
    })


@login_required
def bundle_download(request, bundle_id):
    """Download the bundle PDF, generating it first when out of date."""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)
    bundle.refresh_from_db()

    serve_only = request.headers.get('X-Bundle-Serve-Only') == '1'
    regenerated = False
    pdf_file = None
    if serve_only:
        if not bundle.pdf_is_current():
            return JsonResponse({'error': 'PDF is not ready yet.'}, status=409)
        try:
            pdf_file = bundle.final_pdf.open('rb')
        except Exception:
            return JsonResponse({'error': 'PDF is not ready yet.'}, status=409)
    else:
        success, error, regenerated = _ensure_bundle_final_pdf(
            bundle, user=request.user)
        if not success:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'error': error}, status=400)
            messages.error(request, error)
            return redirect('bundle_edit', bundle_id=bundle.id)

    if pdf_file is None:
        pdf_file = bundle.final_pdf.open('rb')

    response = FileResponse(
        pdf_file,
        content_type='application/pdf'
    )
    response['Content-Disposition'] = f'attachment; filename="{bundle.name}.pdf"'
    if regenerated:
        response['X-Bundle-Regenerated'] = '1'
    return response


def _parse_order_ids(raw_values):
    """Parse repeated form values, also tolerating legacy comma-joined payloads."""
    order_ids = []
    for raw_value in raw_values:
        for value in str(raw_value).split(','):
            value = value.strip()
            if value:
                order_ids.append(int(value))
    return order_ids


def _update_order_safely(queryset, ordered_ids):
    """Update unique order fields without colliding with existing order values."""
    with transaction.atomic():
        max_order = queryset.aggregate(max_order=Max('order'))[
            'max_order'] or 0
        temporary_offset = max_order + len(ordered_ids) + 1

        for index, object_id in enumerate(ordered_ids):
            queryset.filter(id=object_id).update(
                order=temporary_offset + index)

        for index, object_id in enumerate(ordered_ids):
            queryset.filter(id=object_id).update(order=index + 1)


def _section_ordered_document_ids(section):
    return [document.id for document in section.ordered_documents()]


def _open_document_pdf(document, cache=None):
    if cache is not None:
        return open(cache.local_path(document), 'rb')
    return document.file.open('rb')


def _generate_bundle_pdf(bundle, cache=None, progress_callback=None, documents_info=None):
    """Generate the final PDF bundle with index and pagination"""
    from PyPDF2 import PdfWriter, PdfReader
    from PyPDF2.generic import AnnotationBuilder

    if documents_info is None:
        documents_info = _collect_bundle_documents(bundle, cache=cache)
    index_page_count = _estimate_index_page_count(documents_info, bundle)

    current_document_page = index_page_count + 1
    for doc_info in documents_info:
        doc_info['page_start'] = current_document_page
        doc_info['page_end'] = current_document_page + \
            doc_info['page_count'] - 1
        current_document_page = doc_info['page_end'] + 1

    if progress_callback:
        progress_callback(20, 'Generating index page...')
    index_pdf_content, index_links = _generate_index_pdf(
        bundle, documents_info)

    # Create final PDF writer
    writer = PdfWriter()
    page_number = 1

    # Add index page
    index_reader = PdfReader(BytesIO(index_pdf_content))
    for page in index_reader.pages:
        page_with_number = _add_page_number(page, page_number)
        writer.add_page(page_with_number)
        page_number += 1

    # Add documents with page numbers
    doc_total = len(documents_info)
    for doc_index, doc_info in enumerate(documents_info):
        try:
            if progress_callback and doc_total:
                percent = 30 + int(55 * (doc_index + 1) / doc_total)
                label = doc_info['description'][:48] or 'document'
                progress_callback(
                    percent,
                    f'Adding document {doc_index + 1} of {doc_total}: {label}...',
                )
            with _open_document_pdf(doc_info['document'], cache=cache) as doc_file:
                doc_reader = PdfReader(doc_file)
                for page_index in doc_info['page_indices']:
                    page = doc_reader.pages[page_index]
                    # Add page number to each page
                    page_with_number = _add_page_number(page, page_number)
                    writer.add_page(page_with_number)
                    page_number += 1
        except Exception as e:
            logger.exception(
                "Error adding bundle document pages for document %s: %s", doc_info['document'].id, e)

    if progress_callback:
        progress_callback(88, 'Adding bookmarks and links...')

    for link in index_links:
        target_page_index = link['target_page_index']
        if target_page_index < len(writer.pages):
            writer.add_annotation(
                link['source_page_index'],
                AnnotationBuilder.link(
                    rect=link['rect'],
                    target_page_index=target_page_index,
                ),
            )

    _add_bundle_pdf_bookmarks(writer, documents_info)

    if progress_callback:
        progress_callback(92, 'Finalising PDF...')

    for doc_info in documents_info:
        if doc_info['page_start'] > len(writer.pages):
            doc_info['page_start'] = None
            doc_info['page_end'] = None

    for doc_info in documents_info:
        BundleDocument.objects.filter(id=doc_info['document'].id).update(
            page_start=doc_info['page_start'],
            page_end=doc_info['page_end'],
        )

    # Write to bytes
    output_buffer = BytesIO()
    writer.write(output_buffer)
    pdf_content = output_buffer.getvalue()
    output_buffer.close()

    return pdf_content


def _add_bundle_pdf_bookmarks(writer, documents_info):
    """Add PDF outline/bookmarks for sidebar navigation in PDF viewers."""
    writer.add_outline_item('Index', 0)

    current_section = None
    section_parent = None
    serial_number = 1

    for doc_info in documents_info:
        page_start = doc_info.get('page_start')
        if not page_start:
            continue

        target_page = page_start - 1
        if target_page < 0 or target_page >= len(writer.pages):
            continue

        if doc_info['section'] != current_section:
            current_section = doc_info['section']
            section_parent = writer.add_outline_item(
                f'{current_section} (p. {page_start})',
                target_page,
            )

        label = doc_info['description']
        if doc_info['date']:
            label = f'{label} ({doc_info["date"]})'

        writer.add_outline_item(
            f'{serial_number}. {label} (p. {page_start})',
            target_page,
            parent=section_parent,
        )
        serial_number += 1


def _get_pdf_page_count(document, cache=None):
    from PyPDF2 import PdfReader

    with _open_document_pdf(document, cache=cache) as pdf_file:
        reader = PdfReader(pdf_file)
        return len(reader.pages)


def _normalise_document_page_order(document, page_count):
    if not document.page_order:
        return list(range(1, page_count + 1))

    page_order = []
    seen_pages = set()
    for page_number in document.page_order:
        try:
            page_number = int(page_number)
        except (TypeError, ValueError):
            continue

        if 1 <= page_number <= page_count and page_number not in seen_pages:
            page_order.append(page_number)
            seen_pages.add(page_number)

    return page_order or list(range(1, page_count + 1))


def _format_page_order_summary(page_order, page_count):
    if page_order == list(range(1, page_count + 1)):
        return f"All {page_count} page{'s' if page_count != 1 else ''}"
    return "Pages " + ", ".join(str(page_number) for page_number in page_order)


def _get_bundle_document_page_choices(document):
    try:
        page_count = _get_pdf_page_count(document)
    except Exception as e:
        logger.exception(
            "Error reading page count for bundle document %s: %s", document.id, e)
        document.page_count = 0
        return []

    page_order = _normalise_document_page_order(document, page_count)
    included_pages = set(page_order)
    ordered_pages = page_order + [
        page_number
        for page_number in range(1, page_count + 1)
        if page_number not in included_pages
    ]

    document.page_count = page_count
    return [
        {
            'number': page_number,
            'included': page_number in included_pages,
        }
        for page_number in ordered_pages
    ]


def _collect_bundle_documents(bundle, cache=None):
    from PyPDF2 import PdfReader

    documents_info = []
    for section in bundle.sections.all().order_by('order'):
        for document in section.documents.all().order_by('order'):
            try:
                with _open_document_pdf(document, cache=cache) as pdf_file:
                    reader = PdfReader(pdf_file)
                    page_order = _normalise_document_page_order(
                        document, len(reader.pages))
                    page_indices = [page_number -
                                    1 for page_number in page_order]

                if not page_indices:
                    continue

                documents_info.append({
                    'document': document,
                    'section': section.heading,
                    'description': document.description,
                    'date': document.date.strftime('%d/%m/%Y') if document.date else '',
                    'page_count': len(page_indices),
                    'page_indices': page_indices,
                    'page_start': None,
                    'page_end': None,
                })
            except Exception as e:
                logger.exception(
                    "Error processing bundle document %s: %s", document.id, e)
    return documents_info


def _wrap_index_text_lines(canvas, text, max_width, font_name=None, font_size=12):
    if font_name is None:
        font_name = _BUNDLE_INDEX_SERIF_FONT
    text = str(text or '').strip()
    if not text:
        return ['']

    words = text.split()
    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = f'{current} {word}'
        if canvas.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)

    wrapped = []
    for line in lines:
        if canvas.stringWidth(line, font_name, font_size) <= max_width:
            wrapped.append(line)
            continue
        chunk = line
        while chunk:
            while chunk and canvas.stringWidth(chunk, font_name, font_size) > max_width:
                chunk = chunk[:-1]
            wrapped.append(chunk)
            line = line[len(chunk):]
            chunk = line
    return wrapped or ['']


def _estimate_index_text_line_count(text, max_width):
    text = str(text or '').strip()
    if not text:
        return 1

    chars_per_line = max(18, int(max_width / 6))
    lines = 1
    current_len = 0
    for word in text.split():
        word_len = len(word)
        extra = word_len + (1 if current_len else 0)
        if current_len and current_len + extra > chars_per_line:
            lines += 1
            current_len = word_len
        else:
            current_len += extra
    return max(1, lines)


def _estimate_index_page_count(documents_info, bundle):
    from reportlab.lib.pagesizes import A4

    page_width, _page_height = A4
    margin_x = 42
    show_date_column = _bundle_index_show_date_column(documents_info)
    desc_max_width = page_width - \
        (2 * margin_x) - (230 if show_date_column else 80)
    section_max_width = page_width - (2 * margin_x) - 16
    row_height = 28
    rows_per_page = 24
    first_page_rows = rows_per_page - _index_header_row_offset(bundle)
    page_count = 1
    rows_used = 0
    current_section = None

    for doc_info in documents_info:
        rows_needed = _estimate_index_text_line_count(
            doc_info['description'], desc_max_width)
        if current_section != doc_info['section']:
            rows_needed += _estimate_index_text_line_count(
                doc_info['section'], section_max_width)
            current_section = doc_info['section']

        row_units = max(
            1, int((rows_needed * 14 + row_height - 1) / row_height))
        page_capacity = first_page_rows if page_count == 1 else rows_per_page
        if rows_used and rows_used + row_units > page_capacity:
            page_count += 1
            rows_used = 0

        rows_used += row_units

    return page_count


def _index_header_row_offset(bundle):
    if not bundle.is_court_bundle:
        return 3
    claimants, defendants = bundle.court_parties_by_side()
    rows = 1 + max(len(claimants), 1) + 1 + max(len(defendants), 1) + 5
    if len(bundle.court_name or '') > 42:
        rows += 1
    if bundle.hearing_line or bundle.conference_line:
        rows += 1
    return rows


def _court_heading_text(value):
    return str(value or '').strip().upper()


_BUNDLE_INDEX_RULE_GREY = None
_BUNDLE_PAGE_NUMBER_FONT_SIZE = 20
_BUNDLE_PAGE_NUMBER_RIGHT_MARGIN = 18
_BUNDLE_PAGE_NUMBER_BOTTOM_MARGIN = 16
_BUNDLE_INDEX_SERIF_FONT = 'Times-Roman'
_BUNDLE_INDEX_SERIF_FONT_BOLD = 'Times-Bold'
_SEMIBOLD_OFFSET = 0.3


def _bundle_index_rule_color():
    from reportlab.lib import colors

    global _BUNDLE_INDEX_RULE_GREY
    if _BUNDLE_INDEX_RULE_GREY is None:
        _BUNDLE_INDEX_RULE_GREY = colors.grey
    return _BUNDLE_INDEX_RULE_GREY


def _semibold_string_width(canvas, text, font_size):
    return canvas.stringWidth(text, _BUNDLE_INDEX_SERIF_FONT, font_size) + _SEMIBOLD_OFFSET


def _draw_semibold_text(canvas, text, x, y, font_size, align='left'):
    canvas.setFont(_BUNDLE_INDEX_SERIF_FONT, font_size)
    if align == 'left':
        canvas.drawString(x, y, text)
        canvas.drawString(x + _SEMIBOLD_OFFSET, y, text)
    elif align == 'center':
        canvas.drawCentredString(x, y, text)
        canvas.drawCentredString(x + _SEMIBOLD_OFFSET, y, text)
    elif align == 'right':
        canvas.drawRightString(x, y, text)
        canvas.drawRightString(x + _SEMIBOLD_OFFSET, y, text)


def _draw_index_table_header(index_canvas, margin_x, page_width, table_top, show_date_column, date_col_x, page_col_x):
    index_font_size = 12
    index_canvas.setStrokeColor(_bundle_index_rule_color())
    _draw_semibold_text(index_canvas, 'No.', margin_x,
                        table_top, index_font_size)
    _draw_semibold_text(
        index_canvas, 'Description', margin_x + 48, table_top, index_font_size)
    if show_date_column:
        _draw_semibold_text(index_canvas, 'Date', date_col_x,
                            table_top, index_font_size)
    _draw_semibold_text(
        index_canvas, 'Page', page_col_x, table_top, index_font_size, align='right')
    index_canvas.line(margin_x, table_top - 10,
                      page_width - margin_x, table_top - 10)


def _draw_standard_index_header(index_canvas, bundle, margin_x, page_width, top_y):
    _draw_semibold_text(index_canvas, 'Index', margin_x, top_y, 18)
    index_canvas.setFont(_BUNDLE_INDEX_SERIF_FONT, 11)
    index_canvas.drawString(margin_x, top_y - 24, bundle.name)
    table_top = top_y - 52
    index_canvas.setStrokeColor(_bundle_index_rule_color())
    index_canvas.line(margin_x, table_top + 18,
                      page_width - margin_x, table_top + 18)
    return table_top


def _wrap_court_heading_lines(canvas, text, max_width, font_name=None, font_size=16):
    if font_name is None:
        font_name = _BUNDLE_INDEX_SERIF_FONT
    text = str(text or '').strip()
    if not text:
        return []

    words = text.split()
    if not words:
        return [text]

    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = f'{current} {word}'
        if canvas.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)

    wrapped = []
    for line in lines:
        if canvas.stringWidth(line, font_name, font_size) <= max_width:
            wrapped.append(line)
            continue
        chunk = line
        while chunk and canvas.stringWidth(chunk, font_name, font_size) > max_width:
            chunk = chunk[:-1]
        if chunk and len(chunk) < len(line):
            chunk = chunk[:-3].rstrip() + '...'
        wrapped.append(chunk or line[:1])
    return wrapped


def _draw_court_index_header(index_canvas, bundle, margin_x, page_width, top_y):
    claimants, defendants = bundle.court_parties_by_side()
    if not claimants:
        claimants = [{'name': '', 'role': 'Claimant 1'}]
    if not defendants:
        defendants = [{'name': '', 'role': 'Defendant 1'}]

    y = top_y
    case_label = (
        'CASE NO.'
        if bundle.case_number_type == Bundle.CASE_NUMBER_CASE
        else 'CLAIM NO.'
    )
    row1_font_size = 16
    party_font_size = 14
    index_title_font_size = 14
    footer_font_size = 10
    row1_leading = 20
    party_row_leading = 30

    case_text = (
        f'{case_label} {bundle.case_number.upper()}'
        if bundle.case_number else ''
    )
    case_width = (
        _semibold_string_width(index_canvas, case_text, row1_font_size) + 16
        if case_text else 0
    )
    max_court_width = page_width - (2 * margin_x) - case_width
    court_lines = _wrap_court_heading_lines(
        index_canvas,
        _court_heading_text(bundle.court_name),
        max(max_court_width, 120),
        font_size=row1_font_size,
    )
    if not court_lines and bundle.court_name:
        court_lines = [_court_heading_text(bundle.court_name)]

    for line_index, line in enumerate(court_lines):
        _draw_semibold_text(
            index_canvas,
            line,
            margin_x,
            y - (line_index * row1_leading),
            row1_font_size,
        )
    if case_text:
        _draw_semibold_text(
            index_canvas,
            case_text,
            page_width - margin_x,
            y,
            row1_font_size,
            align='right',
        )
    y -= row1_leading * max(len(court_lines), 1) + 14

    for party in claimants:
        if party.get('name'):
            _draw_semibold_text(
                index_canvas,
                _court_heading_text(party['name']),
                page_width / 2,
                y,
                party_font_size,
                align='center',
            )
        if party.get('role'):
            _draw_semibold_text(
                index_canvas,
                _court_heading_text(party['role']),
                page_width - margin_x,
                y,
                party_font_size,
                align='right',
            )
        y -= party_row_leading

    _draw_semibold_text(
        index_canvas, '-V-', page_width / 2, y, party_font_size, align='center')
    y -= party_row_leading

    for party in defendants:
        if party.get('name'):
            _draw_semibold_text(
                index_canvas,
                _court_heading_text(party['name']),
                page_width / 2,
                y,
                party_font_size,
                align='center',
            )
        if party.get('role'):
            _draw_semibold_text(
                index_canvas,
                _court_heading_text(party['role']),
                page_width - margin_x,
                y,
                party_font_size,
                align='right',
            )
        y -= party_row_leading

    y -= 8
    index_canvas.setStrokeColor(_bundle_index_rule_color())
    index_canvas.line(margin_x, y, page_width - margin_x, y)
    y -= 20

    _draw_semibold_text(
        index_canvas,
        _court_heading_text(bundle.index_title or 'Index to the Bundle'),
        page_width / 2,
        y,
        index_title_font_size,
        align='center',
    )
    y -= 10

    footer_parts = []
    if bundle.hearing_line.strip():
        footer_parts.append(_court_heading_text(bundle.hearing_line))
    if bundle.conference_line.strip():
        footer_parts.append(_court_heading_text(bundle.conference_line))
    if footer_parts:
        _draw_semibold_text(
            index_canvas,
            '   '.join(footer_parts),
            page_width / 2,
            y,
            footer_font_size,
            align='center',
        )
        y -= 12

    y -= 6
    index_canvas.line(margin_x, y, page_width - margin_x, y)
    return y - 24


def _bundle_index_show_date_column(documents_info):
    return any(doc_info['date'] for doc_info in documents_info)


def _index_font_vertical_metrics(font_name, font_size):
    from reportlab.pdfbase import pdfmetrics

    ascender = pdfmetrics.getAscent(font_name) / 1000.0 * font_size
    descender = abs(pdfmetrics.getDescent(font_name)) / 1000.0 * font_size
    return ascender, descender


def _index_row_min_height(line_count, line_leading, font_size=12, font_name=None, min_height=28, padding=8):
    if font_name is None:
        font_name = _BUNDLE_INDEX_SERIF_FONT
    line_count = max(1, line_count)
    ascender, descender = _index_font_vertical_metrics(font_name, font_size)
    visual_height = ascender + (line_count - 1) * line_leading + descender
    return max(min_height, int(visual_height + padding))


def _index_row_text_baselines(row_top, row_height, line_count, line_leading, font_size=12, font_name=None):
    if font_name is None:
        font_name = _BUNDLE_INDEX_SERIF_FONT
    line_count = max(1, line_count)
    ascender, descender = _index_font_vertical_metrics(font_name, font_size)
    baseline_span = (line_count - 1) * line_leading
    visual_height = ascender + baseline_span + descender
    row_center = row_top - (row_height / 2)
    first_baseline = row_center + (visual_height / 2) - ascender
    single_baseline = row_center - ((ascender - descender) / 2)
    return first_baseline, single_baseline


def _generate_index_pdf(bundle, documents_info):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    show_date_column = _bundle_index_show_date_column(documents_info)
    page_width, page_height = A4
    margin_x = 42
    top_y = page_height - 42
    row_height = 28
    index_font_size = 12
    bottom_y = 50
    page_col_x = page_width - margin_x
    date_col_x = page_width - margin_x - 150
    desc_max_width = page_width - \
        (2 * margin_x) - (230 if show_date_column else 80)
    section_max_width = page_width - (2 * margin_x) - 16
    desc_line_leading = 14

    buffer = BytesIO()
    index_canvas = canvas.Canvas(buffer, pagesize=A4)
    links = []
    page_index = 0
    serial_number = 1
    current_section = None

    def draw_page_header(full_header=True):
        nonlocal y
        if full_header and bundle.is_court_bundle:
            table_top = _draw_court_index_header(
                index_canvas, bundle, margin_x, page_width, top_y)
        elif full_header:
            table_top = _draw_standard_index_header(
                index_canvas, bundle, margin_x, page_width, top_y)
        else:
            table_top = top_y - 24
            index_canvas.setStrokeColor(_bundle_index_rule_color())
            index_canvas.line(margin_x, table_top + 18,
                              page_width - margin_x, table_top + 18)

        _draw_index_table_header(
            index_canvas, margin_x, page_width, table_top,
            show_date_column, date_col_x, page_col_x)
        y = table_top - row_height

    def new_page():
        nonlocal page_index
        index_canvas.showPage()
        page_index += 1
        draw_page_header(full_header=False)

    draw_page_header(full_header=True)

    for doc_info in documents_info:
        desc_lines = _wrap_index_text_lines(
            index_canvas,
            doc_info['description'],
            desc_max_width,
            font_size=index_font_size,
        )
        doc_row_height = _index_row_min_height(
            len(desc_lines),
            desc_line_leading,
            index_font_size,
        )

        section_lines = []
        section_row_height = row_height
        if current_section != doc_info['section']:
            section_lines = _wrap_index_text_lines(
                index_canvas,
                doc_info['section'],
                section_max_width,
                font_name=_BUNDLE_INDEX_SERIF_FONT,
                font_size=index_font_size,
            )
            section_row_height = _index_row_min_height(
                len(section_lines),
                desc_line_leading,
                index_font_size,
                font_name=_BUNDLE_INDEX_SERIF_FONT_BOLD,
            )

        needed_height = (
            section_row_height + doc_row_height
            if current_section != doc_info['section']
            else doc_row_height
        )
        if y - needed_height < bottom_y:
            new_page()

        if current_section != doc_info['section']:
            current_section = doc_info['section']
            row_top = y
            row_bottom = y - section_row_height
            section_first_baseline, _section_single_baseline = _index_row_text_baselines(
                row_top,
                section_row_height,
                len(section_lines),
                desc_line_leading,
                index_font_size,
                font_name=_BUNDLE_INDEX_SERIF_FONT_BOLD,
            )
            index_canvas.setFillColor(colors.whitesmoke)
            index_canvas.rect(
                margin_x,
                row_bottom,
                page_width - (2 * margin_x),
                section_row_height,
                fill=True,
                stroke=False,
            )
            index_canvas.setFillColor(colors.black)
            index_canvas.setFont(_BUNDLE_INDEX_SERIF_FONT_BOLD, index_font_size)
            for line_index, line in enumerate(section_lines):
                index_canvas.drawString(
                    margin_x + 8,
                    section_first_baseline - (line_index * desc_line_leading),
                    line,
                )
            y -= section_row_height

        page_range = (
            str(doc_info['page_start'])
            if doc_info['page_start'] == doc_info['page_end']
            else f"{doc_info['page_start']}-{doc_info['page_end']}"
        )
        row_top = y
        row_bottom = y - doc_row_height
        desc_first_baseline, single_baseline = _index_row_text_baselines(
            row_top,
            doc_row_height,
            len(desc_lines),
            desc_line_leading,
            index_font_size,
        )

        index_canvas.setStrokeColor(_bundle_index_rule_color())
        index_canvas.line(margin_x, row_bottom,
                          page_width - margin_x, row_bottom)
        index_canvas.setFillColor(colors.black)
        index_canvas.setFont(_BUNDLE_INDEX_SERIF_FONT, index_font_size)
        index_canvas.drawString(margin_x, single_baseline, str(serial_number))
        for line_index, line in enumerate(desc_lines):
            index_canvas.drawString(
                margin_x + 48,
                desc_first_baseline - (line_index * desc_line_leading),
                line,
            )
        if show_date_column:
            index_canvas.drawString(
                date_col_x, single_baseline, doc_info['date'])
        index_canvas.drawRightString(page_col_x, single_baseline, page_range)
        index_canvas.setFillColor(colors.black)

        links.append({
            'source_page_index': page_index,
            'target_page_index': doc_info['page_start'] - 1,
            'rect': (margin_x, row_bottom, page_width - margin_x, row_top),
        })

        serial_number += 1
        y -= doc_row_height

    index_canvas.save()
    return buffer.getvalue(), links


def _add_page_number(page, page_number):
    """Add page number to bottom-right of PDF page"""
    from PyPDF2 import PdfReader
    from reportlab.pdfgen import canvas

    width = float(page.mediabox.width)
    height = float(page.mediabox.height)

    # Create an in-memory PDF with just the page number.
    number_buffer = BytesIO()
    temp_canvas = canvas.Canvas(number_buffer, pagesize=(width, height))

    temp_canvas.setFont(_BUNDLE_INDEX_SERIF_FONT_BOLD,
                        _BUNDLE_PAGE_NUMBER_FONT_SIZE)
    temp_canvas.setFillColorRGB(0, 0, 0)

    x_position = width - _BUNDLE_PAGE_NUMBER_RIGHT_MARGIN
    y_position = _BUNDLE_PAGE_NUMBER_BOTTOM_MARGIN

    temp_canvas.drawRightString(x_position, y_position, str(page_number))
    temp_canvas.save()
    number_buffer.seek(0)

    number_reader = PdfReader(number_buffer)
    page.merge_page(number_reader.pages[0])

    return page


@login_required
def bundle_delete(request, bundle_id):
    """Delete a bundle"""
    bundle = get_object_or_404(Bundle, id=bundle_id, created_by=request.user)

    if request.method == 'POST':
        # Delete all associated files
        for section in bundle.sections.all():
            for document in section.documents.all():
                if document.file:
                    try:
                        default_storage.delete(document.file.name)
                    except:
                        pass

        # Delete final PDF if exists
        if bundle.final_pdf:
            try:
                default_storage.delete(bundle.final_pdf.name)
            except:
                pass

        bundle_name = bundle.name
        bundle_file_number = bundle.file_number.file_number if bundle.file_number else None
        if bundle.file_number_id:
            log_deleted_on_parent(
                request.user,
                bundle.file_number,
                'bundle',
                bundle_name,
            )
        bundle.delete()

        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or 'application/json' in request.headers.get('Accept', ''):
            return JsonResponse({'success': True})

        messages.success(
            request, f'Bundle "{bundle_name}" deleted successfully.')
        if bundle_file_number:
            return redirect('home', file_number=bundle_file_number)
        return redirect('user_dashboard')

    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required
@require_POST
def update_comment(request):
    """Update comment for a file via AJAX"""
    try:
        data = json.loads(request.body)
        file_number = data.get('file_number')
        comment = data.get('comment', '')

        if not file_number:
            return JsonResponse({'success': False, 'error': 'File number is required'}, status=400)

        # Get the WIP object
        wip = get_object_or_404(WIP, file_number=file_number)

        # Store old value for modification tracking
        old_comment = wip.comments

        # Update the comment
        wip.comments = comment
        wip.save()

        # Create modification record
        changes = {
            'comments': {
                'old_value': old_comment,
                'new_value': comment
            }
        }
        create_modification(request.user, wip, changes)

        return JsonResponse({
            'success': True,
            'message': 'Comment updated successfully'
        })

    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)
