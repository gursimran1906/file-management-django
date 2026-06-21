"""Export firm policies to an editable Word (.docx) document.

Converts the Quill rich-text stored on each ``PolicyVersion`` into native Word
paragraphs and runs, so the output is fully editable rather than a flattened
PDF/image. The input is Quill 1.3.7 HTML (``version.content.html``) — the same
source ``download_policy_pdf`` renders for the single-policy PDF.

Supported formatting: headings (h1-h6), bold/italic/underline/strikethrough,
ordered/bullet lists (incl. ``ql-indent-N`` nesting), blockquotes, code blocks,
paragraph alignment (``ql-align-*``), and hyperlinks. Inline images are not
carried across (a marker is inserted so nothing is silently dropped).
"""
import io
import re
import zipfile

from bs4 import BeautifulSoup, NavigableString, Tag
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

HEADING_LEVELS = {f"h{i}": i for i in range(1, 7)}
BLOCK_TAGS = {"p", "div", "blockquote", "pre", "ul", "ol"} | set(HEADING_LEVELS)
HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _add_hyperlink(paragraph, url, text, bold=False, italic=False):
    """Append a real, clickable hyperlink run to *paragraph* (python-docx has
    no high-level API for this, so the run XML is built by hand)."""
    r_id = paragraph.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    run.append(rpr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _format_for(tag_name, fmt):
    """Return a copy of *fmt* with the formatting implied by *tag_name* set."""
    fmt = dict(fmt)
    if tag_name in ("strong", "b"):
        fmt["bold"] = True
    elif tag_name in ("em", "i"):
        fmt["italic"] = True
    elif tag_name == "u":
        fmt["underline"] = True
    elif tag_name in ("s", "strike", "del"):
        fmt["strike"] = True
    return fmt


def _render_inline(paragraph, node, fmt):
    """Walk inline content under *node*, appending formatted runs to *paragraph*."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                run.bold = fmt.get("bold", False)
                run.italic = fmt.get("italic", False)
                run.underline = fmt.get("underline", False)
                if fmt.get("strike"):
                    run.font.strike = True
            continue
        if not isinstance(child, Tag):
            continue

        name = child.name.lower()
        if name == "br":
            paragraph.add_run().add_break()
        elif name == "a":
            href = (child.get("href") or "").strip()
            text = child.get_text()
            if href and text:
                _add_hyperlink(
                    paragraph, href, text,
                    bold=fmt.get("bold", False), italic=fmt.get("italic", False),
                )
            else:
                _render_inline(paragraph, child, fmt)
        elif name == "img":
            paragraph.add_run("[image omitted]").italic = True
        else:
            _render_inline(paragraph, child, _format_for(name, fmt))


def _indent_level(tag):
    for cls in tag.get("class") or []:
        match = re.match(r"ql-indent-(\d+)", cls)
        if match:
            return int(match.group(1))
    return 0


def _apply_alignment(paragraph, tag):
    classes = tag.get("class") or []
    if "ql-align-center" in classes:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif "ql-align-right" in classes:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif "ql-align-justify" in classes:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_list_item(document, li, base_style):
    """Add one ``<li>`` as a list paragraph, honouring ql-indent nesting."""
    level = _indent_level(li)
    style = base_style
    if level:
        # Word ships built-in "List Bullet 2/3" / "List Number 2/3" styles.
        style = f"{base_style} {min(level + 1, 3)}"
    try:
        paragraph = document.add_paragraph(style=style)
    except KeyError:
        paragraph = document.add_paragraph(style=base_style)
    _render_inline(paragraph, li, {})


def _add_block(document, tag):
    name = tag.name.lower()
    if name in HEADING_LEVELS:
        paragraph = document.add_paragraph(style=f"Heading {HEADING_LEVELS[name]}")
        _apply_alignment(paragraph, tag)
        _render_inline(paragraph, tag, {})
    elif name == "blockquote":
        paragraph = document.add_paragraph(style="Quote")
        _render_inline(paragraph, tag, {})
    elif name == "pre":
        paragraph = document.add_paragraph()
        run = paragraph.add_run(tag.get_text())
        run.font.name = "Courier New"
    elif name in ("ul", "ol"):
        base_style = "List Bullet" if name == "ul" else "List Number"
        for li in tag.find_all("li", recursive=False):
            _add_list_item(document, li, base_style)
    else:  # <p>, <div>, anything else block-ish
        paragraph = document.add_paragraph()
        _apply_alignment(paragraph, tag)
        level = _indent_level(tag)
        if level:
            paragraph.paragraph_format.left_indent = Pt(18 * level)
        _render_inline(paragraph, tag, {})


def html_to_docx(document, html):
    """Append the converted Quill *html* to an open python-docx *document*."""
    soup = BeautifulSoup(html or "", "html.parser")
    root = soup.body or soup
    for child in root.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                document.add_paragraph(text)
        elif isinstance(child, Tag):
            if child.name.lower() in BLOCK_TAGS:
                _add_block(document, child)
            else:  # stray top-level inline content (e.g. bare <strong>...)
                _render_inline(document.add_paragraph(), child, {})


def _add_policy(document, policy, version, heading_level=0):
    """Append one policy (title, metadata line, content) to *document*."""
    document.add_heading(policy.description, level=heading_level)

    if version is None:
        document.add_paragraph("No versions available.")
        return

    meta = document.add_paragraph()
    timestamp = version.timestamp
    if timestamp is not None and timezone.is_aware(timestamp):
        timestamp = timezone.localtime(timestamp)
    bits = [f"Version {version.version_number}"]
    if timestamp is not None:
        bits.append(f"Updated {timestamp.strftime('%d/%m/%Y')}")
    if version.changes_by:
        bits.append(f"By {version.changes_by}")
    meta.add_run("  ·  ".join(bits)).italic = True

    html_to_docx(document, version.content.html)


def build_policy_docx(policy, version):
    """Build a standalone .docx for a single policy and return the bytes."""
    document = Document()
    _add_policy(document, policy, version)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _safe_filename(name):
    """Turn a policy description into a filesystem-safe base name."""
    cleaned = re.sub(r"[^\w\s-]", "", name or "").strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned or "policy"


def build_policies_zip(policies):
    """Build one .docx per policy and bundle them into a .zip; return the bytes.

    *policies* is an iterable of ``(Policy, PolicyVersion | None)`` tuples. Each
    file is named after the policy description; duplicate names get a numeric
    suffix so nothing is overwritten inside the archive.
    """
    buffer = io.BytesIO()
    used = {}
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for policy, version in policies:
            base = _safe_filename(policy.description)
            used[base] = used.get(base, 0) + 1
            suffix = "" if used[base] == 1 else f"_{used[base]}"
            archive.writestr(f"{base}{suffix}.docx", build_policy_docx(policy, version))
    return buffer.getvalue()
